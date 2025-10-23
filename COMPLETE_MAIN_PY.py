from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import hashlib
import tempfile
import os
from docx import Document

app = Flask(__name__)

# FIXED CORS - Allows GitHub Pages URLs
CORS(app, 
     origins=[
         "https://tuyisingize750.github.io",
         "https://tuyisingize750.github.io/rtb-document-planner",
         "https://schemesession.netlify.app",
         "http://localhost:5173",
         "http://localhost:8000"
     ],
     methods=["GET", "POST", "OPTIONS", "PUT"],
     allow_headers=["Content-Type", "Authorization"],
     supports_credentials=False)

@app.after_request
def after_request(response):
    origin = request.headers.get('Origin')
    allowed_origins = [
        "https://tuyisingize750.github.io",
        "https://tuyisingize750.github.io/rtb-document-planner", 
        "https://schemesession.netlify.app",
        "http://localhost:5173",
        "http://localhost:8000"
    ]
    if origin in allowed_origins:
        response.headers.add('Access-Control-Allow-Origin', origin)
    response.headers.add('Access-Control-Allow-Headers', 'Content-Type,Authorization')
    response.headers.add('Access-Control-Allow-Methods', 'GET,POST,OPTIONS,PUT')
    return response

# Enhanced user storage with registration support
users = {
    "+250796014803": {
        "name": "Test Teacher",
        "password": hashlib.sha256("teacher123".encode()).hexdigest(),
        "phone": "+250796014803",
        "role": "user",
        "institution": "Test School"
    },
    "+250789751597": {
        "name": "Administrator", 
        "password": hashlib.sha256("admin123".encode()).hexdigest(),
        "phone": "+250789751597",
        "role": "admin",
        "institution": "RTB"
    }
}

documents = {}
doc_counter = 1

@app.route('/', methods=['GET', 'OPTIONS'])
def home():
    if request.method == 'OPTIONS':
        return '', 204
    return jsonify({
        "message": "RTB Document Planner API",
        "status": "online",
        "version": "2.0",
        "cors": "enabled",
        "users_count": len(users)
    })

# FIXED REGISTRATION ENDPOINT
@app.route('/users/register', methods=['POST', 'OPTIONS'])
def register():
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        data = request.get_json()
        phone = data.get('phone')
        password = data.get('password')
        name = data.get('name')
        institution = data.get('institution', '')
        
        if not all([phone, password, name]):
            return jsonify({"detail": "Phone, password, and name are required"}), 400
        
        if phone in users:
            return jsonify({"detail": "Phone number already registered"}), 400
        
        # Add new user
        users[phone] = {
            "name": name,
            "password": hashlib.sha256(password.encode()).hexdigest(),
            "phone": phone,
            "role": "user",
            "institution": institution
        }
        
        return jsonify({"message": "User registered successfully"}), 201
    except Exception as e:
        return jsonify({"detail": "Registration failed"}), 500

@app.route('/users/login', methods=['POST', 'OPTIONS'])
def login():
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        data = request.get_json()
        phone = data.get('phone')
        password = data.get('password')
        
        if not all([phone, password]):
            return jsonify({"detail": "Phone and password are required"}), 400
        
        if phone in users:
            user = users[phone]
            if hashlib.sha256(password.encode()).hexdigest() == user['password']:
                return jsonify({
                    "message": "Login successful",
                    "user": {
                        "name": user['name'],
                        "phone": user['phone'],
                        "role": user['role'],
                        "institution": user.get('institution', ''),
                        "is_premium": False,
                        "session_plans_limit": 2,
                        "schemes_limit": 2,
                        "session_plans_downloaded": 0,
                        "schemes_downloaded": 0
                    }
                }), 200
        
        return jsonify({"detail": "Invalid credentials"}), 401
    except Exception as e:
        return jsonify({"detail": "Login failed"}), 500

@app.route('/session-plans', methods=['POST', 'OPTIONS'])
def create_session_plan():
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        global doc_counter
        doc_id = doc_counter
        doc_counter += 1
        
        data = request.get_json()
        documents[doc_id] = {
            'type': 'session_plan',
            'data': data
        }
        
        return jsonify({"id": doc_id, "message": "Session plan created successfully"}), 201
    except Exception as e:
        return jsonify({"detail": "Generation failed"}), 500

@app.route('/schemes', methods=['POST', 'OPTIONS'])
def create_scheme():
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        global doc_counter
        doc_id = doc_counter
        doc_counter += 1
        
        data = request.get_json()
        documents[doc_id] = {
            'type': 'scheme',
            'data': data
        }
        
        return jsonify({"id": doc_id, "message": "Scheme of work created successfully"}), 201
    except Exception as e:
        return jsonify({"detail": "Generation failed"}), 500

@app.route('/session-plans/<int:plan_id>/download', methods=['GET', 'OPTIONS'])
def download_session_plan(plan_id):
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        if plan_id not in documents:
            return jsonify({"detail": "Session plan not found"}), 404
        
        data = documents[plan_id]['data']
        
        # Create professional DOCX
        doc = Document()
        doc.add_heading('RTB SESSION PLAN', 0)
        
        # Add content
        doc.add_paragraph(f"Sector: {data.get('sector', 'N/A')}")
        doc.add_paragraph(f"Sub-Sector: {data.get('sub_sector', 'N/A')}")
        doc.add_paragraph(f"Trade: {data.get('trade', 'N/A')}")
        doc.add_paragraph(f"Qualification Title: {data.get('qualification_title', 'N/A')}")
        doc.add_paragraph(f"RQF Level: {data.get('rqf_level', 'N/A')}")
        doc.add_paragraph(f"Module Code & Title: {data.get('module_code_title', 'N/A')}")
        doc.add_paragraph(f"Term: {data.get('term', 'N/A')}")
        doc.add_paragraph(f"Week: {data.get('week', 'N/A')}")
        doc.add_paragraph(f"Date: {data.get('date', 'N/A')}")
        doc.add_paragraph(f"Trainer Name: {data.get('trainer_name', 'N/A')}")
        doc.add_paragraph(f"Class Name: {data.get('class_name', 'N/A')}")
        doc.add_paragraph(f"Number of Trainees: {data.get('number_of_trainees', 'N/A')}")
        doc.add_paragraph(f"Topic of Session: {data.get('topic_of_session', 'N/A')}")
        doc.add_paragraph(f"Duration: {data.get('duration', '40')} minutes")
        
        doc.add_heading('Learning Outcomes', level=1)
        doc.add_paragraph(data.get('learning_outcomes', 'N/A'))
        
        doc.add_heading('Indicative Contents', level=1)
        doc.add_paragraph(data.get('indicative_contents', 'N/A'))
        
        doc.add_heading('Objectives', level=1)
        doc.add_paragraph(data.get('objectives', 'N/A'))
        
        doc.add_heading('Facilitation Techniques', level=1)
        doc.add_paragraph(data.get('facilitation_techniques', 'N/A'))
        
        doc.add_heading('Learning Activities', level=1)
        doc.add_paragraph(data.get('learning_activities', 'N/A'))
        
        doc.add_heading('Resources', level=1)
        doc.add_paragraph(data.get('resources', 'N/A'))
        
        doc.add_heading('Assessment Details', level=1)
        doc.add_paragraph(data.get('assessment_details', 'N/A'))
        
        doc.add_heading('References', level=1)
        doc.add_paragraph(data.get('references', 'N/A'))
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        return send_file(temp_file.name, as_attachment=True, download_name=f"RTB_Session_Plan_{plan_id}.docx")
    except Exception as e:
        return jsonify({"detail": "Download failed"}), 500

@app.route('/schemes-of-work/<int:scheme_id>/download', methods=['GET', 'OPTIONS'])
def download_scheme(scheme_id):
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        if scheme_id not in documents:
            return jsonify({"detail": "Scheme of work not found"}), 404
        
        data = documents[scheme_id]['data']
        
        # Create professional DOCX
        doc = Document()
        doc.add_heading('RTB SCHEME OF WORK', 0)
        
        # Add content
        doc.add_paragraph(f"Province: {data.get('province', 'N/A')}")
        doc.add_paragraph(f"District: {data.get('district', 'N/A')}")
        doc.add_paragraph(f"Sector: {data.get('sector', 'N/A')}")
        doc.add_paragraph(f"School: {data.get('school', 'N/A')}")
        doc.add_paragraph(f"Department/Trade: {data.get('department_trade', 'N/A')}")
        doc.add_paragraph(f"Qualification Title: {data.get('qualification_title', 'N/A')}")
        doc.add_paragraph(f"RQF Level: {data.get('rqf_level', 'N/A')}")
        doc.add_paragraph(f"Module Code & Title: {data.get('module_code_title', 'N/A')}")
        doc.add_paragraph(f"School Year: {data.get('school_year', 'N/A')}")
        doc.add_paragraph(f"Terms: {data.get('terms', 'N/A')}")
        doc.add_paragraph(f"Module Hours: {data.get('module_hours', 'N/A')}")
        doc.add_paragraph(f"Number of Classes: {data.get('number_of_classes', 'N/A')}")
        doc.add_paragraph(f"Class Name: {data.get('class_name', 'N/A')}")
        doc.add_paragraph(f"Trainer Name: {data.get('trainer_name', 'N/A')}")
        
        # Term 1
        if data.get('term1_weeks'):
            doc.add_heading('TERM 1', level=1)
            doc.add_paragraph(f"Weeks: {data.get('term1_weeks', 'N/A')}")
            doc.add_paragraph(f"Learning Outcomes: {data.get('term1_learning_outcomes', 'N/A')}")
            doc.add_paragraph(f"Indicative Contents: {data.get('term1_indicative_contents', 'N/A')}")
            doc.add_paragraph(f"Duration: {data.get('term1_duration', 'N/A')}")
        
        # Term 2
        if data.get('term2_weeks'):
            doc.add_heading('TERM 2', level=1)
            doc.add_paragraph(f"Weeks: {data.get('term2_weeks', 'N/A')}")
            doc.add_paragraph(f"Learning Outcomes: {data.get('term2_learning_outcomes', 'N/A')}")
            doc.add_paragraph(f"Indicative Contents: {data.get('term2_indicative_contents', 'N/A')}")
            doc.add_paragraph(f"Duration: {data.get('term2_duration', 'N/A')}")
        
        # Term 3
        if data.get('term3_weeks'):
            doc.add_heading('TERM 3', level=1)
            doc.add_paragraph(f"Weeks: {data.get('term3_weeks', 'N/A')}")
            doc.add_paragraph(f"Learning Outcomes: {data.get('term3_learning_outcomes', 'N/A')}")
            doc.add_paragraph(f"Indicative Contents: {data.get('term3_indicative_contents', 'N/A')}")
            doc.add_paragraph(f"Duration: {data.get('term3_duration', 'N/A')}")
        
        doc.add_paragraph(f"DOS Name: {data.get('dos_name', 'N/A')}")
        doc.add_paragraph(f"Manager Name: {data.get('manager_name', 'N/A')}")
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        return send_file(temp_file.name, as_attachment=True, download_name=f"RTB_Scheme_of_Work_{scheme_id}.docx")
    except Exception as e:
        return jsonify({"detail": "Download failed"}), 500

# User limits endpoint for dashboard
@app.route('/user-limits/<phone>', methods=['GET', 'OPTIONS'])
def get_user_limits(phone):
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        if phone not in users:
            return jsonify({"detail": "User not found"}), 404
        
        return jsonify({
            "is_premium": False,
            "session_plans_limit": 2,
            "session_plans_downloaded": 0,
            "session_plans_remaining": 2,
            "schemes_limit": 2,
            "schemes_downloaded": 0,
            "schemes_remaining": 2
        }), 200
    except Exception as e:
        return jsonify({"detail": "Failed to get limits"}), 500

if __name__ == '__main__':
    app.run(debug=False)