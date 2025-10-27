"""Analyze Official RTB Templates from DOCS TO REFER TO folder"""
from docx import Document
from docx.shared import Pt, RGBColor
import os

def analyze_session_plan():
    doc = Document('../DOCS TO REFER TO/SESSION PLAN.docx')
    print("=" * 80)
    print("SESSION PLAN OFFICIAL TEMPLATE ANALYSIS")
    print("=" * 80)
    
    # Analyze table structure
    if doc.tables:
        table = doc.tables[0]
        print(f"\nTable: {len(table.rows)} rows × {len(table.columns)} columns")
        
        for i, row in enumerate(table.rows):
            print(f"\n--- Row {i} ---")
            for j, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    # Get cell formatting
                    if cell.paragraphs:
                        para = cell.paragraphs[0]
                        if para.runs:
                            run = para.runs[0]
                            font_name = run.font.name
                            font_size = run.font.size.pt if run.font.size else "N/A"
                            bold = run.font.bold
                            print(f"  Cell[{i},{j}]: '{text[:50]}' | Font: {font_name} {font_size}pt | Bold: {bold}")
                        else:
                            print(f"  Cell[{i},{j}]: '{text[:50]}'")

def analyze_scheme():
    doc = Document('../DOCS TO REFER TO/CSAPA 301 Scheme of work.docx')
    print("\n" + "=" * 80)
    print("SCHEME OF WORK OFFICIAL TEMPLATE ANALYSIS")
    print("=" * 80)
    
    # Analyze header paragraphs
    print("\n--- Header Information ---")
    for i, para in enumerate(doc.paragraphs[:15]):
        if para.text.strip():
            print(f"Para {i}: {para.text}")
    
    # Analyze tables
    print(f"\n--- Tables: {len(doc.tables)} ---")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}: {len(table.rows)} rows × {len(table.columns)} columns")
        
        # Analyze header row
        if table.rows:
            header_row = table.rows[0]
            print("Header cells:")
            for j, cell in enumerate(header_row.cells):
                text = cell.text.strip()
                if text:
                    print(f"  Col {j}: '{text}'")

if __name__ == "__main__":
    analyze_session_plan()
    analyze_scheme()
