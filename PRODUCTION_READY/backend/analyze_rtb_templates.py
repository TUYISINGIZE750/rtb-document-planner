from docx import Document
import os

def analyze_session_plan_template():
    """Analyze RTB session plan template structure"""
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_session_plan_template.docx')
    
    if not os.path.exists(template_path):
        print("Session plan template not found")
        return
    
    doc = Document(template_path)
    print("=== RTB SESSION PLAN TEMPLATE ANALYSIS ===")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row_idx, row in enumerate(table.rows):
            print(f"  Row {row_idx}:")
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"    Cell [{row_idx}][{cell_idx}]: {text[:100]}...")

def analyze_scheme_template():
    """Analyze RTB scheme template structure"""
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_scheme_template.docx')
    
    if not os.path.exists(template_path):
        print("Scheme template not found")
        return
    
    doc = Document(template_path)
    print("\n=== RTB SCHEME TEMPLATE ANALYSIS ===")
    
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        for row_idx, row in enumerate(table.rows):
            print(f"  Row {row_idx}:")
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"    Cell [{row_idx}][{cell_idx}]: {text[:100]}...")

if __name__ == "__main__":
    analyze_session_plan_template()
    analyze_scheme_template()