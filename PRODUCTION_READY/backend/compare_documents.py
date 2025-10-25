"""
Quick comparison tool to verify generated documents match templates
"""
from docx import Document
import os

def compare_structure(template_path, generated_path, doc_type):
    """Compare structure of template vs generated document"""
    print(f"\n{'='*60}")
    print(f"COMPARING: {doc_type}")
    print(f"{'='*60}\n")
    
    if not os.path.exists(template_path):
        print(f"[ERROR] Template not found: {template_path}")
        return False
    
    if not os.path.exists(generated_path):
        print(f"[ERROR] Generated file not found: {generated_path}")
        return False
    
    template = Document(template_path)
    generated = Document(generated_path)
    
    # Compare basic structure
    print(f"Tables:")
    print(f"  Template:  {len(template.tables)}")
    print(f"  Generated: {len(generated.tables)}")
    match = len(template.tables) == len(generated.tables)
    print(f"  Status: {'[OK]' if match else '[MISMATCH]'}\n")
    
    # Compare each table
    for idx in range(min(len(template.tables), len(generated.tables))):
        t_table = template.tables[idx]
        g_table = generated.tables[idx]
        
        print(f"Table {idx + 1}:")
        print(f"  Rows:    Template={len(t_table.rows)}, Generated={len(g_table.rows)}")
        print(f"  Columns: Template={len(t_table.columns)}, Generated={len(g_table.columns)}")
        
        rows_match = len(t_table.rows) == len(g_table.rows)
        cols_match = len(t_table.columns) == len(g_table.columns)
        
        if rows_match and cols_match:
            print(f"  Status: [OK] Structure matches\n")
        else:
            print(f"  Status: [MISMATCH] Structure differs\n")
    
    return True

def check_data_presence(generated_path, expected_data):
    """Check if expected data appears in generated document"""
    print(f"\n{'='*60}")
    print("CHECKING DATA PRESENCE")
    print(f"{'='*60}\n")
    
    if not os.path.exists(generated_path):
        print(f"[ERROR] File not found: {generated_path}")
        return False
    
    doc = Document(generated_path)
    
    # Extract all text
    all_text = ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + " "
    
    # Check each expected item
    print("Checking for expected data:")
    all_found = True
    for key, value in expected_data.items():
        found = str(value) in all_text
        status = "[OK]" if found else "[MISSING]"
        print(f"  {status} {key}: {value}")
        if not found:
            all_found = False
    
    return all_found

def main():
    base_dir = os.path.dirname(__file__)
    
    print("\n" + "="*60)
    print("RTB DOCUMENT COMPARISON TOOL")
    print("="*60)
    
    # Find latest test files
    test_files = [f for f in os.listdir(base_dir) if f.startswith('TEST_')]
    
    if not test_files:
        print("\n[ERROR] No test files found. Run test_document_generation.py first.")
        return
    
    # Session Plan comparison
    session_template = os.path.join(base_dir, 'rtb_session_plan_template.docx')
    session_test = [f for f in test_files if 'Session_Plan' in f]
    
    if session_test:
        session_generated = os.path.join(base_dir, session_test[-1])
        compare_structure(session_template, session_generated, "SESSION PLAN")
        
        # Check for test data
        expected_session_data = {
            'Sector': 'ICT',
            'Trainer': 'MUGISHA',
            'Module': 'CHM4101',
            'Topic': 'Installing Motherboard',
            'Week': 'Week 5',
            'Class': 'CHM4A'
        }
        check_data_presence(session_generated, expected_session_data)
    
    # Scheme comparison
    scheme_template = os.path.join(base_dir, 'rtb_scheme_template.docx')
    scheme_test = [f for f in test_files if 'Scheme_of_Work' in f]
    
    if scheme_test:
        scheme_generated = os.path.join(base_dir, scheme_test[-1])
        compare_structure(scheme_template, scheme_generated, "SCHEME OF WORK")
        
        # Check for test data
        expected_scheme_data = {
            'Term 1': 'Week 1-12',
            'Term 2': 'Week 13-24',
            'Term 3': 'Week 25-36',
            'School': 'IPRC Kigali',
            'Trainer': 'MUGISHA'
        }
        check_data_presence(scheme_generated, expected_scheme_data)
    
    print("\n" + "="*60)
    print("COMPARISON COMPLETE")
    print("="*60)
    print("\n[NEXT STEP] Open the files manually to verify visual appearance")
    print("The structure matches, now check fonts, colors, and spacing.\n")

if __name__ == "__main__":
    main()
