"""Document Generator - Uses Official RTB Templates"""
from official_template_filler import fill_session_plan_official, fill_scheme_official
import logging

logger = logging.getLogger(__name__)

def generate_session_plan_docx(data):
    """Generate session plan using official RTB template"""
    from official_template_filler import fill_session_plan_official
    
    logger.info("Generating session plan with official template")
    
    try:
        file_path = fill_session_plan_official(data)
        logger.info(f"Document created: {file_path}")
        return file_path
    except Exception as e:
        logger.error(f"Template generation failed: {str(e)}")
        # Fallback to simple document
        from docx import Document
        import tempfile
        
        doc = Document()
        doc.add_heading('RTB SESSION PLAN', 0)
        doc.add_paragraph(f"Topic: {data.get('topic_of_session', '')}")
        doc.add_paragraph(f"\nObjectives:\n{data.get('objectives', '')}")
        doc.add_paragraph(f"\nLearning Activities:\n{data.get('learning_activities', '')}")
        doc.add_paragraph(f"\nAssessment:\n{data.get('assessment_details', '')}")
        doc.add_paragraph(f"\nReferences:\n{data.get('references', '')}")
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        return temp_file.name

def generate_scheme_of_work_docx(data):
    """Generate scheme of work using official RTB template"""
    logger.info("=== SCHEME GENERATION START ===")
    logger.info(f"Data keys: {list(data.keys()) if data else 'NO DATA'}")
    
    try:
        if not data:
            raise ValueError("No data provided for scheme generation")
        
        logger.info("Calling fill_scheme_official...")
        file_path = fill_scheme_official(data)
        
        if not file_path:
            raise ValueError("fill_scheme_official returned None")
        
        logger.info(f"Scheme document created: {file_path}")
        
        import os
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Generated file not found: {file_path}")
        
        file_size = os.path.getsize(file_path)
        logger.info(f"File size: {file_size} bytes")
        
        return file_path
        
    except Exception as e:
        logger.error(f"=== SCHEME GENERATION ERROR ===")
        logger.error(f"Error type: {type(e).__name__}")
        logger.error(f"Error message: {str(e)}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        raise
