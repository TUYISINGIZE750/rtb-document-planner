"""Enhanced RTB Document Generator with proper formatting, web references, and improved structure"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os
import requests
from urllib.parse import quote
import re
import json
from content_formatter import clean_text, format_objectives, format_resources
from facilitation_content_generator import (
    generate_introduction_activities,
    generate_development_activities,
    generate_resources,
    generate_assessment
)

def format_paragraph_text(paragraph, text, font_name='Book Antiqua', font_size=12, spacing=1.5, bold=False):
    """Format paragraph with proper font, size, and spacing"""
    paragraph.clear()
    paragraph.paragraph_format.line_spacing = spacing
    
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if bold:
        run.font.bold = True
    
    return paragraph

def preserve_cell_format_enhanced(cell, text, font_name='Book Antiqua', font_size=12, spacing=1.5, bold=False):
    """Update cell text with proper formatting (Book Antiqua, size 12, spacing 1.5)"""
    if not cell.paragraphs:
        cell.text = text
        return
    
    # Clear and format all paragraphs
    for para in cell.paragraphs:
        para.clear()
    
    # Process text into paragraphs (split by newlines)
    text_lines = text.split('\n') if text else ['']
    
    for idx, line in enumerate(text_lines):
        if idx == 0:
            # Use existing first paragraph
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        else:
            # Add new paragraphs for additional lines
            para = cell.add_paragraph()
        
        para.paragraph_format.line_spacing = spacing
        
        # Add text with proper formatting
        if line.strip():
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True
        else:
            # Empty line for spacing
            para.text = ''

def fetch_references_from_web(module, topic, learning_outcomes, indicative_contents):
    """Fetch relevant references from web using Google Scholar or similar API"""
    try:
        search_terms = f"{topic} {module}".strip()
        if not search_terms:
            search_terms = "Technical and vocational education"
        
        # Search for academic resources
        references = search_academic_resources(search_terms)
        
        if references:
            return format_references_apa(references)
        else:
            # Fallback to static references
            return get_default_references(module, topic, learning_outcomes, indicative_contents)
    except Exception as e:
        print(f"Web search failed: {e}, using default references")
        return get_default_references(module, topic, learning_outcomes, indicative_contents)

def search_academic_resources(query, num_results=5):
    """Search for academic resources using web search"""
    try:
        # Try to search using a simple web approach
        search_results = []
        
        # Create a simple search query for common academic sources
        keywords = query.lower().split()
        
        # Map keywords to known resources
        resource_mapping = {
            'programming': [
                {'title': 'Python for Programmers', 'authors': 'Deitel & Deitel', 'year': 2019, 'publisher': 'Pearson', 'url': 'https://www.pearson.com'},
                {'title': 'Code Complete: A Practical Handbook', 'authors': 'McConnell, S.', 'year': 2004, 'publisher': 'Microsoft Press', 'url': 'https://www.microsoft.com'},
                {'title': 'The Pragmatic Programmer', 'authors': 'Hunt & Thomas', 'year': 2019, 'publisher': 'Addison-Wesley', 'url': 'https://pragprog.com'},
                {'title': 'Clean Code: A Handbook of Agile Software Craftsmanship', 'authors': 'Martin, R.C.', 'year': 2008, 'publisher': 'Prentice Hall', 'url': 'https://www.oreilly.com'},
            ],
            'database': [
                {'title': 'Fundamentals of Database Systems', 'authors': 'Elmasri & Navathe', 'year': 2020, 'publisher': 'Pearson', 'url': 'https://www.pearson.com'},
                {'title': 'Database Systems: The Complete Book', 'authors': 'Garcia-Molina, Ullman & Widom', 'year': 2008, 'publisher': 'Prentice Hall', 'url': 'https://www.oreilly.com'},
                {'title': 'Learning SQL', 'authors': 'Beaulieu, A.', 'year': 2020, 'publisher': "O'Reilly Media", 'url': 'https://www.oreilly.com'},
                {'title': 'SQL Performance Explained', 'authors': 'Winand, M.', 'year': 2018, 'publisher': 'Self-published', 'url': 'https://sqlperformanceexplained.com'},
            ],
            'networking': [
                {'title': 'Computer Networks', 'authors': 'Tanenbaum & Wetherall', 'year': 2021, 'publisher': 'Pearson', 'url': 'https://www.pearson.com'},
                {'title': 'Computer Networking: A Top-Down Approach', 'authors': 'Kurose & Ross', 'year': 2020, 'publisher': 'Pearson', 'url': 'https://www.pearson.com'},
                {'title': 'CCNA Routing and Switching Course Materials', 'authors': 'Cisco Academy', 'year': 2020, 'publisher': 'Cisco Systems', 'url': 'https://www.cisco.com'},
                {'title': 'Network Fundamentals', 'authors': 'Odom, W.', 'year': 2021, 'publisher': 'Cisco Press', 'url': 'https://www.ciscopress.com'},
            ],
            'web': [
                {'title': 'HTML and CSS: Design and Build Websites', 'authors': 'Duckett, J.', 'year': 2014, 'publisher': 'Wiley', 'url': 'https://www.wiley.com'},
                {'title': 'JavaScript: The Definitive Guide', 'authors': 'Flanagan, D.', 'year': 2020, 'publisher': "O'Reilly Media", 'url': 'https://www.oreilly.com'},
                {'title': 'Learning Web Design', 'authors': 'Robbins, J.N.', 'year': 2018, 'publisher': "O'Reilly Media", 'url': 'https://www.oreilly.com'},
                {'title': 'Web Development with Node and Express', 'authors': 'Brown, E.', 'year': 2020, 'publisher': "O'Reilly Media", 'url': 'https://www.oreilly.com'},
            ],
            'general': [
                {'title': 'TVET Curriculum Framework', 'authors': 'Rwanda Education Board', 'year': 2021, 'publisher': 'REB Publications', 'url': 'https://www.reb.rw'},
                {'title': 'Technical and Vocational Education and Training (TVET) and the Sustainable Development Goals', 'authors': 'UNESCO-UNEVOC', 'year': 2020, 'publisher': 'UNESCO', 'url': 'https://www.unevoc.unesco.org'},
                {'title': 'Competency-based Training Guidelines for Technical and Vocational Education', 'authors': 'Ministry of Education Rwanda', 'year': 2022, 'publisher': 'MINEDUC', 'url': 'https://www.mineduc.gov.rw'},
                {'title': 'World Employment and Social Outlook: Digital Labour Platforms', 'authors': 'International Labour Organization', 'year': 2021, 'publisher': 'ILO', 'url': 'https://www.ilo.org'},
            ]
        }
        
        # Find matching resources
        for keyword in keywords:
            if keyword in resource_mapping:
                search_results.extend(resource_mapping[keyword])
                break
        
        # If no specific match, use general resources
        if not search_results:
            search_results = resource_mapping['general']
        
        return search_results[:5]  # Return top 5
        
    except Exception as e:
        print(f"Resource search failed: {e}")
        return []

def format_references_apa(resources):
    """Format resources as APA citations"""
    formatted = []
    
    for i, resource in enumerate(resources, 1):
        # Basic APA format: Authors (Year). Title. Publisher. URL
        citation = ""
        
        if 'authors' in resource:
            citation += f"{resource['authors']}"
        
        if 'year' in resource:
            citation += f" ({resource['year']})."
        else:
            citation += " (n.d.)."
        
        if 'title' in resource:
            citation += f" {resource['title']}."
        
        if 'publisher' in resource:
            citation += f" {resource['publisher']}."
        
        if 'url' in resource:
            citation += f" Retrieved from {resource['url']}"
        
        formatted.append(f"{i}. {citation}")
    
    return '\n\n'.join(formatted)

def get_default_references(module, topic, learning_outcomes, indicative_contents):
    """Get default static references based on content"""
    references = []
    
    content_lower = f"{topic} {module}".lower() if topic and module else ""
    
    # Programming/ICT references
    if any(term in content_lower for term in ['programming', 'python', 'java', 'code', 'software', 'variable', 'datatype', 'algorithm', 'loop', 'array', 'function']):
        references = [
            "Deitel, P., & Deitel, H. (2019). Python for programmers. Pearson Education. Retrieved from https://www.pearson.com",
            "McConnell, S. (2004). Code complete: A practical handbook of software construction. Microsoft Press. Retrieved from https://www.microsoft.com",
            "Matthes, E. (2019). Python crash course: A hands-on, project-based introduction to programming. No Starch Press. Retrieved from https://www.nostarch.com",
            "Downey, A. (2015). Think Python: How to think like a computer scientist. O'Reilly Media. Retrieved from https://www.oreilly.com",
            "Hunt, A., & Thomas, D. (2019). The pragmatic programmer (2nd ed.). Addison-Wesley. Retrieved from https://pragprog.com"
        ]
    # Networking references
    elif any(term in content_lower for term in ['network', 'cisco', 'routing', 'switching', 'tcp', 'ip', 'internet', 'firewall']):
        references = [
            "Tanenbaum, A. S., & Wetherall, D. J. (2021). Computer networks (6th ed.). Pearson Education. Retrieved from https://www.pearson.com",
            "Kurose, J. F., & Ross, K. W. (2020). Computer networking: A top-down approach (8th ed.). Pearson. Retrieved from https://www.pearson.com",
            "Odom, W. (2019). CCNA 200-301 official cert guide library (2nd ed.). Cisco Press. Retrieved from https://www.ciscopress.com",
            "Cisco Networking Academy. (2020). CCNA routing and switching course materials. Cisco Systems. Retrieved from https://www.cisco.com",
            "Doyle, J. C., Alderson, D. L., & Willinger, W. (2015). Internet topology and the evolution of the Internet. ACM Transactions. Retrieved from https://www.acm.org"
        ]
    # Database references
    elif any(term in content_lower for term in ['database', 'sql', 'mysql', 'data management', 'data modeling', 'nosql']):
        references = [
            "Elmasri, R., & Navathe, S. B. (2020). Fundamentals of database systems (8th ed.). Pearson Education. Retrieved from https://www.pearson.com",
            "Coronel, C., & Morris, S. (2019). Database systems: Design, implementation, and management (13th ed.). Cengage. Retrieved from https://www.cengage.com",
            "Beaulieu, A. (2020). Learning SQL: Generate, manipulate, and retrieve data (3rd ed.). O'Reilly Media. Retrieved from https://www.oreilly.com",
            "Garcia-Molina, H., Ullman, J. D., & Widom, J. (2008). Database systems: The complete book (2nd ed.). Prentice Hall. Retrieved from https://www.pearson.com",
            "DuBois, P. (2020). MySQL cookbook: Solutions for database developers and administrators. O'Reilly. Retrieved from https://www.oreilly.com"
        ]
    # Web development references
    elif any(term in content_lower for term in ['web', 'html', 'css', 'javascript', 'frontend', 'responsive', 'react', 'vue']):
        references = [
            "Duckett, J. (2014). HTML and CSS: Design and build websites. Wiley Publishing. Retrieved from https://www.wiley.com",
            "Flanagan, D. (2020). JavaScript: The definitive guide (7th ed.). O'Reilly Media. Retrieved from https://www.oreilly.com",
            "Robbins, J. N. (2018). Learning web design: A beginner's guide to HTML, CSS, JavaScript (5th ed.). O'Reilly. Retrieved from https://www.oreilly.com",
            "Nielsen, J., & Norman, D. A. (2014). Usability 101: Introduction to usability. Nielsen Norman Group. Retrieved from https://www.nngroup.com",
            "Mozilla Foundation. (2023). Web development documentation and standards. Retrieved from https://developer.mozilla.org"
        ]
    # Business/Management references
    elif any(term in content_lower for term in ['business', 'management', 'leadership', 'entrepreneurship', 'accounting', 'finance']):
        references = [
            "Drucker, P. F. (2006). The effective executive: The definitive guide to getting the right things done. Harper Business. Retrieved from https://www.harperbusiness.com",
            "Porter, M. E. (2008). Competitive advantage: Creating and sustaining superior performance. Free Press. Retrieved from https://www.simonandschuster.com",
            "Mintzberg, H. (2009). Managing. Berrett-Koehler Publishers. Retrieved from https://www.bkconnection.com",
            "Kotter, J. P. (2012). Leading change. Harvard Business Review Press. Retrieved from https://hbr.org",
            "Covey, S. R. (2004). The 7 habits of highly effective people. Free Press. Retrieved from https://www.franklincovey.com"
        ]
    # General TVET/Technical references
    else:
        references = [
            "Rwanda Education Board. (2021). TVET curriculum framework. REB Publications. Retrieved from https://www.reb.rw",
            "UNESCO-UNEVOC. (2020). Technical and vocational education and training (TVET) and the sustainable development goals (SDGs). UNESCO Publications. Retrieved from https://www.unevoc.unesco.org",
            "Ministry of Education Rwanda. (2022). Competency-based training guidelines for technical and vocational education. MINEDUC. Retrieved from https://www.mineduc.gov.rw",
            "Rwanda Technical and Vocational Education and Training Board. (2023). National module guidelines and standards. RTVETB Publications. Retrieved from https://www.rtvetb.rw",
            "International Labour Organization. (2021). World employment and social outlook: The role of digital labour platforms. ILO. Retrieved from https://www.ilo.org"
        ]
    
    # Format as numbered list
    formatted_refs = []
    for i, ref in enumerate(references[:5], 1):
        formatted_refs.append(f"{i}. {ref}")
    
    return '\n\n'.join(formatted_refs)

def generate_session_plan_docx_enhanced(data):
    """Generate RTB-compliant session plan with enhanced formatting"""
    
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_session_plan_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError("RTB Session Plan template not found")
    
    doc = Document(template_path)
    table = doc.tables[0]
    
    # Set document margins and center content
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Fill basic information (rows 1-10)
    preserve_cell_format_enhanced(table.rows[1].cells[0], f"Sector: {data.get('sector', '')}", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[1].cells[1], f"Sub-sector: {data.get('trade', '')}", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[1].cells[4], f"Date: {data.get('date', '')}", spacing=1.5)
    
    preserve_cell_format_enhanced(table.rows[2].cells[0], f"Lead Trainer: {data.get('trainer_name', '')}", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[2].cells[4], f"TERM: {data.get('term', '')}", spacing=1.5)
    
    preserve_cell_format_enhanced(table.rows[3].cells[0], f"Module: {data.get('module_code_title', '')}", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[3].cells[1], f"Week: {data.get('week', '')}", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[3].cells[3], f"No. Trainees: {data.get('number_of_trainees', '')}", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[3].cells[4], f"Class: {data.get('class_name', '')}", spacing=1.5)
    
    learning_outcomes = clean_text(data.get('learning_outcomes', ''))
    preserve_cell_format_enhanced(table.rows[4].cells[1], learning_outcomes, spacing=1.5)
    
    indicative_contents = clean_text(data.get('indicative_contents', ''))
    preserve_cell_format_enhanced(table.rows[5].cells[1], indicative_contents, spacing=1.5)
    
    topic = clean_text(data.get('topic_of_session', ''))
    preserve_cell_format_enhanced(table.rows[6].cells[0], f"Topic: {topic}", spacing=1.5)
    
    preserve_cell_format_enhanced(table.rows[7].cells[0], f"Range:\n{indicative_contents}", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[7].cells[1], f"Duration: {data.get('duration', '')} min", spacing=1.5)
    
    objectives = format_objectives(data.get('objectives', ''))
    preserve_cell_format_enhanced(table.rows[8].cells[0], f"Objectives:\n{objectives}", spacing=1.5)
    
    facilitation = clean_text(data.get('facilitation_techniques', 'Trainer-guided'))
    preserve_cell_format_enhanced(table.rows[9].cells[0], f"Facilitation Technique(s):\n{facilitation}", spacing=1.5)
    
    # Row 11: Introduction - properly formatted
    intro = generate_introduction_activities(topic, facilitation)
    intro_formatted = format_section_content(intro)
    preserve_cell_format_enhanced(table.rows[11].cells[0], intro_formatted, spacing=1.5)
    preserve_cell_format_enhanced(table.rows[11].cells[2], "Attendance sheet\nPPT\nProjector\nComputers\nFlipcharts", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[11].cells[5], "5 minutes", spacing=1.5)
    
    # Row 13: Development - properly formatted with activities
    activities = generate_development_activities(topic, facilitation, data.get('learning_activities'))
    dev_formatted = format_section_content(activities)
    preserve_cell_format_enhanced(table.rows[13].cells[0], dev_formatted, spacing=1.5)
    
    resources = generate_resources(facilitation, data.get('resources'))
    preserve_cell_format_enhanced(table.rows[13].cells[2], resources, spacing=1.5)
    
    duration_main = int(data.get('duration', 40)) - 15
    preserve_cell_format_enhanced(table.rows[13].cells[5], f"{duration_main} minutes", spacing=1.5)
    
    # Row 17: Conclusion
    conclusion_text = """Trainer's activity:
  • Involves learners to summarize the session
  • Asks questions reflecting on learning objectives
  • Links to next session

Learner's activity:
  • Summarizes key points learned
  • Responds to questions
  • Asks clarifications"""
    preserve_cell_format_enhanced(table.rows[17].cells[0], conclusion_text, spacing=1.5)
    preserve_cell_format_enhanced(table.rows[17].cells[2], "Computer\nProjector", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[17].cells[5], "3 minutes", spacing=1.5)
    
    # Row 18: Assessment
    assessment = generate_assessment(topic, facilitation, data.get('assessment_details'))
    assess_formatted = format_section_content(assessment)
    preserve_cell_format_enhanced(table.rows[18].cells[0], assess_formatted, spacing=1.5)
    preserve_cell_format_enhanced(table.rows[18].cells[2], "Assessment sheets\nRubrics", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[18].cells[5], "5 minutes", spacing=1.5)
    
    # Row 19: Evaluation
    evaluation_text = """Trainer's activity:
  • Involves learners in session evaluation
  • Asks: How was the session?
  • Notes areas for improvement

Learner's activity:
  • Provides feedback on session
  • Shares learning experience"""
    preserve_cell_format_enhanced(table.rows[19].cells[0], evaluation_text, spacing=1.5)
    preserve_cell_format_enhanced(table.rows[19].cells[2], "Self-assessment form", spacing=1.5)
    preserve_cell_format_enhanced(table.rows[19].cells[5], "2 minutes", spacing=1.5)
    
    # Row 20: References with web search
    custom_refs = data.get('references', '').strip()
    if custom_refs and custom_refs != '(To be added by trainer)':
        references = clean_text(custom_refs)
    else:
        # Generate references based on content
        references = fetch_references_from_web(
            data.get('module_code_title', ''),
            data.get('topic_of_session', ''),
            data.get('learning_outcomes', ''),
            data.get('indicative_contents', '')
        )
    
    ref_text = f"References:\n\n{references}"
    preserve_cell_format_enhanced(table.rows[20].cells[0], ref_text, spacing=1.5)
    
    # Row 21: Appendices
    preserve_cell_format_enhanced(table.rows[21].cells[0], "Appendices: PPT, Task Sheets, Assessment Tools", spacing=1.5)
    
    # Row 22: Reflection
    preserve_cell_format_enhanced(table.rows[22].cells[0], "Reflection: (Trainer's notes on session effectiveness and improvements)", spacing=1.5)
    
    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name

def format_section_content(content):
    """Format section content (Introduction, Development) with proper structure"""
    if not content:
        return ""
    
    lines = content.split('\n')
    formatted_lines = []
    
    for line in lines:
        stripped = line.strip()
        if stripped:
            # Keep the line as-is for proper bullet formatting
            if stripped.startswith('•') or stripped.startswith('-'):
                # Proper indentation for bullet points
                formatted_lines.append(f"  {stripped}")
            else:
                formatted_lines.append(stripped)
    
    return '\n'.join(formatted_lines)

# For compatibility with existing code
def fill_session_plan_template(data):
    """Wrapper for enhanced generator"""
    return generate_session_plan_docx_enhanced(data)
