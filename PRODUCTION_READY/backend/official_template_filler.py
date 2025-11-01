"""Official RTB Template Filler - Uses templates from DOCS TO REFER TO folder"""
from docx import Document
from docx.shared import Pt, RGBColor
import os
import tempfile
from datetime import datetime

def set_cell_font(cell, font_name='Bookman Old Style', font_size=12, bold=False):
    """Set font for all paragraphs and runs in a cell"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
        if not paragraph.runs and paragraph.text:
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ''
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold

def set_cell_text_with_bold_label(cell, label, value, font_name='Bookman Old Style', font_size=12):
    """Set cell text with bold label and normal value"""
    cell.text = ''
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    
    label_run = p.add_run(label)
    label_run.font.name = font_name
    label_run.font.size = Pt(font_size)
    label_run.font.bold = True
    
    value_run = p.add_run(value)
    value_run.font.name = font_name
    value_run.font.size = Pt(font_size)
    value_run.font.bold = False

def fill_session_plan_official(data):
    """Fill RTB Session plan template.docx from RTB Templates folder"""
    import sys
    import logging
    logger = logging.getLogger(__name__)
    
    if data is None:
        logger.error("❌ Data is None!")
        raise ValueError("Data cannot be None")
    
    logger.info(f"🐍 Python version: {sys.version}")
    logger.info(f"📂 Current file: {__file__}")
    logger.info(f"📂 Current dir: {os.getcwd()}")
    logger.info(f"📊 Data type: {type(data)}")
    logger.info(f"📊 Data keys: {list(data.keys()) if isinstance(data, dict) else 'NOT A DICT'}")
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    logger.info(f"📂 Base dir: {base_dir}")
    
    # List files in base directory
    try:
        files = os.listdir(base_dir)
        logger.info(f"📂 Files in base dir: {files[:10]}")
    except Exception as e:
        logger.error(f"❌ Cannot list base dir: {e}")
    
    # Check if RTB Templates folder exists
    rtb_folder = os.path.join(base_dir, 'RTB Templates')
    logger.info(f"📂 RTB Templates path: {rtb_folder}")
    logger.info(f"📂 RTB Templates exists: {os.path.exists(rtb_folder)}")
    
    if os.path.exists(rtb_folder):
        try:
            template_files = os.listdir(rtb_folder)
            logger.info(f"📂 Files in RTB Templates: {template_files}")
        except Exception as e:
            logger.error(f"❌ Cannot list RTB Templates: {e}")
    
    template_path = os.path.join(base_dir, 'RTB Templates', 'RTB Session plan template.docx')
    logger.info(f"📂 Full template path: {template_path}")
    logger.info(f"📂 Template exists: {os.path.exists(template_path)}")
    
    if not os.path.exists(template_path):
        logger.warning(f"⚠️ Template not found, creating simple document")
        # Fallback: create simple document
        doc = Document()
        doc.add_heading('RTB Session Plan', 0)
        doc.add_paragraph(f"Topic: {data.get('topic_of_session', '')}")
        doc.add_paragraph(f"Objectives: {data.get('objectives', '')}")
        doc.add_paragraph(f"Learning Activities: {data.get('learning_activities', '')}")
        doc.add_paragraph(f"Assessment: {data.get('assessment_details', '')}")
        doc.add_paragraph(f"References: {data.get('references', '')}")
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        logger.info(f"✅ Fallback document created: {temp_file.name}")
        return temp_file.name
    
    try:
        doc = Document(template_path)
        logger.info(f"✅ Template loaded, tables: {len(doc.tables)}")
        
        if not doc.tables:
            raise Exception("No table found in template")
    except Exception as e:
        logger.error(f"❌ Error loading template: {str(e)}")
        raise
    
    if not doc.tables:
        logger.error("❌ No tables in template!")
        raise Exception("Template has no tables")
    
    table = doc.tables[0]
    logger.info(f"✅ Table found with {len(table.rows)} rows")
    
    # Row 0: Bold headers
    set_cell_text_with_bold_label(table.rows[0].cells[0], "Sector: ", data.get('sector', '').strip())
    set_cell_text_with_bold_label(table.rows[0].cells[1], "Trade: ", data.get('trade', '').strip())
    set_cell_text_with_bold_label(table.rows[0].cells[4], "Level: ", data.get('level', 'Level 4').strip())
    set_cell_text_with_bold_label(table.rows[0].cells[6], "Date: ", data.get('date', '').strip())
    
    # Row 1: Bold headers
    set_cell_text_with_bold_label(table.rows[1].cells[0], "Trainer name: ", data.get('trainer_name', '').strip())
    cell = table.rows[1].cells[6]
    cell.text = ''
    p = cell.paragraphs[0]
    r1 = p.add_run("School year: ")
    r1.font.bold = True
    r1.font.name = 'Bookman Old Style'
    r1.font.size = Pt(12)
    r2 = p.add_run(data.get('school_year', '2024-2025'))
    r2.font.name = 'Bookman Old Style'
    r2.font.size = Pt(12)
    p.add_run('\n')
    r3 = p.add_run("Term: ")
    r3.font.bold = True
    r3.font.name = 'Bookman Old Style'
    r3.font.size = Pt(12)
    r4 = p.add_run(data.get('term', ''))
    r4.font.name = 'Bookman Old Style'
    r4.font.size = Pt(12)
    
    # Row 2: Bold headers
    set_cell_text_with_bold_label(table.rows[2].cells[0], "Module (Code&Name): ", data.get('module_code_title', '').strip())
    set_cell_text_with_bold_label(table.rows[2].cells[1], "Week: ", data.get('week', '1').strip())
    set_cell_text_with_bold_label(table.rows[2].cells[5], "No. Trainees: ", data.get('number_of_trainees', '').strip())
    set_cell_text_with_bold_label(table.rows[2].cells[6], "Class(es): ", data.get('class_name', '').strip())
    
    # Row 3: Learning Outcome
    set_cell_font(table.rows[3].cells[0], bold=True)
    table.rows[3].cells[1].text = data.get('learning_outcomes', '').strip()
    set_cell_font(table.rows[3].cells[1])
    
    # Row 4: Indicative content
    set_cell_font(table.rows[4].cells[0], bold=True)
    table.rows[4].cells[1].text = data.get('indicative_contents', '').strip()
    set_cell_font(table.rows[4].cells[1])
    
    # Row 5: Topic
    set_cell_text_with_bold_label(table.rows[5].cells[0], "Topic of the session: ", data.get('topic_of_session', '').strip())
    
    # Row 6: Range and Duration
    range_val = data.get('range', '').strip()
    set_cell_text_with_bold_label(table.rows[6].cells[0], "Range: ", range_val)
    set_cell_text_with_bold_label(table.rows[6].cells[2], "Duration of the session: ", data.get('duration', '').strip())
    
    # Row 7: Objectives
    objectives = data.get('objectives', '').strip()
    logger.info(f"📝 Filling objectives: {objectives[:100] if objectives else 'EMPTY'}")
    set_cell_text_with_bold_label(table.rows[7].cells[0], "Objectives:\n", objectives)
    
    # Row 8: Facilitation techniques
    set_cell_text_with_bold_label(table.rows[8].cells[0], "Facilitation technique(s): ", data.get('facilitation_techniques', '').strip())
    
    # Parse learning activities and resources
    learning_acts = data.get('learning_activities', '')
    resources_text = data.get('resources', '')
    
    # Split activities into sections
    sections = learning_acts.split('\n\n') if learning_acts else []
    intro = ''
    dev = ''
    conclusion = ''
    
    for section in sections:
        if 'Introduction:' in section:
            intro = section.replace('Introduction:', '').strip()
        elif 'Development:' in section:
            dev = section.replace('Development:', '').strip()
        elif 'Conclusion:' in section:
            conclusion = section.replace('Conclusion:', '').strip()
    
    # Parse resources
    resource_lines = [r.strip() for r in resources_text.split('\n') if r.strip()] if resources_text else []
    all_resources = ', '.join(resource_lines) if resource_lines else 'Whiteboard, markers, handouts'
    
    # Row 9: Bold headers
    set_cell_font(table.rows[9].cells[0], bold=True)
    set_cell_font(table.rows[9].cells[3], bold=True)
    set_cell_font(table.rows[9].cells[7], bold=True)
    
    # Row 10: Introduction
    set_cell_text_with_bold_label(table.rows[10].cells[0], "Trainer's activity: ", f"{intro}\n\n")
    cell = table.rows[10].cells[0]
    p = cell.paragraphs[0]
    r = p.add_run("Learner's activity: ")
    r.font.bold = True
    r.font.name = 'Bookman Old Style'
    r.font.size = Pt(12)
    r2 = p.add_run("Participate actively and ask questions")
    r2.font.name = 'Bookman Old Style'
    r2.font.size = Pt(12)
    table.rows[10].cells[3].text = all_resources
    set_cell_font(table.rows[10].cells[3])
    table.rows[10].cells[7].text = "5 min"
    set_cell_font(table.rows[10].cells[7])
    
    # Row 11: Bold header
    set_cell_font(table.rows[11].cells[0], bold=True)
    
    # Row 12: Development
    cell = table.rows[12].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    r1 = p.add_run("Step 1:\n")
    r1.font.bold = True
    r1.font.name = 'Bookman Old Style'
    r1.font.size = Pt(12)
    r2 = p.add_run("Trainer's activity: ")
    r2.font.bold = True
    r2.font.name = 'Bookman Old Style'
    r2.font.size = Pt(12)
    r3 = p.add_run(f"{dev}\n\n")
    r3.font.name = 'Bookman Old Style'
    r3.font.size = Pt(12)
    r4 = p.add_run("Learner's activity: ")
    r4.font.bold = True
    r4.font.name = 'Bookman Old Style'
    r4.font.size = Pt(12)
    r5 = p.add_run("Practice and apply concepts")
    r5.font.name = 'Bookman Old Style'
    r5.font.size = Pt(12)
    table.rows[12].cells[3].text = all_resources
    set_cell_font(table.rows[12].cells[3])
    table.rows[12].cells[7].text = "30 min"
    set_cell_font(table.rows[12].cells[7])
    
    # Row 15: Bold header
    set_cell_font(table.rows[15].cells[0], bold=True)
    
    # Row 16: Conclusion
    set_cell_text_with_bold_label(table.rows[16].cells[0], "Summary: ", conclusion)
    table.rows[16].cells[3].text = all_resources
    set_cell_font(table.rows[16].cells[3])
    table.rows[16].cells[7].text = "5 min"
    set_cell_font(table.rows[16].cells[7])
    
    # Row 17: Assessment - bold header
    assessment = data.get('assessment_details', '').strip()
    table.rows[17].cells[0].text = assessment
    set_cell_font(table.rows[17].cells[0])
    table.rows[17].cells[3].text = "Assessment sheets"
    set_cell_font(table.rows[17].cells[3])
    table.rows[17].cells[7].text = "5 min"
    set_cell_font(table.rows[17].cells[7])
    
    # Row 19: References
    references = data.get('references', '').strip()
    set_cell_text_with_bold_label(table.rows[19].cells[0], "References:\n", references)
    
    # Row 20: Appendices
    appendix = data.get('appendix', '').strip()
    if not appendix:
        appendix = "PPT, Task Sheets, Assessment"
    set_cell_text_with_bold_label(table.rows[20].cells[0], "Appendices:\n", appendix)
    
    # Save to temp file
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        logger.info(f"💾 Saving to: {temp_file.name}")
        doc.save(temp_file.name)
        temp_file.close()
        logger.info(f"✅ Document saved successfully")
        return temp_file.name
    except Exception as e:
        logger.error(f"❌ Error saving document: {str(e)}")
        raise

def fill_scheme_official(data):
    """Fill Scheme of work.docx from RTB Templates folder"""
    template_path = os.path.join(os.path.dirname(__file__), 'RTB Templates', 'Scheme of work.docx')
    doc = Document(template_path)
    
    # Update header paragraphs
    if len(doc.paragraphs) > 3:
        doc.paragraphs[0].text = data.get('province', '')
        doc.paragraphs[1].text = data.get('district', '')
        doc.paragraphs[2].text = data.get('sector', '')
        doc.paragraphs[3].text = data.get('school', '')
    
    # Fill Term 1 table
    if len(doc.tables) > 0:
        table1 = doc.tables[0]
        term1_weeks = data.get('term1_weeks', '').split('\n')
        term1_outcomes = data.get('term1_learning_outcomes', '').split('\n')
        term1_contents = data.get('term1_indicative_contents', '').split('\n')
        term1_duration = data.get('term1_duration', '').split('\n')
        term1_place = data.get('term1_learning_place', '').split('\n')
        
        for i in range(1, min(len(table1.rows), len(term1_weeks) + 1)):
            if i - 1 < len(term1_weeks):
                table1.rows[i].cells[0].text = term1_weeks[i-1]
            if i - 1 < len(term1_outcomes):
                table1.rows[i].cells[1].text = term1_outcomes[i-1]
            if i - 1 < len(term1_contents):
                table1.rows[i].cells[3].text = term1_contents[i-1]
            if i - 1 < len(term1_duration):
                table1.rows[i].cells[4].text = term1_duration[i-1]
            if i - 1 < len(term1_place):
                table1.rows[i].cells[7].text = term1_place[i-1]
    
    # Fill Term 2 table
    if len(doc.tables) > 1:
        table2 = doc.tables[1]
        term2_weeks = data.get('term2_weeks', '').split('\n')
        term2_outcomes = data.get('term2_learning_outcomes', '').split('\n')
        term2_contents = data.get('term2_indicative_contents', '').split('\n')
        term2_duration = data.get('term2_duration', '').split('\n')
        term2_place = data.get('term2_learning_place', '').split('\n')
        
        for i in range(1, min(len(table2.rows), len(term2_weeks) + 1)):
            if i - 1 < len(term2_weeks):
                table2.rows[i].cells[0].text = term2_weeks[i-1]
            if i - 1 < len(term2_outcomes):
                table2.rows[i].cells[1].text = term2_outcomes[i-1]
            if i - 1 < len(term2_contents):
                table2.rows[i].cells[3].text = term2_contents[i-1]
            if i - 1 < len(term2_duration):
                table2.rows[i].cells[4].text = term2_duration[i-1]
            if i - 1 < len(term2_place):
                table2.rows[i].cells[7].text = term2_place[i-1]
    
    # Fill Term 3 table
    if len(doc.tables) > 2:
        table3 = doc.tables[2]
        term3_weeks = data.get('term3_weeks', '').split('\n')
        term3_outcomes = data.get('term3_learning_outcomes', '').split('\n')
        term3_contents = data.get('term3_indicative_contents', '').split('\n')
        term3_duration = data.get('term3_duration', '').split('\n')
        term3_place = data.get('term3_learning_place', '').split('\n')
        
        for i in range(1, min(len(table3.rows), len(term3_weeks) + 1)):
            if i - 1 < len(term3_weeks):
                table3.rows[i].cells[0].text = term3_weeks[i-1]
            if i - 1 < len(term3_outcomes):
                table3.rows[i].cells[1].text = term3_outcomes[i-1]
            if i - 1 < len(term3_contents):
                table3.rows[i].cells[3].text = term3_contents[i-1]
            if i - 1 < len(term3_duration):
                table3.rows[i].cells[4].text = term3_duration[i-1]
            if i - 1 < len(term3_place):
                table3.rows[i].cells[7].text = term3_place[i-1]
    
    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    return temp_file.name
