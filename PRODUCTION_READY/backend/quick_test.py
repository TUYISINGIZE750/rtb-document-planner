#!/usr/bin/env python
# -*- coding: utf-8 -*-
from docx import Document
import os
import sys

template_path = r'c:\Users\PC\Music\Scheme of work and session plan planner\PRODUCTION_READY\RTB Templates\RTB Session plan template.docx'
generated_path = r'c:\Users\PC\Music\Scheme of work and session plan planner\PRODUCTION_READY\backend\TEST_Session_Plan_20251025_094124.docx'

print("TEMPLATE vs PREVIOUSLY GENERATED - QUICK COMPARISON")
print("=" * 80)

try:
    if os.path.exists(template_path) and os.path.exists(generated_path):
        template = Document(template_path)
        generated = Document(generated_path)
        
        t_table = template.tables[0]
        g_table = generated.tables[0]
        
        print(f"Template: {len(t_table.rows)} rows x {len(t_table.rows[0].cells)} cols")
        print(f"Generated: {len(g_table.rows)} rows x {len(g_table.rows[0].cells)} cols")
        
        if len(t_table.rows) == len(g_table.rows):
            print("✓ SAME number of rows")
        else:
            print(f"✗ Different rows: {len(t_table.rows)} vs {len(g_table.rows)}")
        
        print("\n" + "=" * 80)
        
        for row_idx in range(min(8, len(t_table.rows))):
            t_row = t_table.rows[row_idx]
            g_row = g_table.rows[row_idx]
            
            print(f"\nRow {row_idx}: ", end="")
            
            differences = 0
            for col_idx in range(min(len(t_row.cells), len(g_row.cells))):
                t_text = t_row.cells[col_idx].text.strip()[:30]
                g_text = g_row.cells[col_idx].text.strip()[:30]
                
                if t_text != g_text:
                    differences += 1
                    if differences <= 2:
                        print(f"\n  Col {col_idx}: Template='{t_text}' vs Generated='{g_text}'")
            
            if differences == 0:
                print("✓ IDENTICAL")
            else:
                print(f"  Total {differences} cells differ")
        
        print("\n" + "=" * 80)
        print("\nConclusion: Check if the data is being filled in the correct cells")
    else:
        print("ERROR: Files not found")
        if not os.path.exists(template_path):
            print(f"  Template: {template_path}")
        if not os.path.exists(generated_path):
            print(f"  Generated: {generated_path}")

except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()
