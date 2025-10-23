from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

class RTBDocumentGenerator:
    """Generate RTB-compliant documents using official templates"""
    
    def __init__(self):
        self.template_dir = os.path.dirname(__file__)
    
    def generate_session_plan(self, data):
        """Generate session plan using RTB template"""
        template_path = os.path.join(self.template_dir, 'rtb_session_plan_template.docx')
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            # Fill template with data
            self._fill_session_plan_template(doc, data)
        else:
            # Fallback: create from scratch
            doc = self._create_session_plan_from_scratch(data)
        
        return doc
    
    def _fill_session_plan_template(self, doc, data):
        """Fill the RTB template with user data"""
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Replace common placeholders
                    text = cell.text
                    
                    # Map data fields to template
                    replacements = {
                        'SECTOR': data.get('sector', ''),
                        'TRADE': data.get('trade', ''),
                        'RQF LEVEL': data.get('rqf_level', ''),
                        'TRAINER NAME': data.get('trainer_name', ''),
                        'MODULE CODE': data.get('module_code_title', ''),
                        'TERM': data.get('term', ''),
                        'WEEK': data.get('week', ''),
                        'DATE': data.get('date', ''),
                        'NUMBER OF TRAINEES': str(data.get('number_of_trainees', '')),
                        'CLASS NAME': data.get('class_name', ''),
                        'DURATION': f"{data.get('duration', '')} minutes",
                        'TOPIC': data.get('topic_of_session', ''),
                        'LEARNING OUTCOMES': data.get('learning_outcomes', ''),
                        'INDICATIVE CONTENTS': data.get('indicative_contents', ''),
                        'FACILITATION': data.get('facilitation_techniques', '')
                    }
                    
                    for placeholder, value in replacements.items():
                        if placeholder.lower() in text.lower():
                            # Clear cell and add new content
                            cell.text = value
                            # Preserve formatting
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(11)
    
    def _create_session_plan_from_scratch(self, data):
        """Create RTB session plan from scratch if template not found"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # Add RTB header
        header_table = doc.add_table(rows=3, cols=3)
        header_table.style = 'Table Grid'
        
        # Row 1: Logo, Title, Code
        header_table.cell(0, 0).text = 'RTB LOGO'
        header_table.cell(0, 1).text = 'RWANDA TECHNICAL BOARD (RTB)\nSESSION PLAN'
        header_table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_table.cell(0, 2).text = f'Code: {data.get("module_code_title", "")}'
        
        # Merge cells for title
        self._merge_cells(header_table, 0, 1, 0, 1)
        
        # Row 2: Institution details
        header_table.cell(1, 0).text = f'Sector: {data.get("sector", "")}'
        header_table.cell(1, 1).text = f'Trade: {data.get("trade", "")}'
        header_table.cell(1, 2).text = f'Level: {data.get("rqf_level", "")}'
        
        # Row 3: Session details
        header_table.cell(2, 0).text = f'Term: {data.get("term", "")}'
        header_table.cell(2, 1).text = f'Week: {data.get("week", "")}'
        header_table.cell(2, 2).text = f'Date: {data.get("date", "")}'
        
        doc.add_paragraph()
        
        # Main content table
        content_table = doc.add_table(rows=10, cols=2)
        content_table.style = 'Table Grid'
        
        # Set column widths
        content_table.columns[0].width = Inches(2.0)
        content_table.columns[1].width = Inches(5.0)
        
        # Fill content
        fields = [
            ('Trainer Name', data.get('trainer_name', '')),
            ('Class Name', data.get('class_name', '')),
            ('Number of Trainees', str(data.get('number_of_trainees', ''))),
            ('Duration', f"{data.get('duration', '')} minutes"),
            ('Topic of Session', data.get('topic_of_session', '')),
            ('Learning Outcomes', data.get('learning_outcomes', '')),
            ('Indicative Contents', data.get('indicative_contents', '')),
            ('Facilitation Techniques', data.get('facilitation_techniques', '')),
            ('Assessment Methods', 'Formative and Summative'),
            ('Resources Required', 'As per module requirements')
        ]
        
        for idx, (label, value) in enumerate(fields):
            content_table.cell(idx, 0).text = label
            content_table.cell(idx, 1).text = value
            
            # Bold labels
            for paragraph in content_table.cell(idx, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
            
            # Format values
            for paragraph in content_table.cell(idx, 1).paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
        
        return doc
    
    def generate_scheme_of_work(self, data):
        """Generate scheme of work using RTB template"""
        template_path = os.path.join(self.template_dir, 'rtb_scheme_template.docx')
        
        if os.path.exists(template_path):
            doc = Document(template_path)
            self._fill_scheme_template(doc, data)
        else:
            doc = self._create_scheme_from_scratch(data)
        
        return doc
    
    def _fill_scheme_template(self, doc, data):
        """Fill scheme template with data"""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    
                    # Replace placeholders
                    replacements = {
                        'PROVINCE': data.get('province', ''),
                        'DISTRICT': data.get('district', ''),
                        'SECTOR': data.get('sector', ''),
                        'SCHOOL': data.get('school', ''),
                        'TRADE': data.get('trade', ''),
                        'QUALIFICATION': data.get('qualification_title', ''),
                        'RQF LEVEL': data.get('rqf_level', ''),
                        'MODULE': data.get('module_code_title', ''),
                        'YEAR': data.get('school_year', ''),
                        'TRAINER': data.get('trainer_name', '')
                    }
                    
                    for placeholder, value in replacements.items():
                        if placeholder.lower() in text.lower():
                            cell.text = value
    
    def _create_scheme_from_scratch(self, data):
        """Create scheme from scratch"""
        doc = Document()
        
        # Add header
        doc.add_heading('SCHEME OF WORK', 0)
        
        # Add institution details
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Table Grid'
        
        fields = [
            ('Province', data.get('province', '')),
            ('District', data.get('district', '')),
            ('School', data.get('school', '')),
            ('Sector/Trade', data.get('trade', '')),
            ('Module', data.get('module_code_title', '')),
            ('Level', data.get('rqf_level', '')),
            ('School Year', data.get('school_year', '')),
            ('Trainer', data.get('trainer_name', ''))
        ]
        
        for idx, (label, value) in enumerate(fields):
            info_table.cell(idx, 0).text = label
            info_table.cell(idx, 1).text = value
        
        return doc
    
    def _merge_cells(self, table, start_row, start_col, end_row, end_col):
        """Merge cells in table"""
        cell = table.cell(start_row, start_col)
        for row in range(start_row, end_row + 1):
            for col in range(start_col, end_col + 1):
                if row != start_row or col != start_col:
                    cell.merge(table.cell(row, col))
