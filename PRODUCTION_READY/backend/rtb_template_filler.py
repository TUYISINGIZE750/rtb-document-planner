from docx import Document
import tempfile
import os

def fill_session_plan_template(data):
    """Fill RTB session plan template with ONLY teacher's actual data"""
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_session_plan_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError("RTB Session Plan template not found")
    
    doc = Document(template_path)
    table = doc.tables[0]
    
    # CLEAR all template data first - fill with teacher data ONLY
    # Header (rows 1-3)
    table.rows[1].cells[0].text = f"Sector :    {data.get('sector', '')}"
    table.rows[1].cells[1].text = f"Sub-sector: {data.get('trade', '')}"
    table.rows[1].cells[4].text = f"Date : {data.get('date', '')}"
    
    table.rows[2].cells[0].text = f"Lead Trainer's name : {data.get('trainer_name', '')}"
    table.rows[2].cells[4].text = f"TERM : {data.get('term', '')}"
    
    table.rows[3].cells[0].text = f"Module(Code&Name): {data.get('module_code_title', '')}"
    table.rows[3].cells[1].text = f"Week : {data.get('week', '')}"
    table.rows[3].cells[3].text = f"No. Trainees: {data.get('number_of_trainees', '')}"
    table.rows[3].cells[4].text = f"Class(es): {data.get('class_name', '')}"
    
    # Learning content (rows 4-6) - TEACHER DATA ONLY
    table.rows[4].cells[0].text = "Learning outcome:"
    table.rows[4].cells[1].text = data.get('learning_outcomes', '')
    
    table.rows[5].cells[0].text = "Indicative contents:"
    table.rows[5].cells[1].text = data.get('indicative_contents', '')
    
    table.rows[6].cells[0].text = f"Topic of the session: {data.get('topic_of_session', '')}"
    
    # Range and duration (row 7) - TEACHER DATA ONLY
    range_text = data.get('indicative_contents', '')
    table.rows[7].cells[0].text = f"Range: \n{range_text}"
    table.rows[7].cells[1].text = f"Duration of the session: {data.get('duration', '')}min"
    
    # Objectives (row 8) - AI GENERATED FROM TEACHER DATA
    objectives = data.get('objectives', '')
    if objectives:
        table.rows[8].cells[0].text = f"Objectives: By the end of this session every learner should be able to:\n{objectives}"
    
    # Facilitation technique (row 9) - TEACHER SELECTION
    table.rows[9].cells[0].text = f"Facilitation technique(s):   {data.get('facilitation_techniques', '')}"
    
    # Activities (rows 11, 13) - USE AI GENERATED CONTENT FROM TEACHER DATA
    # Introduction activity
    intro_activity = f"Trainer's activity: \n\t• Greets and makes roll call\n\t• Involves learners to set ground rules\n\t• Reviews previous session\n\t• Announces topic: {data.get('topic_of_session', '')}\n\t• Explains objectives\n\nLearner's activity: \n\t• Greets and replies to roll call\n\t• Participates in setting ground rules\n\t• Participates in review\n\t• Reads and participates in explaining objectives\n\t• Asks clarifications if any"
    
    table.rows[11].cells[0].text = intro_activity
    table.rows[11].cells[2].text = "Attendance sheet\nPPT\nProjector\nComputers\nFlipchartwhiteboard\nMarker pen"
    table.rows[11].cells[5].text = "5 minutes"
    
    # Main activities - FORMAT AI GENERATED ACTIVITIES FOR RTB TEMPLATE
    activities = data.get('learning_activities', '')
    if activities:
        # Format the AI activities to match RTB structure with Trainer/Learner activities
        # Extract key points and format properly
        formatted_activities = activities.replace("STRUCTURE:", "").replace("RESOURCES NEEDED:", "")
        # Keep only the main activity steps, remove resource section
        if "RESOURCES NEEDED:" in formatted_activities:
            formatted_activities = formatted_activities.split("RESOURCES NEEDED:")[0]
        table.rows[13].cells[0].text = formatted_activities.strip()
    else:
        # Fallback if no AI activities
        table.rows[13].cells[0].text = f"Development:\nTrainer's activity:\n\t• Presents {data.get('topic_of_session', '')} using {data.get('facilitation_techniques', '')} technique\n\t• Demonstrates key concepts\n\t• Provides examples\n\nLearner's activity:\n\t• Engages in {data.get('topic_of_session', '')} activities\n\t• Practices skills\n\t• Asks questions"
    
    # Resources - EXTRACT FROM AI GENERATED CONTENT OR USE SEPARATE RESOURCES
    resources = data.get('resources', '')
    if resources:
        # Clean up resources formatting
        resources_clean = resources.replace("• ", "").replace("\n\n", "\n")
        table.rows[13].cells[2].text = resources_clean
    else:
        table.rows[13].cells[2].text = "Computer\nProjector\nPPT\nInstalled operating system"
    
    # Duration calculation
    try:
        total_duration = int(data.get('duration', 80))
        duration_main = total_duration - 15  # Subtract intro, conclusion, assessment, evaluation
    except:
        duration_main = 65
    
    table.rows[13].cells[5].text = f"{duration_main}\nminutes"
    
    # Conclusion (row 17)
    conclusion_text = f"Conclusion\nTrainer's activity:\n\t• Guides learners to summarize key points about {data.get('topic_of_session', '')}\n\t• Reviews learning objectives achievement\n\nLearner's activity:\n\t• Summarizes main concepts learned\n\t• Reflects on objectives achieved"
    table.rows[17].cells[0].text = conclusion_text
    table.rows[17].cells[2].text = "Computer\nProjector"
    table.rows[17].cells[5].text = "3 minutes"
    
    # Assessment (row 18) - USE AI GENERATED ASSESSMENT
    assessment = data.get('assessment_details', '')
    if assessment:
        # Use AI-generated assessment with proper formatting
        assessment_text = f"Assessment/Assignment\nTrainer's activity: \n\t{assessment}\n\nLearner's activity:\n\t• Receives assessment\n\t• Answers questions"
        table.rows[18].cells[0].text = assessment_text
    else:
        # Fallback assessment
        table.rows[18].cells[0].text = f"Assessment/Assignment\nTrainer's activity: \n\t• Gives assessment on {data.get('topic_of_session', '')}\n\nLearner's activity:\n\t• Completes assessment"
    table.rows[18].cells[2].text = "Assessment sheets"
    table.rows[18].cells[5].text = "5 minutes"
    
    # Evaluation (row 19)
    evaluation_text = "Evaluation of the session:\nTrainer's activity: \n\t• Involves learners in session evaluation\n\t• Asks: How was the session? What to improve?\n\t• Links current session to next one\n\nLearner's activity:\n\t• Answers evaluation questions\n\t• Understands what will be covered in next session"
    table.rows[19].cells[0].text = evaluation_text
    table.rows[19].cells[2].text = "Self-assessment form"
    table.rows[19].cells[5].text = "2 minutes"
    
    # References (row 20) - USE TEACHER'S REFERENCES IF PROVIDED
    references = data.get('references', '')
    if references:
        table.rows[20].cells[0].text = f"References:\nBibliography\n\n{references}"
    else:
        table.rows[20].cells[0].text = "References:\nBibliography\n\n(To be added by trainer)"
    
    # Appendices (row 21)
    table.rows[21].cells[0].text = "Appendices: PPT, Task Sheets, Assessment"
    
    # Reflection (row 22)
    table.rows[22].cells[0].text = "Reflection: (To be completed after session)"
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name

def fill_scheme_template(data):
    """Fill RTB scheme of work template with exact user data matching RTB format"""
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_scheme_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError("RTB Scheme template not found")
    
    doc = Document(template_path)
    
    # Fill header information in first table if it exists
    if len(doc.tables) > 0:
        header_table = doc.tables[0]
        if len(header_table.rows) > 2:
            # Fill weeks, learning outcomes, duration, indicative contents
            header_table.rows[2].cells[0].text = data.get('term1_weeks', f"{data.get('date', '')} - End of Term 1")
            header_table.rows[2].cells[1].text = data.get('term1_learning_outcomes', data.get('learning_outcomes', ''))
            header_table.rows[2].cells[2].text = data.get('term1_duration', f"{data.get('duration', '40')} hours")
            header_table.rows[2].cells[3].text = data.get('term1_indicative_contents', data.get('indicative_contents', ''))
            
            # Fill learning activities and resources
            if len(header_table.rows[2].cells) > 4:
                activities = data.get('learning_activities', 'Practical exercises, Group discussions, Individual assignments')
                header_table.rows[2].cells[4].text = activities[:200] + "..." if len(activities) > 200 else activities
                
                resources = data.get('resources', 'Computers, Projector, Textbooks, Internet access')
                header_table.rows[2].cells[5].text = resources[:200] + "..." if len(resources) > 200 else resources
                
                assessment = data.get('assessment_details', 'Continuous assessment, Practical tests, Portfolio assessment')
                header_table.rows[2].cells[6].text = assessment[:150] + "..." if len(assessment) > 150 else assessment
                
                header_table.rows[2].cells[7].text = "Classroom/Lab"
                header_table.rows[2].cells[8].text = "Completed successfully"
    
    # Fill additional term tables if they exist
    if len(doc.tables) > 1:
        table2 = doc.tables[1]
        if len(table2.rows) > 2:
            table2.rows[2].cells[0].text = data.get('term2_weeks', 'Term 2 weeks')
            table2.rows[2].cells[1].text = data.get('term2_learning_outcomes', data.get('learning_outcomes', ''))
            table2.rows[2].cells[2].text = data.get('term2_duration', f"{data.get('duration', '40')} hours")
            table2.rows[2].cells[3].text = data.get('term2_indicative_contents', data.get('indicative_contents', ''))
    
    if len(doc.tables) > 2:
        table3 = doc.tables[2]
        if len(table3.rows) > 2:
            table3.rows[2].cells[0].text = data.get('term3_weeks', 'Term 3 weeks')
            table3.rows[2].cells[1].text = data.get('term3_learning_outcomes', data.get('learning_outcomes', ''))
            table3.rows[2].cells[2].text = data.get('term3_duration', f"{data.get('duration', '40')} hours")
            table3.rows[2].cells[3].text = data.get('term3_indicative_contents', data.get('indicative_contents', ''))
    
    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name
