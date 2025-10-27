"""RTB Template Filler - 100% Compliance
Generates session plans that look exactly like official RTB templates
with intelligent content generation based on user input"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from datetime import datetime

from smart_content_generator import SmartSessionPlanContentGenerator


class RTBSessionPlanFiller:
    """Wrapper class for RTB template filling"""
    
    def fill_and_generate(self, data):
        """Fill and generate RTB session plan document"""
        return fill_session_plan_template_100_percent(data)


class RTBSchemeOfWorkFiller:
    """Wrapper class for RTB scheme of work template filling"""
    
    def fill_and_generate(self, data):
        """Fill and generate RTB scheme of work document"""
        return fill_scheme_of_work_template_100_percent(data)


def preserve_cell_format(cell, new_text, font_name='Book Antiqua', font_size=12, spacing=1.5):
    """Update cell text while preserving professional RTB formatting"""
    if not cell.paragraphs:
        cell.text = new_text
        return
    
    cell.paragraphs[0].clear()
    
    text_lines = new_text.split('\n') if new_text else ['']
    
    for idx, line in enumerate(text_lines):
        if idx == 0:
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        else:
            para = cell.add_paragraph()
        
        para.paragraph_format.line_spacing = spacing
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        if line.strip():
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            para.text = ''

def fill_session_plan_template_100_percent(data):
    """
    Generate RTB-compliant session plan matching official RTB template 100%
    
    Accepts:
        data (dict): User-provided session plan data with keys:
            - sector, trade, trainer_name, module_code_title
            - week, term, date, topic_of_session, duration
            - learning_outcomes, facilitation_techniques, indicative_contents
            - resources, number_of_trainees, class_name, rqf_level
    
    Returns:
        str: Path to generated DOCX file
    """
    
    try:
        # Load RTB template
        template_path = os.path.join(os.path.dirname(__file__), 'rtb_session_plan_template.docx')
        
        if not os.path.exists(template_path):
            # If template doesn't exist, create from scratch with RTB structure
            doc = _create_rtb_session_plan_from_scratch(data)
        else:
            doc = Document(template_path)
            doc = _fill_rtb_template(doc, data)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
        
    except Exception as e:
        print(f"Error in fill_session_plan_template_100_percent: {e}")
        # Fallback to creation from scratch
        return _create_rtb_session_plan_from_scratch(data)

def _fill_rtb_template(doc, data):
    """Fill existing RTB template with user data - ONLY fills template table cells"""
    
    # Set document margins to RTB standard
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Get main content table
    if not doc.tables or len(doc.tables) == 0:
        return _create_rtb_session_plan_from_scratch(data)
    
    table = doc.tables[0]
    
    try:
        # Row 1: Sector, Sub-sector (merged 1-3), Date (merged 4-5)
        if len(table.rows) > 1 and len(table.rows[1].cells) >= 6:
            preserve_cell_format(table.rows[1].cells[0], data.get('sector', ''))
            preserve_cell_format(table.rows[1].cells[1], data.get('trade', ''))
            preserve_cell_format(table.rows[1].cells[4], data.get('date', ''))
        
        # Row 2: Lead Trainer (merged 0-3), TERM (merged 4-5)
        if len(table.rows) > 2 and len(table.rows[2].cells) >= 6:
            preserve_cell_format(table.rows[2].cells[0], data.get('trainer_name', ''))
            preserve_cell_format(table.rows[2].cells[4], data.get('term', ''))
        
        # Row 3: Module, Week (merged 1-2), No. Trainees, Class (merged 4-5)
        if len(table.rows) > 3 and len(table.rows[3].cells) >= 6:
            preserve_cell_format(table.rows[3].cells[0], data.get('module_code_title', ''))
            preserve_cell_format(table.rows[3].cells[1], data.get('week', ''))
            preserve_cell_format(table.rows[3].cells[3], str(data.get('number_of_trainees', '')))
            preserve_cell_format(table.rows[3].cells[4], data.get('class_name', ''))
        
        # Row 4: Learning outcomes
        if len(table.rows) > 4 and len(table.rows[4].cells) >= 1:
            preserve_cell_format(table.rows[4].cells[0], data.get('learning_outcomes', ''))
        
        # Row 5: Indicative contents
        if len(table.rows) > 5 and len(table.rows[5].cells) >= 1:
            preserve_cell_format(table.rows[5].cells[0], data.get('indicative_contents', ''))
        
        # Row 6: Topic of session (all columns merged)
        if len(table.rows) > 6 and len(table.rows[6].cells) >= 1:
            preserve_cell_format(table.rows[6].cells[0], data.get('topic_of_session', ''))
        
        # Row 7: Range/RQF Level and Duration
        if len(table.rows) > 7 and len(table.rows[7].cells) >= 2:
            preserve_cell_format(table.rows[7].cells[0], data.get('rqf_level', ''))
            duration = data.get('duration', '')
            preserve_cell_format(table.rows[7].cells[1], f"{duration} minutes" if duration else '')
        
        # Row 8: Objectives (same as learning outcomes)
        if len(table.rows) > 8 and len(table.rows[8].cells) >= 1:
            preserve_cell_format(table.rows[8].cells[0], data.get('learning_outcomes', ''))
        
        # Row 9: Facilitation techniques
        if len(table.rows) > 9 and len(table.rows[9].cells) >= 1:
            preserve_cell_format(table.rows[9].cells[0], data.get('facilitation_techniques', ''))
        
        # Row 10: Introduction header (cell[0])
        if len(table.rows) > 10 and len(table.rows[10].cells) >= 1:
            preserve_cell_format(table.rows[10].cells[0], 'Introduction')
        
        # Row 11: Introduction content - Resources (cells 2-4) and Duration (cell 5)
        if len(table.rows) > 11 and len(table.rows[11].cells) >= 6:
            resources = data.get('resources', '')
            preserve_cell_format(table.rows[11].cells[2], resources)
            preserve_cell_format(table.rows[11].cells[5], '5 minutes')
        
        # Row 12: Development header
        if len(table.rows) > 12 and len(table.rows[12].cells) >= 1:
            preserve_cell_format(table.rows[12].cells[0], 'Development/Body')
        
        # Row 13-15: Development content rows
        for row_num in [13, 14, 15]:
            if len(table.rows) > row_num and len(table.rows[row_num].cells) >= 6:
                resources = data.get('resources', '')
                preserve_cell_format(table.rows[row_num].cells[2], resources)
                preserve_cell_format(table.rows[row_num].cells[5], '25 minutes')
        
        # Row 16: Conclusion header
        if len(table.rows) > 16 and len(table.rows[16].cells) >= 1:
            preserve_cell_format(table.rows[16].cells[0], 'Conclusion')
        
        # Row 17: Conclusion content
        if len(table.rows) > 17 and len(table.rows[17].cells) >= 6:
            preserve_cell_format(table.rows[17].cells[2], data.get('resources', ''))
            preserve_cell_format(table.rows[17].cells[5], '3 minutes')
        
        # Row 18: Assessment/Assignment
        if len(table.rows) > 18 and len(table.rows[18].cells) >= 6:
            preserve_cell_format(table.rows[18].cells[2], 'Assessment sheets')
            preserve_cell_format(table.rows[18].cells[5], '5 minutes')
        
        # Row 19: Evaluation of the session
        if len(table.rows) > 19 and len(table.rows[19].cells) >= 6:
            preserve_cell_format(table.rows[19].cells[2], 'Self-assessment form')
            preserve_cell_format(table.rows[19].cells[5], '2 minutes')
        
        # Row 20: References
        if len(table.rows) > 20 and len(table.rows[20].cells) >= 1:
            refs = _generate_apa_references(
                data.get('module_code_title', ''),
                data.get('topic_of_session', ''),
                data.get('learning_outcomes', ''),
                data.get('indicative_contents', '')
            )
            preserve_cell_format(table.rows[20].cells[0], refs)
        
        # Row 21: Appendices
        if len(table.rows) > 21 and len(table.rows[21].cells) >= 1:
            preserve_cell_format(table.rows[21].cells[0], 'Appendices: PPT, Task Sheets, Assessment tools, Materials')
        
        # Row 22: Reflection
        if len(table.rows) > 22 and len(table.rows[22].cells) >= 1:
            preserve_cell_format(table.rows[22].cells[0], 'Reflection: Space for trainer to reflect on session effectiveness and improvements')
            
    except Exception as e:
        print(f"Error filling RTB template table: {e}")
    
    return doc

def _create_rtb_session_plan_from_scratch(data):
    """Create RTB session plan document from scratch if template not available"""
    doc = Document()
    
    # Set margins to RTB standard
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Title
    title = doc.add_heading('RWANDA TECHNICAL BOARD\nSESSION PLAN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_runs = title.runs
    for run in title_runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(14)
    
    doc.add_paragraph()
    
    # Header Information Table
    header_table = doc.add_table(rows=4, cols=2)
    header_table.style = 'Table Grid'
    
    header_data = [
        ('Sector', data.get('sector', '')),
        ('Sub-Sector/Trade', data.get('trade', '')),
        ('Module', data.get('module_code_title', '')),
        ('Trainer Name', data.get('trainer_name', '')),
    ]
    
    for idx, (label, value) in enumerate(header_data):
        preserve_cell_format(header_table.rows[idx].cells[0], label)
        preserve_cell_format(header_table.rows[idx].cells[1], str(value))
    
    doc.add_paragraph()
    
    # Session Information Table
    session_table = doc.add_table(rows=6, cols=2)
    session_table.style = 'Table Grid'
    
    session_data = [
        ('Topic of Session', data.get('topic_of_session', '')),
        ('Duration (minutes)', str(data.get('duration', ''))),
        ('Week', data.get('week', '')),
        ('Term', data.get('term', '')),
        ('Date', data.get('date', '')),
        ('Number of Trainees', str(data.get('number_of_trainees', ''))),
    ]
    
    for idx, (label, value) in enumerate(session_data):
        preserve_cell_format(session_table.rows[idx].cells[0], label)
        preserve_cell_format(session_table.rows[idx].cells[1], value)
    
    doc.add_paragraph()
    
    # Learning Outcomes
    doc.add_heading('Learning Outcomes', level=1)
    outcomes_para = doc.add_paragraph(data.get('learning_outcomes', ''))
    outcomes_para.paragraph_format.line_spacing = 1.5
    for run in outcomes_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Facilitation Techniques
    doc.add_heading('Facilitation Techniques', level=1)
    facilitation_para = doc.add_paragraph(data.get('facilitation_techniques', ''))
    facilitation_para.paragraph_format.line_spacing = 1.5
    for run in facilitation_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Indicative Contents
    doc.add_heading('Indicative Contents', level=1)
    contents_para = doc.add_paragraph(data.get('indicative_contents', ''))
    contents_para.paragraph_format.line_spacing = 1.5
    for run in contents_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Generate and add intelligent content sections
    topic = data.get('topic_of_session', '')
    facilitation = data.get('facilitation_techniques', 'Trainer-guided instruction')
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    intro_content = SmartSessionPlanContentGenerator.generate_introduction_section(topic, facilitation)
    intro_para = doc.add_paragraph(intro_content)
    intro_para.paragraph_format.line_spacing = 1.5
    for run in intro_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Development
    doc.add_heading('Development', level=1)
    dev_content = SmartSessionPlanContentGenerator.generate_development_section(topic, facilitation)
    dev_para = doc.add_paragraph(dev_content)
    dev_para.paragraph_format.line_spacing = 1.5
    for run in dev_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    conclusion_content = SmartSessionPlanContentGenerator.generate_conclusion_section(topic)
    conclusion_para = doc.add_paragraph(conclusion_content)
    conclusion_para.paragraph_format.line_spacing = 1.5
    for run in conclusion_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Assessment
    doc.add_heading('Assessment', level=1)
    assessment_content = SmartSessionPlanContentGenerator.generate_assessment_section(topic, facilitation)
    assessment_para = doc.add_paragraph(assessment_content)
    assessment_para.paragraph_format.line_spacing = 1.5
    for run in assessment_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Evaluation
    doc.add_heading('Evaluation', level=1)
    evaluation_content = SmartSessionPlanContentGenerator.generate_evaluation_section(topic)
    evaluation_para = doc.add_paragraph(evaluation_content)
    evaluation_para.paragraph_format.line_spacing = 1.5
    for run in evaluation_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # Resources
    doc.add_heading('Resources', level=1)
    resources_text = data.get('resources', '')
    if resources_text:
        for line in resources_text.split('\n'):
            if line.strip():
                resource_para = doc.add_paragraph(f"• {line.strip()}", style='List Bullet')
                resource_para.paragraph_format.line_spacing = 1.5
                for run in resource_para.runs:
                    run.font.name = 'Book Antiqua'
                    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    # References
    doc.add_heading('References', level=1)
    references = _generate_apa_references(
        data.get('module_code_title', ''),
        topic,
        data.get('learning_outcomes', ''),
        data.get('indicative_contents', '')
    )
    ref_para = doc.add_paragraph(references)
    ref_para.paragraph_format.line_spacing = 1.5
    for run in ref_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    return doc

def _generate_apa_references(module, topic, learning_outcomes, indicative_contents):
    """Generate APA-formatted references based on content"""
    content_lower = f"{topic} {module}".lower() if topic and module else ""
    references = []
    
    # Programming/ICT references
    if any(term in content_lower for term in ['programming', 'python', 'java', 'code', 'software', 'algorithm', 'loop', 'array', 'function', 'c programming']):
        references = [
            "Deitel, P., & Deitel, H. (2019). Python for programmers. Pearson Education.",
            "McConnell, S. (2004). Code complete: A practical handbook of software construction. Microsoft Press.",
            "Matthes, E. (2019). Python crash course: A hands-on, project-based introduction to programming. No Starch Press.",
            "Downey, A. (2015). Think Python: How to think like a computer scientist. O'Reilly Media.",
            "Hunt, A., & Thomas, D. (2019). The pragmatic programmer (2nd ed.). Addison-Wesley."
        ]
    # Networking references
    elif any(term in content_lower for term in ['network', 'cisco', 'routing', 'switching', 'tcp', 'ip', 'internet', 'firewall']):
        references = [
            "Tanenbaum, A. S., & Wetherall, D. J. (2021). Computer networks (6th ed.). Pearson Education.",
            "Kurose, J. F., & Ross, K. W. (2020). Computer networking: A top-down approach (8th ed.). Pearson.",
            "Odom, W. (2019). CCNA 200-301 official cert guide library (2nd ed.). Cisco Press.",
            "Cisco Networking Academy. (2020). CCNA routing and switching course materials. Cisco Systems.",
            "Doyle, J. C., Alderson, D. L., & Willinger, W. (2015). Internet topology and the evolution of the Internet. ACM Transactions."
        ]
    # Database references
    elif any(term in content_lower for term in ['database', 'sql', 'mysql', 'data management', 'data modeling', 'nosql']):
        references = [
            "Elmasri, R., & Navathe, S. B. (2020). Fundamentals of database systems (8th ed.). Pearson Education.",
            "Coronel, C., & Morris, S. (2019). Database systems: Design, implementation, and management (13th ed.). Cengage.",
            "Beaulieu, A. (2020). Learning SQL: Generate, manipulate, and retrieve data (3rd ed.). O'Reilly Media.",
            "Garcia-Molina, H., Ullman, J. D., & Widom, J. (2008). Database systems: The complete book (2nd ed.). Prentice Hall.",
            "DuBois, P. (2020). MySQL cookbook: Solutions for database developers and administrators. O'Reilly."
        ]
    # Web development references
    elif any(term in content_lower for term in ['web', 'html', 'css', 'javascript', 'frontend', 'responsive', 'react', 'vue']):
        references = [
            "Duckett, J. (2014). HTML and CSS: Design and build websites. Wiley Publishing.",
            "Flanagan, D. (2020). JavaScript: The definitive guide (7th ed.). O'Reilly Media.",
            "Robbins, J. N. (2018). Learning web design: A beginner's guide to HTML, CSS, JavaScript (5th ed.). O'Reilly.",
            "Nielsen, J., & Norman, D. A. (2014). Usability 101: Introduction to usability. Nielsen Norman Group.",
            "Mozilla Foundation. (2023). Web development documentation and standards. Retrieved from https://developer.mozilla.org"
        ]
    # Business/Management references
    elif any(term in content_lower for term in ['business', 'management', 'leadership', 'entrepreneurship', 'accounting', 'finance']):
        references = [
            "Drucker, P. F. (2006). The effective executive: The definitive guide to getting the right things done. Harper Business.",
            "Porter, M. E. (2008). Competitive advantage: Creating and sustaining superior performance. Free Press.",
            "Mintzberg, H. (2009). Managing. Berrett-Koehler Publishers.",
            "Kotter, J. P. (2012). Leading change. Harvard Business Review Press.",
            "Covey, S. R. (2004). The 7 habits of highly effective people. Free Press."
        ]
    else:
        # Default TVET references
        references = [
            "Rwanda Education Board. (2021). TVET curriculum framework. REB Publications.",
            "UNESCO-UNEVOC. (2020). Technical and vocational education and training (TVET) and the sustainable development goals (SDGs). UNESCO Publications.",
            "Ministry of Education Rwanda. (2022). Competency-based training guidelines for technical and vocational education. MINEDUC.",
            "Rwanda Technical and Vocational Education and Training Board. (2023). National module guidelines and standards. RTVETB Publications.",
            "International Labour Organization. (2021). World employment and social outlook: The role of digital labour platforms. ILO."
        ]
    
    # Format as numbered list
    formatted_refs = []
    for i, ref in enumerate(references[:5], 1):
        formatted_refs.append(f"{i}. {ref}")
    
    return '\n\n'.join(formatted_refs)


def fill_scheme_of_work_template_100_percent(data):
    """
    Generate RTB-compliant scheme of work matching official RTB template 100%
    
    Accepts:
        data (dict): User-provided scheme of work data with keys for each term
    
    Returns:
        str: Path to generated DOCX file
    """
    
    try:
        template_path = os.path.join(os.path.dirname(__file__), 'rtb_scheme_template.docx')
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")
        
        doc = Document(template_path)
        
        # Fill each term table
        terms_data = data.get('terms', [])
        for term_idx, table in enumerate(doc.tables):
            if term_idx < len(terms_data):
                _fill_scheme_table(table, terms_data[term_idx])
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
        
    except Exception as e:
        print(f"Error in fill_scheme_of_work_template_100_percent: {e}")
        raise


def _fill_scheme_table(table, term_data):
    """Fill a term table in scheme of work (9 columns)"""
    try:
        # Table structure:
        # Row 0-1: Headers (preserved)
        # Row 2+: Data rows
        
        # Get data items for this term (each item fills one row starting from row 2)
        items = term_data.get('items', [])
        
        for idx, item in enumerate(items):
            row_num = 2 + idx
            if row_num >= len(table.rows):
                break
            
            row = table.rows[row_num]
            
            # Fill cells based on columns
            # [0]: Weeks/Dates
            if len(row.cells) > 0:
                preserve_cell_format(row.cells[0], item.get('weeks', ''))
            
            # [1]: Learning Outcome (LO)
            if len(row.cells) > 1:
                preserve_cell_format(row.cells[1], item.get('learning_outcome', ''))
            
            # [2]: Duration
            if len(row.cells) > 2:
                preserve_cell_format(row.cells[2], item.get('duration', ''))
            
            # [3]: Indicative Content (IC)
            if len(row.cells) > 3:
                preserve_cell_format(row.cells[3], item.get('indicative_content', ''))
            
            # [4]: Learning Activities
            if len(row.cells) > 4:
                preserve_cell_format(row.cells[4], item.get('learning_activities', ''))
            
            # [5]: Resources
            if len(row.cells) > 5:
                preserve_cell_format(row.cells[5], item.get('resources', ''))
            
            # [6]: Evidences of Assessment
            if len(row.cells) > 6:
                preserve_cell_format(row.cells[6], item.get('evidence_of_assessment', ''))
            
            # [7]: Learning Place
            if len(row.cells) > 7:
                preserve_cell_format(row.cells[7], item.get('learning_place', ''))
            
            # [8]: Observation
            if len(row.cells) > 8:
                preserve_cell_format(row.cells[8], item.get('observation', ''))
                
    except Exception as e:
        print(f"Error filling scheme table: {e}")
