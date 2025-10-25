"""RTB Template Filler - Preserves exact formatting, fonts, colspan, rowspan"""
from docx import Document
from docx.shared import Pt
import tempfile
import os
from facilitation_content_generator import (
    generate_introduction_activities,
    generate_development_activities,
    generate_resources,
    generate_assessment
)
from content_formatter import clean_text, format_objectives, format_resources, generate_references

def preserve_cell_format(cell, new_text):
    """Update cell text while preserving all formatting"""
    if not cell.paragraphs:
        cell.text = new_text
        return
    
    # Clear existing text but keep formatting
    para = cell.paragraphs[0]
    if para.runs:
        # Keep first run's formatting
        run = para.runs[0]
        font_name = run.font.name
        font_size = run.font.size
        bold = run.font.bold
        
        # Clear all runs
        for run in para.runs:
            run.text = ''
        
        # Add new text with preserved formatting
        new_run = para.runs[0] if para.runs else para.add_run()
        new_run.text = new_text
        if font_name:
            new_run.font.name = font_name
        if font_size:
            new_run.font.size = font_size
        if bold is not None:
            new_run.font.bold = bold
    else:
        para.text = new_text

def fill_session_plan_template(data):
    """Fill session plan - preserves exact template formatting"""
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_session_plan_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError("RTB Session Plan template not found")
    
    doc = Document(template_path)
    table = doc.tables[0]
    
    # Row 1: Sector, Sub-sector (colspan=3), Date (colspan=2)
    preserve_cell_format(table.rows[1].cells[0], f"Sector :    {data.get('sector', '')}")
    preserve_cell_format(table.rows[1].cells[1], f"Sub-sector: {data.get('trade', '')}")
    preserve_cell_format(table.rows[1].cells[4], f"Date : {data.get('date', '')}")
    
    # Row 2: Trainer (colspan=4), Term (colspan=2)
    preserve_cell_format(table.rows[2].cells[0], f"Lead Trainer's name : {data.get('trainer_name', '')}")
    preserve_cell_format(table.rows[2].cells[4], f"TERM : {data.get('term', '')}")
    
    # Row 3: Module, Week (colspan=2), Trainees, Class (colspan=2)
    preserve_cell_format(table.rows[3].cells[0], f"Module(Code&Name): {data.get('module_code_title', '')}")
    preserve_cell_format(table.rows[3].cells[1], f"Week : {data.get('week', '')}")
    preserve_cell_format(table.rows[3].cells[3], f"No. Trainees: {data.get('number_of_trainees', '')}")
    preserve_cell_format(table.rows[3].cells[4], f"Class(es): {data.get('class_name', '')}")
    
    # Row 4: Learning outcome
    learning_outcomes = clean_text(data.get('learning_outcomes', ''))
    preserve_cell_format(table.rows[4].cells[1], learning_outcomes)
    
    # Row 5: Indicative contents
    indicative_contents = clean_text(data.get('indicative_contents', ''))
    preserve_cell_format(table.rows[5].cells[1], indicative_contents)
    
    # Row 6: Topic
    topic = clean_text(data.get('topic_of_session', ''))
    preserve_cell_format(table.rows[6].cells[0], f"Topic of the session: {topic}")
    
    # Row 7: Range, Duration
    preserve_cell_format(table.rows[7].cells[0], f"Range:\n{indicative_contents}")
    preserve_cell_format(table.rows[7].cells[1], f"Duration of the session: {data.get('duration', '')}min")
    
    # Row 8: Objectives - SMART formatted
    objectives = format_objectives(data.get('objectives', ''))
    preserve_cell_format(table.rows[8].cells[0], f"Objectives: By the end of this session every learner should be able to:\n{objectives}")
    
    # Row 9: Facilitation
    facilitation = clean_text(data.get('facilitation_techniques', ''))
    preserve_cell_format(table.rows[9].cells[0], f"Facilitation technique(s): {facilitation}")
    
    # Row 11: Introduction - clean formatting
    topic = data.get('topic_of_session', '')
    facilitation = data.get('facilitation_techniques', 'Trainer-guided')
    intro = clean_text(generate_introduction_activities(topic, facilitation))
    preserve_cell_format(table.rows[11].cells[0], intro)
    preserve_cell_format(table.rows[11].cells[2], "Attendance sheet\nPPT\nProjector\nComputers\nFlipchart\nWhiteboard\nMarker pen")
    preserve_cell_format(table.rows[11].cells[5], "5 minutes")
    
    # Row 13: Development - based on facilitation technique
    activities = generate_development_activities(topic, facilitation, data.get('learning_activities'))
    preserve_cell_format(table.rows[13].cells[0], clean_text(activities))
    resources = generate_resources(facilitation, data.get('resources'))
    formatted_resources = format_resources(resources)
    preserve_cell_format(table.rows[13].cells[2], formatted_resources)
    duration_main = int(data.get('duration', 40)) - 15
    preserve_cell_format(table.rows[13].cells[5], f"{duration_main}\nminutes")
    
    # Row 17: Conclusion - clean formatting
    conclusion = "Trainer's activity:\n• Involves learners to summarize the session\n• Asks questions reflecting on learning objectives\n• Links to next session\n\nLearner's activity:\n• Summarizes key points learned\n• Responds to questions\n• Asks clarifications"
    preserve_cell_format(table.rows[17].cells[0], conclusion)
    preserve_cell_format(table.rows[17].cells[2], "Computer\nProjector")
    preserve_cell_format(table.rows[17].cells[5], "3 minutes")
    
    # Row 18: Assessment - based on facilitation technique
    assessment = generate_assessment(topic, facilitation, data.get('assessment_details'))
    preserve_cell_format(table.rows[18].cells[0], assessment)
    preserve_cell_format(table.rows[18].cells[2], "Assessment sheets")
    preserve_cell_format(table.rows[18].cells[5], "5 minutes")
    
    # Row 19: Evaluation - clean formatting
    evaluation = "Trainer's activity:\n• Involves learners in session evaluation\n• Asks: How was the session?\n• Notes areas for improvement\n\nLearner's activity:\n• Provides feedback on session\n• Shares learning experience"
    preserve_cell_format(table.rows[19].cells[0], evaluation)
    preserve_cell_format(table.rows[19].cells[2], "Self-assessment form")
    preserve_cell_format(table.rows[19].cells[5], "2 minutes")
    
    # Row 20: References - AI generated based on content
    custom_refs = data.get('references', '').strip()
    if custom_refs and custom_refs != '(To be added by trainer)':
        references = clean_text(custom_refs)
    else:
        # Generate references based on module content
        references = generate_references(
            data.get('module_code_title', ''),
            data.get('topic_of_session', ''),
            data.get('learning_outcomes', ''),
            data.get('indicative_contents', '')
        )
    preserve_cell_format(table.rows[20].cells[0], f"References:\n\n{references}")
    
    # Row 21: Appendices
    preserve_cell_format(table.rows[21].cells[0], "Appendices: PPT, Task Sheets, assessment")
    
    # Row 22: Reflection
    preserve_cell_format(table.rows[22].cells[0], "Reflection:")
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name

def fill_scheme_template(data):
    """Fill scheme - preserves exact template formatting"""
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_scheme_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError("RTB Scheme template not found")
    
    doc = Document(template_path)
    
    # Fill all 3 term tables
    for term_idx in range(min(3, len(doc.tables))):
        term_num = term_idx + 1
        table = doc.tables[term_idx]
        
        if len(table.rows) > 2:
            # Row 2 is the data row
            weeks = clean_text(data.get(f'term{term_num}_weeks', ''))
            outcomes = clean_text(data.get(f'term{term_num}_learning_outcomes', ''))
            duration = clean_text(data.get(f'term{term_num}_duration', ''))
            contents = clean_text(data.get(f'term{term_num}_indicative_contents', ''))
            learning_place = clean_text(data.get(f'term{term_num}_learning_place', ''))
            
            preserve_cell_format(table.rows[2].cells[0], weeks)
            preserve_cell_format(table.rows[2].cells[1], outcomes)
            preserve_cell_format(table.rows[2].cells[2], duration)
            preserve_cell_format(table.rows[2].cells[3], contents)
            # Add learning place if column exists (check table structure)
            if len(table.rows[2].cells) > 4:
                preserve_cell_format(table.rows[2].cells[4], learning_place)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name
