"""
Test document generation and compare with official RTB templates
"""
from document_generator import generate_session_plan_docx, generate_scheme_of_work_docx
from docx import Document
import os
from datetime import datetime

def test_session_plan():
    """Test session plan generation with sample data"""
    print("\n" + "="*60)
    print("TESTING SESSION PLAN GENERATION")
    print("="*60 + "\n")
    
    # Sample data that a teacher would input
    test_data = {
        'sector': 'ICT',
        'trade': 'Computer Hardware Maintenance',
        'rqf_level': 'Level 4',
        'trainer_name': 'John MUGISHA',
        'module_code_title': 'CHM4101 - Computer Assembly',
        'term': 'Term 1',
        'week': 'Week 5',
        'date': '15/01/2025',
        'class_name': 'CHM4A',
        'number_of_trainees': '25',
        'learning_outcomes': 'By the end of this module, learners will be able to assemble desktop computers correctly',
        'indicative_contents': 'Computer components, Assembly procedures, Safety precautions, Testing procedures',
        'topic_of_session': 'Installing Motherboard and Power Supply',
        'duration': '40',
        'objectives': '1. Identify motherboard components\n2. Install motherboard correctly\n3. Connect power supply safely',
        'facilitation_techniques': 'Demonstration, Hands-on practice, Group work',
        'learning_activities': 'Students will practice installing motherboards in computer cases under supervision',
        'resources': 'Computer cases, Motherboards, Power supplies, Screwdrivers, Anti-static wristbands',
        'assessment_details': 'Practical assessment: Students assemble a complete system\nWritten test: Safety procedures',
        'references': 'Computer Hardware Maintenance Manual, RTB Curriculum Guide'
    }
    
    print("Test Data:")
    for key, value in test_data.items():
        print(f"  {key}: {value[:50]}..." if len(str(value)) > 50 else f"  {key}: {value}")
    
    print("\n[*] Generating session plan...")
    try:
        output_file = generate_session_plan_docx(test_data)
        print(f"[OK] Generated: {output_file}")
        
        # Analyze generated document
        doc = Document(output_file)
        print(f"\n[*] Generated document has:")
        print(f"    - {len(doc.tables)} table(s)")
        print(f"    - {len(doc.paragraphs)} paragraph(s)")
        
        if len(doc.tables) > 0:
            table = doc.tables[0]
            print(f"    - Main table: {len(table.rows)} rows x {len(table.columns)} columns")
            
            # Check if key data is present
            print(f"\n[*] Checking if user data appears in document...")
            doc_text = ""
            for row in table.rows:
                for cell in row.cells:
                    doc_text += cell.text + " "
            
            checks = {
                'Sector (ICT)': 'ICT' in doc_text,
                'Trainer name': 'MUGISHA' in doc_text,
                'Module code': 'CHM4101' in doc_text,
                'Topic': 'Installing Motherboard' in doc_text,
                'Learning outcomes': 'assemble desktop computers' in doc_text,
                'Duration': '40' in doc_text
            }
            
            for check_name, result in checks.items():
                status = "[OK]" if result else "[FAIL]"
                print(f"    {status} {check_name}")
        
        # Save with descriptive name
        test_output = os.path.join(os.path.dirname(__file__), 
                                   f"TEST_Session_Plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(test_output)
        print(f"\n[OK] Test file saved: {test_output}")
        print(f"\n[ACTION REQUIRED] Please:")
        print(f"    1. Open: {test_output}")
        print(f"    2. Open: rtb_session_plan_template.docx")
        print(f"    3. Compare them side-by-side")
        print(f"    4. Check if ALL data is in correct positions")
        
        return True
        
    except Exception as e:
        print(f"[ERROR] Generation failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_scheme_of_work():
    """Test scheme of work generation with sample data"""
    print("\n" + "="*60)
    print("TESTING SCHEME OF WORK GENERATION")
    print("="*60 + "\n")
    
    # Sample data for scheme
    test_data = {
        'province': 'Kigali City',
        'district': 'Gasabo',
        'sector': 'ICT',
        'school': 'IPRC Kigali',
        'department_trade': 'Computer Hardware Maintenance',
        'qualification_title': 'Advanced Diploma in Computer Hardware Maintenance',
        'rqf_level': 'Level 4',
        'module_code_title': 'CHM4101 - Computer Assembly and Maintenance',
        'school_year': '2024-2025',
        'terms': '3',
        'module_hours': '120',
        'number_of_classes': '2',
        'class_name': 'CHM4A, CHM4B',
        'trainer_name': 'John MUGISHA',
        
        # Term 1
        'term1_weeks': 'Week 1-12 (Sept 8 - Dec 19, 2024)',
        'term1_learning_outcomes': 'LO1: Prepare tools, materials and equipment\nLO2: Assemble desktop computers',
        'term1_indicative_contents': 'IC1.1: Workplace preparation\nIC1.2: Tool identification\nIC2.1: Computer components\nIC2.2: Assembly procedures',
        'term1_duration': '40 hours',
        
        # Term 2
        'term2_weeks': 'Week 13-24 (Jan 5 - Apr 3, 2025)',
        'term2_learning_outcomes': 'LO3: Install operating systems\nLO4: Configure computer systems',
        'term2_indicative_contents': 'IC3.1: OS installation\nIC3.2: Driver installation\nIC4.1: BIOS configuration\nIC4.2: System optimization',
        'term2_duration': '40 hours',
        
        # Term 3
        'term3_weeks': 'Week 25-36 (Apr 20 - Jul 3, 2025)',
        'term3_learning_outcomes': 'LO5: Troubleshoot hardware issues\nLO6: Maintain computer systems',
        'term3_indicative_contents': 'IC5.1: Diagnostic procedures\nIC5.2: Component replacement\nIC6.1: Preventive maintenance\nIC6.2: Documentation',
        'term3_duration': '40 hours',
        
        'dos_name': 'Dr. UWASE Marie',
        'manager_name': 'Eng. NKUSI Patrick'
    }
    
    print("Test Data:")
    print(f"  School: {test_data['school']}")
    print(f"  Module: {test_data['module_code_title']}")
    print(f"  Trainer: {test_data['trainer_name']}")
    print(f"  Terms: {test_data['terms']}")
    
    print("\n[*] Generating scheme of work...")
    try:
        output_file = generate_scheme_of_work_docx(test_data)
        print(f"[OK] Generated: {output_file}")
        
        # Analyze generated document
        doc = Document(output_file)
        print(f"\n[*] Generated document has:")
        print(f"    - {len(doc.tables)} table(s)")
        print(f"    - {len(doc.paragraphs)} paragraph(s)")
        
        # Check each table
        for idx, table in enumerate(doc.tables):
            print(f"    - Table {idx+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        
        # Save with descriptive name
        test_output = os.path.join(os.path.dirname(__file__), 
                                   f"TEST_Scheme_of_Work_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        doc.save(test_output)
        print(f"\n[OK] Test file saved: {test_output}")
        print(f"\n[ACTION REQUIRED] Please:")
        print(f"    1. Open: {test_output}")
        print(f"    2. Open: rtb_scheme_template.docx")
        print(f"    3. Compare them side-by-side")
        print(f"    4. Check if ALL 3 terms are filled correctly")
        
        return True
        
    except Exception as e:
        print(f"[ERROR] Generation failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    print("\n" + "="*60)
    print("RTB DOCUMENT GENERATION TEST SUITE")
    print("="*60)
    
    # Test session plan
    session_result = test_session_plan()
    
    # Test scheme of work
    scheme_result = test_scheme_of_work()
    
    # Summary
    print("\n" + "="*60)
    print("TEST SUMMARY")
    print("="*60)
    print(f"Session Plan: {'[PASS]' if session_result else '[FAIL]'}")
    print(f"Scheme of Work: {'[PASS]' if scheme_result else '[FAIL]'}")
    
    if session_result and scheme_result:
        print("\n[OK] Both documents generated successfully!")
        print("[ACTION] Now manually compare with official templates")
    else:
        print("\n[ERROR] Some tests failed. Check errors above.")
    
    print("\n" + "="*60)

if __name__ == "__main__":
    main()
