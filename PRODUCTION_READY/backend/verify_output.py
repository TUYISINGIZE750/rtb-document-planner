"""Verify the actual generated document"""
from official_template_filler import fill_scheme_official
from docx import Document

test_data = {
    'province': 'Kigali City',
    'district': 'Gasabo',
    'sector': 'ICT & MULTIMEDIA',
    'school': 'IPRC Kigali',
    'trainer_name': 'John Doe',
    'department_trade': 'ICT Department',
    'module_code_title': 'ICT101 - Web Development',
    'school_year': '2024-2025',
    'module_hours': '120',
    'qualification_title': 'Advanced Diploma in ICT',
    'terms': 'Term 1, 2, 3',
    'rqf_level': 'Level 5',
    'number_of_classes': '3',
    'class_name': 'Year 2 ICT',
    
    'term1_weeks': '1-4',
    'term1_learning_outcomes': 'LO1: Create HTML pages\nLO2: Style with CSS\nLO3: Add JavaScript interactivity',
    'term1_indicative_contents': 'IC1: HTML5 tags and structure\nIC2: CSS selectors and properties\nIC3: JavaScript basics and DOM',
    'term1_duration': '40 hours',
    'term1_learning_place': 'Computer Lab',
    
    'term2_weeks': '5-8',
    'term2_learning_outcomes': 'LO4: Build responsive layouts\nLO5: Implement forms',
    'term2_indicative_contents': 'IC4: Flexbox and Grid\nIC5: Form validation',
    'term2_duration': '40 hours',
    'term2_learning_place': 'Computer Lab',
    
    'term3_weeks': '9-12',
    'term3_learning_outcomes': 'LO6: Deploy web applications',
    'term3_indicative_contents': 'IC6: Hosting and deployment',
    'term3_duration': '40 hours',
    'term3_learning_place': 'Computer Lab'
}

print("Generating document...")
output_file = fill_scheme_official(test_data)

# Read the generated document
print("\nReading generated document...")
doc = Document(output_file)

# Check Term 1 table by reading XML
print("\n=== TERM 1 TABLE - XML LEVEL CHECK ===")
t1 = doc.tables[1]
for i in range(2, 5):
    row = t1.rows[i]
    cell1 = row.cells[1]
    
    # Get all text from all paragraphs
    all_text = []
    for para in cell1.paragraphs:
        all_text.append(para.text)
    
    print(f"Row {i}:")
    print(f"  cell.text = '{cell1.text}'")
    print(f"  Paragraphs: {all_text}")
    print(f"  Cell ID: {id(cell1._element)}")

print(f"\nGenerated file: {output_file}")
print("Open this file in Word to verify the actual content!")
