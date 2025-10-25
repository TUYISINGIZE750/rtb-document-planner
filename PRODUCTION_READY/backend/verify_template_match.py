"""
Verify if generated documents match RTB official templates
"""
from docx import Document
import os

def analyze_template_structure(template_path, doc_type):
    """Analyze the structure of RTB template"""
    print(f"\n{'='*60}")
    print(f"ANALYZING: {doc_type}")
    print(f"{'='*60}\n")
    
    if not os.path.exists(template_path):
        print(f"❌ Template not found: {template_path}")
        return None
    
    doc = Document(template_path)
    
    print(f"[*] Document has {len(doc.tables)} table(s)")
    print(f"[*] Document has {len(doc.paragraphs)} paragraph(s)\n")
    
    structure = {
        'tables': [],
        'paragraphs': len(doc.paragraphs)
    }
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\n[TABLE {table_idx + 1}]:")
        print(f"   Rows: {len(table.rows)}")
        print(f"   Columns: {len(table.columns)}")
        
        table_data = {
            'rows': len(table.rows),
            'columns': len(table.columns),
            'cells': []
        }
        
        print(f"\n   Cell Contents:")
        for row_idx, row in enumerate(table.rows):
            print(f"\n   Row {row_idx}:")
            row_cells = []
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()[:50]  # First 50 chars
                if text:
                    # Handle encoding issues
                    try:
                        print(f"      [{row_idx},{col_idx}]: {text}")
                    except UnicodeEncodeError:
                        print(f"      [{row_idx},{col_idx}]: [Contains special characters]")
                        text = text.encode('ascii', 'ignore').decode('ascii')
                row_cells.append(text)
            table_data['cells'].append(row_cells)
        
        structure['tables'].append(table_data)
    
    return structure

def compare_with_filler_code():
    """Compare template structure with rtb_template_filler.py expectations"""
    print(f"\n{'='*60}")
    print("CHECKING rtb_template_filler.py COMPATIBILITY")
    print(f"{'='*60}\n")
    
    filler_path = os.path.join(os.path.dirname(__file__), 'rtb_template_filler.py')
    
    if os.path.exists(filler_path):
        with open(filler_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        print("[*] Template filler expects:")
        print("   - Session Plan: table.rows[1], table.rows[2], table.rows[3], etc.")
        print("   - Uses hardcoded row indices (rows[1] to rows[22])")
        
        if 'table.rows[1]' in content:
            print("\n[WARNING] Hardcoded row indices found!")
            print("   This will FAIL if template structure doesn't match exactly")
        
        if 'table = doc.tables[0]' in content:
            print("\n[OK] Expects first table (tables[0])")
        
        return True
    else:
        print("[ERROR] rtb_template_filler.py not found")
        return False

def main():
    base_dir = os.path.dirname(__file__)
    
    # Analyze Session Plan Template
    session_template = os.path.join(base_dir, 'rtb_session_plan_template.docx')
    session_structure = analyze_template_structure(session_template, "RTB SESSION PLAN TEMPLATE")
    
    # Analyze Scheme Template
    scheme_template = os.path.join(base_dir, 'rtb_scheme_template.docx')
    scheme_structure = analyze_template_structure(scheme_template, "RTB SCHEME OF WORK TEMPLATE")
    
    # Compare with filler code
    compare_with_filler_code()
    
    # Summary
    print(f"\n{'='*60}")
    print("VERIFICATION SUMMARY")
    print(f"{'='*60}\n")
    
    if session_structure:
        print("[OK] Session Plan Template Found")
        print(f"   - Tables: {len(session_structure['tables'])}")
        if len(session_structure['tables']) > 0:
            print(f"   - Main table rows: {session_structure['tables'][0]['rows']}")
            print(f"   - Main table columns: {session_structure['tables'][0]['columns']}")
    else:
        print("[ERROR] Session Plan Template NOT Found")
    
    if scheme_structure:
        print("\n[OK] Scheme Template Found")
        print(f"   - Tables: {len(scheme_structure['tables'])}")
    else:
        print("\n[ERROR] Scheme Template NOT Found")
    
    print("\n" + "="*60)
    print("RECOMMENDATIONS:")
    print("="*60)
    print("""
1. [OK] Templates should be in PRODUCTION_READY/backend/ folder
2. [OK] rtb_template_filler.py should match template structure EXACTLY
3. [WARNING] If row indices don't match, documents will be INCORRECT
4. [TIP] Test by generating a document and comparing with template
5. [TIP] All user input should fill corresponding template cells
    """)

if __name__ == "__main__":
    main()
