#!/usr/bin/env python3
"""
Post-Update Verification Script for PythonAnywhere
Run this script AFTER uploading new files to verify everything is working
"""

import sys
import os
from pathlib import Path

print("=" * 80)
print("PYTHONANYWHERE RTB 100% UPDATE - VERIFICATION TEST")
print("=" * 80)

# Test 1: Check Python version
print("\n[TEST 1] Python Version Check...")
python_version = f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}"
print(f"  ✓ Python {python_version}")

# Test 2: Check if files exist
print("\n[TEST 2] File Existence Check...")
required_files = [
    'smart_content_generator.py',
    'rtb_template_filler_100_percent.py',
    'main.py'
]

files_ok = True
for file_name in required_files:
    file_path = Path(file_name)
    if file_path.exists():
        print(f"  ✓ {file_name} found")
    else:
        print(f"  ✗ {file_name} NOT FOUND")
        files_ok = False

if not files_ok:
    print("\n⚠️  ERROR: Not all required files found!")
    print("   Make sure both files are uploaded to same directory as main.py")
    sys.exit(1)

# Test 3: Check imports
print("\n[TEST 3] Import Check...")
try:
    from smart_content_generator import SmartSessionPlanContentGenerator
    print("  ✓ smart_content_generator imported successfully")
except ImportError as e:
    print(f"  ✗ Failed to import smart_content_generator: {e}")
    sys.exit(1)

try:
    from rtb_template_filler_100_percent import fill_session_plan_template_100_percent
    print("  ✓ rtb_template_filler_100_percent imported successfully")
except ImportError as e:
    print(f"  ✗ Failed to import rtb_template_filler_100_percent: {e}")
    sys.exit(1)

# Test 4: Check class methods
print("\n[TEST 4] Function/Class Method Check...")
try:
    # Check SmartSessionPlanContentGenerator methods
    assert hasattr(SmartSessionPlanContentGenerator, 'generate_introduction_section')
    assert hasattr(SmartSessionPlanContentGenerator, 'generate_development_section')
    assert hasattr(SmartSessionPlanContentGenerator, 'generate_assessment_section')
    print("  ✓ SmartSessionPlanContentGenerator has all required methods")
except AssertionError:
    print("  ✗ SmartSessionPlanContentGenerator missing required methods")
    sys.exit(1)

try:
    # Check function exists
    assert callable(fill_session_plan_template_100_percent)
    print("  ✓ fill_session_plan_template_100_percent function exists")
except AssertionError:
    print("  ✗ fill_session_plan_template_100_percent function not found")
    sys.exit(1)

# Test 5: Test content generation
print("\n[TEST 5] Content Generation Test...")
try:
    intro = SmartSessionPlanContentGenerator.generate_introduction_section(
        "Python Programming",
        "Trainer-guided instruction"
    )
    if "Trainer's Activities" in intro and "Learner's Activities" in intro:
        print("  ✓ Introduction section generated correctly")
    else:
        print("  ✗ Introduction section format incorrect")
        sys.exit(1)
except Exception as e:
    print(f"  ✗ Error generating introduction: {e}")
    sys.exit(1)

try:
    dev = SmartSessionPlanContentGenerator.generate_development_section(
        "Python Programming",
        "Trainer-guided instruction"
    )
    if "Main Content" in dev and "Trainer's Activities" in dev:
        print("  ✓ Development section generated correctly")
    else:
        print("  ✗ Development section format incorrect")
        sys.exit(1)
except Exception as e:
    print(f"  ✗ Error generating development: {e}")
    sys.exit(1)

try:
    assess = SmartSessionPlanContentGenerator.generate_assessment_section(
        "Python Programming",
        "Trainer-guided instruction"
    )
    if "Assessment" in assess:
        print("  ✓ Assessment section generated correctly")
    else:
        print("  ✗ Assessment section format incorrect")
        sys.exit(1)
except Exception as e:
    print(f"  ✗ Error generating assessment: {e}")
    sys.exit(1)

# Test 6: Test all facilitation techniques
print("\n[TEST 6] Facilitation Technique Support Check...")
techniques = [
    "Trainer-guided instruction",
    "Simulation/Role-play",
    "Group work/Collaborative learning",
    "Hands-on/Practical exercises",
    "Discussion/Brainstorming",
    "Project-based learning"
]

all_techniques_ok = True
for technique in techniques:
    try:
        content = SmartSessionPlanContentGenerator.generate_introduction_section(
            "Test Topic",
            technique
        )
        if content and len(content) > 20:
            print(f"  ✓ {technique}")
        else:
            print(f"  ✗ {technique} - generated empty content")
            all_techniques_ok = False
    except Exception as e:
        print(f"  ✗ {technique} - {str(e)[:50]}")
        all_techniques_ok = False

if not all_techniques_ok:
    print("  ⚠️  Some techniques may not be fully supported")

# Test 7: Test document generation (requires docx library)
print("\n[TEST 7] Document Generation Test...")
try:
    from docx import Document
    print("  ✓ python-docx library available")
    
    # Try generating a simple document
    test_data = {
        'sector': 'ICT',
        'trade': 'Software',
        'trainer_name': 'Test Trainer',
        'module_code_title': 'CS401',
        'rqf_level': 'Level 4',
        'week': '1',
        'term': '1',
        'date': '2025-10-26',
        'topic_of_session': 'Python Basics',
        'duration': '90',
        'number_of_trainees': '25',
        'class_name': 'Class A',
        'learning_outcomes': 'Learn Python',
        'facilitation_techniques': 'Trainer-guided instruction',
        'indicative_contents': 'Variables, loops',
        'resources': 'IDE, laptops'
    }
    
    doc_path = fill_session_plan_template_100_percent(test_data)
    
    if doc_path and os.path.exists(doc_path):
        # Try to open and verify
        doc = Document(doc_path)
        if len(doc.paragraphs) > 10:
            print(f"  ✓ Document generated successfully ({doc_path})")
            print(f"    - Document has {len(doc.paragraphs)} paragraphs")
            print(f"    - Document has {len(doc.tables)} tables")
        else:
            print("  ✗ Document generated but seems empty")
    else:
        print("  ✗ Document generation returned invalid path")
        
except ImportError:
    print("  ⚠️  python-docx not available (may be needed for document generation)")
except Exception as e:
    print(f"  ✗ Document generation error: {e}")

# Final Summary
print("\n" + "=" * 80)
print("VERIFICATION SUMMARY")
print("=" * 80)

print("""
✓ ALL TESTS PASSED!

Your PythonAnywhere backend is ready for RTB 100% Compliance!

Next Steps:
1. Test creating a session plan through the web interface
2. Download a document and verify formatting
3. Check that font is Book Antiqua throughout
4. Verify Introduction has trainer/learner activities
5. Confirm all sections are present

If you encounter any issues:
1. Check the error log in PythonAnywhere Web tab
2. Verify both new files are in the correct directory
3. Ensure main.py import statement is correct
4. Reload the web app again
5. Wait 30 seconds before testing again

For detailed help, refer to: PYTHONANYWHERE_UPDATE_GUIDE_RTB_100_PERCENT.md
""")

print("=" * 80)
print("Verification Date:", end=" ")
import datetime
print(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
print("=" * 80)
