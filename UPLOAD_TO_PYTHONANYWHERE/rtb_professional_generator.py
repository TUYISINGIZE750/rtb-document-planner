from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def merge_cells_in_row(table, row_idx, start_col, end_col):
    """Merge cells horizontally"""
    cell = table.rows[row_idx].cells[start_col]
    for col in range(start_col + 1, end_col + 1):
        cell.merge(table.rows[row_idx].cells[col])
    return cell

def generate_rtb_session_plan(data):
    """Generate RTB-compliant session plan matching official template"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Main table: 23 rows x 6 columns
    table = doc.add_table(rows=23, cols=6)
    table.style = 'Table Grid'
    
    # Row 0: Empty header row
    
    # Row 1: Sector, Sub-sector, Date
    merge_cells_in_row(table, 1, 1, 3).text = f"Sub-sector: {data.get('sub_sector', 'N/A')}"
    merge_cells_in_row(table, 1, 4, 5).text = f"Date : {data.get('date', 'N/A')}"
    table.rows[1].cells[0].text = f"Sector :    {data.get('sector', 'N/A')}"
    
    # Row 2: Trainer name, Term
    merge_cells_in_row(table, 2, 0, 3).text = f"Lead Trainer's name : {data.get('trainer_name', 'N/A')}"
    merge_cells_in_row(table, 2, 4, 5).text = f"TERM : {data.get('term', 'I')}"
    
    # Row 3: Module, Week, Trainees, Class
    table.rows[3].cells[0].text = f"Module(Code&Name): {data.get('module_code', 'N/A')}"
    merge_cells_in_row(table, 3, 1, 2).text = f"Week : {data.get('week', 'I')}"
    table.rows[3].cells[3].text = f"No. Trainees: {data.get('num_trainees', '0')}"
    merge_cells_in_row(table, 3, 4, 5).text = f"Class(es): {data.get('classes', '1')}"
    
    # Row 4: Learning outcome
    table.rows[4].cells[0].text = "Learning outcome:"
    merge_cells_in_row(table, 4, 1, 5).text = data.get('learning_outcome', 'N/A')
    
    # Row 5: Indicative contents
    table.rows[5].cells[0].text = "Indicative contents:"
    merge_cells_in_row(table, 5, 1, 5).text = data.get('indicative_content', 'N/A')
    
    # Row 6: Topic
    merge_cells_in_row(table, 6, 0, 5).text = f"Topic of the session: {data.get('topic_of_session', 'N/A')}"
    
    # Row 7: Range and Duration
    table.rows[7].cells[0].text = f"Range: \n{data.get('range', 'N/A')}"
    merge_cells_in_row(table, 7, 1, 5).text = f"Duration of the session: {data.get('duration', '40')}min"
    
    # Row 8: Objectives
    merge_cells_in_row(table, 8, 0, 5).text = f"Objectives: By the end of this session every learner should be able to:\n{data.get('objectives', 'N/A')}"
    
    # Row 9: Facilitation technique
    merge_cells_in_row(table, 9, 0, 5).text = f"Facilitation technique(s):   {data.get('facilitation_technique', 'N/A')}"
    
    # Row 10: Headers for Introduction
    table.rows[10].cells[0].text = "Introduction"
    merge_cells_in_row(table, 10, 1, 1).text = "Introduction"
    merge_cells_in_row(table, 10, 2, 4).text = "Resources"
    table.rows[10].cells[5].text = "Duration"
    
    # Row 11: Introduction content
    merge_cells_in_row(table, 11, 0, 1).text = f"Trainer's activity: \n{data.get('intro_trainer_activity', 'Greets and Make roll calls')}"
    merge_cells_in_row(table, 11, 2, 4).text = data.get('intro_resources', 'Attendance sheet\nPPT\nProjector')
    table.rows[11].cells[5].text = data.get('intro_duration', '5 minutes')
    
    # Row 12: Development header
    merge_cells_in_row(table, 12, 0, 5).text = "Development/Body"
    
    # Rows 13-15: Development steps
    for i, step_num in enumerate([1, 2, 3]):
        row_idx = 13 + i
        step_data = data.get(f'step{step_num}', {})
        merge_cells_in_row(table, row_idx, 0, 1).text = f"Step {step_num}: {step_data.get('title', 'N/A')}\nTrainer's activity: \n{step_data.get('activity', 'N/A')}"
        merge_cells_in_row(table, row_idx, 2, 4).text = step_data.get('resources', 'Computer\nprojector\nPPT')
        table.rows[row_idx].cells[5].text = step_data.get('duration', '25\nminutes')
    
    # Row 16: Conclusion header
    merge_cells_in_row(table, 16, 0, 5).text = "Conclusion"
    
    # Row 17: Summary
    merge_cells_in_row(table, 17, 0, 1).text = f"Summary:\n{data.get('summary', 'The trainer involves the learners to summarize the session')}"
    merge_cells_in_row(table, 17, 2, 4).text = data.get('summary_resources', 'Computer\nprojector')
    table.rows[17].cells[5].text = data.get('summary_duration', '3 minutes')
    
    # Row 18: Assessment
    merge_cells_in_row(table, 18, 0, 1).text = f"Assessment/Assignment\nTrainer's activity: \n{data.get('assessment_activity', 'Trainer gives learners assessment')}"
    merge_cells_in_row(table, 18, 2, 4).text = data.get('assessment_resources', 'Assessment sheets')
    table.rows[18].cells[5].text = data.get('assessment_duration', '5 minutes')
    
    # Row 19: Evaluation
    merge_cells_in_row(table, 19, 0, 1).text = f"Evaluation of the session:\nTrainer's activity: \n{data.get('evaluation_activity', 'Trainer involves learners in evaluation')}"
    merge_cells_in_row(table, 19, 2, 4).text = data.get('evaluation_resources', 'Self-assessment form')
    table.rows[19].cells[5].text = data.get('evaluation_duration', '2minutes')
    
    # Row 20: References
    merge_cells_in_row(table, 20, 0, 5).text = f"References:\n{data.get('references', 'N/A')}"
    
    # Row 21: Appendices
    merge_cells_in_row(table, 21, 0, 5).text = f"Appendices: {data.get('appendices', 'PPT, Task Sheets, assessment')}"
    
    # Row 22: Reflection
    merge_cells_in_row(table, 22, 0, 5).text = f"Reflection :\n{data.get('reflection', '')}"
    
    # Format all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    
    return doc

def generate_rtb_scheme_of_work(data):
    """Generate RTB-compliant scheme of work matching official template"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.page_width = Cm(29.7)  # A4 landscape
        section.page_height = Cm(21.0)
    
    # Header information
    doc.add_heading('SCHEME OF WORK', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # School info table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('School:', data.get('school', 'N/A')),
        ('Module Code & Title:', data.get('module_code_title', 'N/A')),
        ('Trainer Name:', data.get('trainer_name', 'N/A')),
        ('School Year:', data.get('school_year', 'N/A')),
        ('Term:', data.get('term', 'I'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Main scheme table: 9 columns
    scheme_table = doc.add_table(rows=2, cols=9)
    scheme_table.style = 'Table Grid'
    
    # Header row 1
    headers1 = ['Weeks', 'Competence code and name', '', '', 'Learning Activities', 
                'Resources (Equipment, tools, and materials)', 'Evidences of formative assessment', 
                'Learning Place', 'Observation']
    for i, header in enumerate(headers1):
        scheme_table.rows[0].cells[i].text = header
    
    # Header row 2
    headers2 = ['Weeks', 'Learning outcome (LO)', 'Duration', 'Indicative content (IC)', 
                'Learning Activities', 'Resources (Equipment, tools, and materials)', 
                'Evidences of formative assessment', 'Learning Place', 'Observation']
    for i, header in enumerate(headers2):
        scheme_table.rows[1].cells[i].text = header
        scheme_table.rows[1].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add weeks data
    weeks = data.get('weeks', [])
    for week in weeks:
        row = scheme_table.add_row()
        row.cells[0].text = week.get('week_range', '')
        row.cells[1].text = week.get('learning_outcome', '')
        row.cells[2].text = week.get('duration', '')
        row.cells[3].text = week.get('indicative_content', '')
        row.cells[4].text = week.get('learning_activities', '')
        row.cells[5].text = week.get('resources', '')
        row.cells[6].text = week.get('assessment_evidence', '')
        row.cells[7].text = week.get('learning_place', '')
        row.cells[8].text = week.get('observation', '')
    
    # Format all cells
    for row in scheme_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
    
    return doc
