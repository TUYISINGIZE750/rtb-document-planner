from docx import Document
import json

doc = Document(r"RTB Templates\TUYISINGIZE LEONARD_SESSION PLAN_S4F3.docx")

print("=== DOCUMENT STRUCTURE ANALYSIS ===\n")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total tables: {len(doc.tables)}\n")

# Analyze paragraphs
print("=== FIRST 20 PARAGRAPHS ===")
for i, para in enumerate(doc.paragraphs[:20]):
    text = para.text.strip()
    if text:
        print(f"{i}: {text[:80]}")

# Analyze tables
print("\n=== TABLES ANALYSIS ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}: {len(table.rows)} rows × {len(table.columns)} cols")
    
    # Show all rows
    for r_idx, row in enumerate(table.rows):
        cells_text = []
        for col_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()[:30]
            cells_text.append(f"C{col_idx}:{cell_text if cell_text else '(empty)'}")
        print(f"  Row {r_idx}: {' | '.join(cells_text)}")
