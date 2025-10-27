from docx import Document
import os

template_dir = r"c:\Users\PC\Music\Scheme of work and session plan planner\PRODUCTION_READY\RTB Templates"

session_template = os.path.join(template_dir, "RTB Session plan template.docx")
scheme_template = os.path.join(template_dir, "Scheme of work.docx")

print("=" * 80)
print("SESSION PLAN TEMPLATE ANALYSIS")
print("=" * 80)

try:
    doc = Document(session_template)
    print(f"Number of tables: {len(doc.tables)}")
    
    if doc.tables:
        table = doc.tables[0]
        print(f"Main table: {len(table.rows)} rows x {len(table.columns)} columns")
        print("\nFirst 10 rows structure:")
        for i, row in enumerate(table.rows[:10]):
            cells = [cell.text[:40] for cell in row.cells]
            print(f"Row {i}: {cells}")
except Exception as e:
    print(f"Error analyzing session template: {e}")

print("\n" + "=" * 80)
print("SCHEME OF WORK TEMPLATE ANALYSIS")
print("=" * 80)

try:
    doc = Document(scheme_template)
    print(f"Number of tables: {len(doc.tables)}")
    
    for t_idx, table in enumerate(doc.tables[:3]):
        print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} columns")
        if len(table.rows) >= 2:
            print(f"  Header row (Row 0): {[cell.text[:30] for cell in table.rows[0].cells[:5]]}")
            if len(table.rows) > 1:
                print(f"  Data row (Row 1): {[cell.text[:30] for cell in table.rows[1].cells[:5]]}")
except Exception as e:
    print(f"Error analyzing scheme template: {e}")
