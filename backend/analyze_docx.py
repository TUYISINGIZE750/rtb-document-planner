from docx import Document
import sys

def analyze_docx(file_path):
    try:
        doc = Document(file_path)
        
        print("=== DOCUMENT STRUCTURE ANALYSIS ===")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")
        
        print("\n=== PARAGRAPHS ===")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"Para {i}: {para.text[:100]}...")
        
        print("\n=== TABLES ===")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} columns")
            for j, row in enumerate(table.rows):
                print(f"  Row {j}:")
                for k, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"    Cell {k}: {cell.text.strip()}")
        
        print("\n=== STYLES ===")
        for style in doc.styles:
            print(f"Style: {style.name}")
            
    except Exception as e:
        print(f"Error analyzing document: {e}")

if __name__ == "__main__":
    analyze_docx("../GENCP 401 Session plan.docx")