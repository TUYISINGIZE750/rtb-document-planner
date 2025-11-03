"""Analyze official RTB Scheme of Work template structure"""
from docx import Document
import os

template_path = os.path.join(os.path.dirname(__file__), 'RTB Templates', 'Scheme of work.docx')

if not os.path.exists(template_path):
    print(f"Template not found: {template_path}")
    exit(1)

doc = Document(template_path)

print("=" * 80)
print("OFFICIAL RTB SCHEME OF WORK TEMPLATE ANALYSIS")
print("=" * 80)

print(f"\nTotal tables: {len(doc.tables)}")

for idx, table in enumerate(doc.tables):
    print(f"\n{'=' * 80}")
    print(f"TABLE {idx + 1}: {len(table.rows)} rows x {len(table.columns)} columns")
    print(f"{'=' * 80}")
    
    for row_idx, row in enumerate(table.rows):
        print(f"\nRow {row_idx}:")
        for cell_idx, cell in enumerate(row.cells):
            # Check if this cell is merged
            tc = cell._element
            tcPr = tc.tcPr
            
            gridSpan = None
            vMerge = None
            
            if tcPr is not None:
                gridSpan_elem = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                if gridSpan_elem is not None:
                    gridSpan = gridSpan_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                
                vMerge_elem = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                if vMerge_elem is not None:
                    vMerge_val = vMerge_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    vMerge = vMerge_val if vMerge_val else 'continue'
            
            text = cell.text.strip()[:50]
            merge_info = []
            if gridSpan:
                merge_info.append(f"colspan={gridSpan}")
            if vMerge:
                merge_info.append(f"vmerge={vMerge}")
            
            merge_str = f" [{', '.join(merge_info)}]" if merge_info else ""
            print(f"  Cell {cell_idx}: '{text}'{merge_str}")

print("\n" + "=" * 80)
