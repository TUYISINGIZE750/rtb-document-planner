"""Analyze RTB template structure to verify colspan/rowspan"""
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def analyze_table_structure(docx_path, table_name):
    print(f"\n{'='*70}")
    print(f"ANALYZING: {table_name}")
    print(f"{'='*70}")
    
    doc = Document(docx_path)
    
    if not doc.tables:
        print("ERROR: No tables found!")
        return
    
    table = doc.tables[0]
    print(f"\nTotal Rows: {len(table.rows)}")
    print(f"Total Columns (first row): {len(table.rows[0].cells)}")
    
    print(f"\n{'Row':<5} {'Cells':<7} {'Cell Details'}")
    print("-" * 70)
    
    for row_idx, row in enumerate(table.rows):
        cells_info = []
        for cell_idx, cell in enumerate(row.cells):
            # Get grid span (colspan)
            tc = cell._element
            tcPr = tc.tcPr
            
            gridSpan = 1
            vMerge = None
            
            if tcPr is not None:
                gridSpan_elem = tcPr.find(qn('w:gridSpan'))
                if gridSpan_elem is not None:
                    gridSpan = int(gridSpan_elem.get(qn('w:val')))
                
                vMerge_elem = tcPr.find(qn('w:vMerge'))
                if vMerge_elem is not None:
                    vMerge_val = vMerge_elem.get(qn('w:val'))
                    vMerge = 'restart' if vMerge_val == 'restart' else 'continue'
            
            text_preview = cell.text[:20].replace('\n', ' ') if cell.text else ''
            
            if gridSpan > 1 or vMerge:
                cells_info.append(f"[{cell_idx}:span={gridSpan},vmerge={vMerge}]")
            else:
                cells_info.append(f"[{cell_idx}]")
        
        print(f"{row_idx:<5} {len(row.cells):<7} {' '.join(cells_info)}")
    
    # Detailed analysis of merged cells
    print(f"\n{'='*70}")
    print("MERGED CELLS DETAILS")
    print(f"{'='*70}")
    
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            tc = cell._element
            tcPr = tc.tcPr
            
            if tcPr is not None:
                gridSpan_elem = tcPr.find(qn('w:gridSpan'))
                vMerge_elem = tcPr.find(qn('w:vMerge'))
                
                if gridSpan_elem is not None or vMerge_elem is not None:
                    gridSpan = int(gridSpan_elem.get(qn('w:val'))) if gridSpan_elem is not None else 1
                    vMerge = vMerge_elem.get(qn('w:val')) if vMerge_elem is not None else None
                    
                    text = cell.text[:30].replace('\n', ' ')
                    print(f"Row {row_idx}, Cell {cell_idx}: colspan={gridSpan}, vmerge={vMerge}")
                    print(f"  Text: '{text}'")

if __name__ == "__main__":
    print("\n" + "="*70)
    print("RTB TEMPLATE STRUCTURE ANALYSIS")
    print("="*70)
    
    # Analyze Session Plan
    analyze_table_structure(
        'RTB Templates/RTB Session plan template.docx',
        'SESSION PLAN TEMPLATE'
    )
    
    # Analyze Scheme of Work
    print("\n\n")
    analyze_table_structure(
        'RTB Templates/Scheme of work.docx',
        'SCHEME OF WORK TEMPLATE'
    )
    
    print("\n" + "="*70)
    print("ANALYSIS COMPLETE")
    print("="*70)
