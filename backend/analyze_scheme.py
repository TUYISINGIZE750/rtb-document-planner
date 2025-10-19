from docx import Document
import sys

def analyze_docx(file_path):
    try:
        doc = Document(file_path)
        
        print("=== SCHEME OF WORK STRUCTURE ANALYSIS ===")
        print(f"Total paragraphs: {len(doc.paragraphs)}")
        print(f"Total tables: {len(doc.tables)}")
        
        print("\n=== TABLES ===")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i}: {len(table.rows)} rows x {len(table.columns)} columns")
            for j, row in enumerate(table.rows[:10]):  # Show first 10 rows
                print(f"  Row {j}:")
                for k, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"    Cell {k}: {cell.text.strip()[:100]}...")
            
    except Exception as e:
        print(f"Error analyzing document: {e}")

if __name__ == "__main__":
    analyze_docx("../SCHEME_OF_WORK_OF_C_PROGRAMMING_L4CSA.docx")