"""Analyze template structure including colspan, rowspan, fonts, and formatting"""
from docx import Document
from docx.shared import Pt
import os

def analyze_cell_properties(cell):
    """Extract detailed cell properties"""
    props = {
        'text': cell.text[:30] if cell.text else '',
        'font_name': None,
        'font_size': None,
        'bold': False,
        'bg_color': None
    }
    
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            run = para.runs[0]
            if run.font.name:
                props['font_name'] = run.font.name
            if run.font.size:
                props['font_size'] = run.font.size.pt
            props['bold'] = run.font.bold
    
    # Check background color
    try:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shd = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        if shd is not None:
            props['bg_color'] = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
    except:
        pass
    
    return props

def analyze_merged_cells(table):
    """Detect merged cells (colspan/rowspan)"""
    merged = []
    processed = set()
    
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell_id = id(cell._element)
            if cell_id in processed:
                continue
            processed.add(cell_id)
            
            # Check if cell spans multiple columns or rows
            grid_span = 1
            v_merge = False
            
            try:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                gridSpan = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}gridSpan')
                if gridSpan is not None:
                    grid_span = int(gridSpan.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
                
                vMerge = tcPr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vMerge')
                if vMerge is not None:
                    v_merge = True
            except:
                pass
            
            if grid_span > 1 or v_merge:
                merged.append({
                    'row': row_idx,
                    'col': col_idx,
                    'colspan': grid_span,
                    'vmerge': v_merge,
                    'text': cell.text[:20]
                })
    
    return merged

def analyze_template(template_path, name):
    """Complete template analysis"""
    print(f"\n{'='*70}")
    print(f"ANALYZING: {name}")
    print(f"{'='*70}\n")
    
    if not os.path.exists(template_path):
        print(f"[ERROR] Not found: {template_path}")
        return None
    
    doc = Document(template_path)
    
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTABLE {table_idx + 1}:")
        print(f"  Rows: {len(table.rows)}, Columns: {len(table.columns)}")
        
        # Check for merged cells
        merged = analyze_merged_cells(table)
        if merged:
            print(f"\n  MERGED CELLS:")
            for m in merged:
                print(f"    Row {m['row']}, Col {m['col']}: colspan={m['colspan']}, vmerge={m['vmerge']} - '{m['text']}'")
        
        # Analyze first few rows for formatting
        print(f"\n  CELL FORMATTING (first 5 rows):")
        for row_idx in range(min(5, len(table.rows))):
            row = table.rows[row_idx]
            print(f"\n    Row {row_idx}:")
            for col_idx in range(min(6, len(row.cells))):
                cell = row.cells[col_idx]
                props = analyze_cell_properties(cell)
                if props['text'] or props['bg_color']:
                    print(f"      [{row_idx},{col_idx}]: {props['text'][:20]}")
                    if props['font_name']:
                        print(f"        Font: {props['font_name']}, Size: {props['font_size']}, Bold: {props['bold']}")
                    if props['bg_color']:
                        print(f"        Background: {props['bg_color']}")

def main():
    base_dir = os.path.dirname(__file__)
    
    # Analyze Session Plan
    session_template = os.path.join(base_dir, 'rtb_session_plan_template.docx')
    analyze_template(session_template, "SESSION PLAN TEMPLATE")
    
    # Analyze Scheme
    scheme_template = os.path.join(base_dir, 'rtb_scheme_template.docx')
    analyze_template(scheme_template, "SCHEME OF WORK TEMPLATE")
    
    print(f"\n{'='*70}")
    print("ANALYSIS COMPLETE")
    print(f"{'='*70}\n")

if __name__ == "__main__":
    main()
