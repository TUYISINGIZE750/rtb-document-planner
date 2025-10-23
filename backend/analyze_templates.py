from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

def analyze_template(file_path, doc_type):
    doc = Document(file_path)
    
    print(f"\n{'='*60}")
    print(f"{doc_type.upper()} TEMPLATE ANALYSIS")
    print(f"{'='*60}")
    
    print(f"\nTotal Tables: {len(doc.tables)}")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    
    for i, table in enumerate(doc.tables):
        print(f"\n--- Table {i} ---")
        print(f"Dimensions: {len(table.rows)} rows x {len(table.columns)} columns")
        
        for row_idx, row in enumerate(table.rows):
            print(f"\nRow {row_idx}:")
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    print(f"  [{row_idx},{col_idx}]: {text[:80]}")

if __name__ == "__main__":
    analyze_template("RTB Templates/TUYISINGIZE LEONARD_SESSION PLAN_S4F3.docx", "Session Plan")
    analyze_template("RTB Templates/CSAPA 301 Scheme of work.docx", "Scheme of Work")
