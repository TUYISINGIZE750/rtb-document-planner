from docx import Document

doc = Document('RTB Templates/Scheme of work.docx')
print(f'Total tables: {len(doc.tables)}')

print('\n=== TABLE 0 (Info Table) ===')
t0 = doc.tables[0]
print(f'Rows: {len(t0.rows)}, Cols: {len(t0.rows[0].cells)}')
for i in range(len(t0.rows)):
    print(f'\nRow {i}:')
    for j, cell in enumerate(t0.rows[i].cells[:5]):
        print(f'  Cell {j}: {repr(cell.text[:50])}')

print('\n=== TABLE 1 (Term 1) ===')
t1 = doc.tables[1]
print(f'Rows: {len(t1.rows)}, Cols: {len(t1.rows[0].cells)}')
print('Row 0:', [c.text[:20] for c in t1.rows[0].cells[:5]])
print('Row 1:', [c.text[:20] for c in t1.rows[1].cells[:5]])

print('\n=== TABLE 2 (Terms 2 & 3 or Signatures) ===')
t2 = doc.tables[2]
print(f'Rows: {len(t2.rows)}, Cols: {len(t2.rows[0].cells)}')
if len(t2.rows) > 0:
    print('Row 0:', [c.text[:30] for c in t2.rows[0].cells[:5]])
