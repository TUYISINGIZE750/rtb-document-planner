from docx import Document

doc = Document('RTB Templates/Scheme of work.docx')
print(f'Total tables: {len(doc.tables)}')

print('\n=== HEADER TABLE (Table 1) ===')
h = doc.tables[1]
print(f'Rows: {len(h.rows)}')
for i in range(min(8, len(h.rows))):
    cells_text = [c.text.strip()[:40] for c in h.rows[i].cells[:5]]
    print(f'Row {i}: {cells_text}')

print('\n=== TERM 1 TABLE (Table 2) ===')
t1 = doc.tables[2]
print(f'Rows: {len(t1.rows)}, Cols: {len(t1.rows[0].cells)}')
print('\nRow 0 (Header 1):')
for i, cell in enumerate(t1.rows[0].cells):
    print(f'  Cell {i}: "{cell.text.strip()}"')
print('\nRow 1 (Header 2):')
for i, cell in enumerate(t1.rows[1].cells):
    print(f'  Cell {i}: "{cell.text.strip()}"')
