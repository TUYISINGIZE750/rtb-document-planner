from docx import Document

doc = Document('RTB Templates/Scheme of work.docx')
print(f'Total tables: {len(doc.tables)}')

for table_idx in range(len(doc.tables)):
    print(f'\n=== TABLE {table_idx} ===')
    t = doc.tables[table_idx]
    print(f'Rows: {len(t.rows)}, Cols: {len(t.rows[0].cells) if t.rows else 0}')
    if table_idx == 0:
        print('This is the HEADER INFO table')
    elif table_idx == 1:
        print('This is the TERM 1 table')
        print('\nRow 0 (Header 1):')
        for i, cell in enumerate(t.rows[0].cells[:5]):
            print(f'  Cell {i}: {repr(cell.text[:30])}')
        print('\nRow 1 (Header 2):')
        for i, cell in enumerate(t.rows[1].cells[:5]):
            print(f'  Cell {i}: {repr(cell.text[:30])}')
    elif table_idx == 2:
        print('This is the TERM 2 & 3 combined table')
