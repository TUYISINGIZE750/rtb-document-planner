"""Check RTB template structure"""
from docx import Document
import os

try:
    template_path = 'rtb_session_plan_template.docx'
    if os.path.exists(template_path):
        doc = Document(template_path)
        print(f"Total tables in document: {len(doc.tables)}")
        
        if doc.tables:
            table = doc.tables[0]
            print(f"\nMain table structure:")
            print(f"  Rows: {len(table.rows)}")
            print(f"  Columns: {len(table.columns)}")
            
            print("\nFirst 25 rows content:")
            for i in range(min(25, len(table.rows))):
                row = table.rows[i]
                cells_content = [cell.text[:40].replace('\n', ' ') for cell in row.cells[:3]]
                print(f"  Row {i:2d}: {cells_content}")
    else:
        print("Template file not found")
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
