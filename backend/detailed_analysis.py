from docx import Document
from docx.shared import Inches
import sys

def analyze_rtb_design(file_path):
    try:
        doc = Document(file_path)
        
        print("=== DETAILED RTB DESIGN ANALYSIS ===")
        
        # Check headers and footers
        print(f"\nHeaders: {len(doc.sections[0].header.paragraphs)} paragraphs")
        for i, para in enumerate(doc.sections[0].header.paragraphs):
            if para.text.strip():
                print(f"Header {i}: {para.text}")
        
        # Check for images in header
        print(f"\nHeader images/shapes:")
        for rel in doc.sections[0].header.part.rels.values():
            if "image" in rel.target_ref:
                print(f"Image found: {rel.target_ref}")
        
        # Analyze table structure in detail
        if doc.tables:
            table = doc.tables[0]
            print(f"\nTable details:")
            print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
            
            # Check cell borders and styling
            print(f"\nTable style: {table.style.name if table.style else 'None'}")
            
            # Check for merged cells
            print(f"\nCell merge analysis:")
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    if hasattr(cell, '_tc') and cell._tc.grid_span > 1:
                        print(f"Row {i}, Cell {j}: spans {cell._tc.grid_span} columns")
        
        # Check document margins and page setup
        section = doc.sections[0]
        print(f"\nPage setup:")
        print(f"Page width: {section.page_width}")
        print(f"Page height: {section.page_height}")
        print(f"Left margin: {section.left_margin}")
        print(f"Right margin: {section.right_margin}")
        print(f"Top margin: {section.top_margin}")
        print(f"Bottom margin: {section.bottom_margin}")
        
        # Check for any embedded objects
        print(f"\nDocument relationships:")
        for rel in doc.part.rels.values():
            print(f"Relationship: {rel.reltype} -> {rel.target_ref}")
            
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    analyze_rtb_design("../GENCP 401 Session plan.docx")