from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import logging

logger = logging.getLogger(__name__)

try:
    from rtb_template_filler import fill_session_plan_template, fill_scheme_template
    HAS_BETTER_FILLER = True
except ImportError:
    HAS_BETTER_FILLER = False
    logger.warning("RTB template filler not available, using fallback")

def generate_session_plan_docx(data):
    """Generate RTB-compliant session plan DOCX"""
    if HAS_BETTER_FILLER:
        try:
            return fill_session_plan_template(data)
        except Exception as e:
            logger.error(f"Template filler failed: {e}, using fallback")
    
    return _generate_session_plan_fallback(data)

def generate_scheme_of_work_docx(data):
    """Generate RTB-compliant scheme of work DOCX"""
    if HAS_BETTER_FILLER:
        try:
            return fill_scheme_template(data)
        except Exception as e:
            logger.error(f"Scheme template filler failed: {e}, using fallback")
    
    return _generate_scheme_of_work_fallback(data)

def _generate_session_plan_fallback(data):
    """Fallback session plan generation with proper RTB format"""
    try:
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        title = doc.add_paragraph()
        title_run = title.add_run('RWANDA TECHNICAL BOARD (RTB)\nSESSION PLAN')
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Book Antiqua'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing = 1.5
        
        doc.add_paragraph()
        
        header_table = doc.add_table(rows=3, cols=3)
        header_table.style = 'Table Grid'
        
        for row_idx, (label1, val1, label2, val2) in enumerate([
            ("Code", data.get('module_code_title', ''), "Sector", data.get('sector', '')),
            ("Level", data.get('rqf_level', ''), "Term", data.get('term', '')),
            ("Date", data.get('date', ''), "Trainer", data.get('trainer_name', ''))
        ]):
            cells = header_table.rows[row_idx].cells
            _set_cell_text(cells[0], f"{label1}: {val1}")
            _set_cell_text(cells[1], f"{label2}: {val2}")
        
        doc.add_paragraph()
        
        content_table = doc.add_table(rows=11, cols=2)
        content_table.style = 'Table Grid'
        
        content_data = [
            ("Sector", data.get('sector', '')),
            ("Trade", data.get('trade', '')),
            ("RQF Level", data.get('rqf_level', '')),
            ("Module", data.get('module_code_title', '')),
            ("Class", data.get('class_name', '')),
            ("Trainees", str(data.get('number_of_trainees', ''))),
            ("Topic", data.get('topic_of_session', '')),
            ("Duration", f"{data.get('duration', '')} min"),
            ("Learning Outcomes", _clean(data.get('learning_outcomes', ''))),
            ("Contents", _clean(data.get('indicative_contents', ''))),
            ("Facilitation", _clean(data.get('facilitation_techniques', ''))),
        ]
        
        for i, (label, value) in enumerate(content_data):
            _set_cell_text(content_table.rows[i].cells[0], label, bold=True)
            _set_cell_text(content_table.rows[i].cells[1], value)
        
        doc.add_paragraph()
        
        for section_name, section_key in [
            ("Resources", "resources"),
            ("Learning Activities", "learning_activities"),
            ("Assessment", "assessment_details"),
            ("References", "references")
        ]:
            heading = doc.add_paragraph()
            heading_run = heading.add_run(section_name)
            heading_run.bold = True
            heading_run.font.size = Pt(12)
            heading_run.font.name = 'Book Antiqua'
            heading.paragraph_format.line_spacing = 1.5
            
            content = _clean(data.get(section_key, ''))
            if content:
                para = doc.add_paragraph(content)
                para.paragraph_format.line_spacing = 1.5
                for run in para.runs:
                    run.font.name = 'Book Antiqua'
                    run.font.size = Pt(12)
            
            doc.add_paragraph()
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        logger.info(f"Fallback session plan generated: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Fallback generation error: {e}")
        return _error_document()

def _generate_scheme_of_work_fallback(data):
    """Fallback scheme of work generation"""
    try:
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        title = doc.add_paragraph()
        title_run = title.add_run('RWANDA TECHNICAL BOARD (RTB)\nSCHEME OF WORK')
        title_run.bold = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Book Antiqua'
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing = 1.5
        
        doc.add_paragraph()
        
        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        _set_cell_text(info_table.rows[0].cells[0], "Module", bold=True)
        _set_cell_text(info_table.rows[0].cells[1], data.get('module_code_title', ''))
        _set_cell_text(info_table.rows[1].cells[0], "Trainer", bold=True)
        _set_cell_text(info_table.rows[1].cells[1], data.get('trainer_name', ''))
        _set_cell_text(info_table.rows[2].cells[0], "School Year", bold=True)
        _set_cell_text(info_table.rows[2].cells[1], data.get('school_year', ''))
        
        doc.add_paragraph()
        
        for term_num in [1, 2, 3]:
            if data.get(f'term{term_num}_weeks'):
                doc.add_heading(f'Term {term_num}', level=2)
                
                term_table = doc.add_table(rows=2, cols=4)
                term_table.style = 'Table Grid'
                
                headers = ['Weeks', 'Learning Outcomes', 'Contents', 'Duration']
                for i, header in enumerate(headers):
                    _set_cell_text(term_table.rows[0].cells[i], header, bold=True)
                
                _set_cell_text(term_table.rows[1].cells[0], data.get(f'term{term_num}_weeks', ''))
                _set_cell_text(term_table.rows[1].cells[1], _clean(data.get(f'term{term_num}_learning_outcomes', '')))
                _set_cell_text(term_table.rows[1].cells[2], _clean(data.get(f'term{term_num}_indicative_contents', '')))
                _set_cell_text(term_table.rows[1].cells[3], data.get(f'term{term_num}_duration', ''))
                
                doc.add_paragraph()
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        logger.info(f"Fallback scheme generated: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Scheme fallback error: {e}")
        return _error_document()

def _set_cell_text(cell, text, bold=False):
    """Set cell text with formatting"""
    cell.text = ''
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    para.text = ''
    if text:
        run = para.add_run(str(text))
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
        if bold:
            run.bold = True
    para.paragraph_format.line_spacing = 1.5

def _clean(text):
    """Clean and validate text"""
    if not text:
        return ""
    return str(text).strip()

def _error_document():
    """Generate error document"""
    try:
        doc = Document()
        doc.add_heading('Error Generating Document', 0)
        doc.add_paragraph('An error occurred while generating your document.')
        doc.add_paragraph('Please contact support if this issue persists.')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        return temp_file.name
    except:
        return None

def generate_session_plan_pdf(data):
    """PDF generation not implemented"""
    raise NotImplementedError("PDF generation coming soon")

def generate_scheme_of_work_pdf(data):
    """PDF generation not implemented"""
    raise NotImplementedError("PDF generation coming soon")
