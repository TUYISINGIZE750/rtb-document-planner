"""
RTB Document Generator - CLEAN VERSION
Uses RTB template ONLY - no fallbacks to prevent unstructured documents
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import tempfile
import os
import logging

logger = logging.getLogger(__name__)

def set_cell_background(cell, fill):
    """Set cell background color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def preserve_cell_format(cell, new_text, font_name='Book Antiqua', font_size=12, spacing=1.5):
    """Update cell text while preserving RTB formatting"""
    if not cell.paragraphs:
        cell.text = new_text
        return
    
    cell.paragraphs[0].clear()
    text_lines = str(new_text).split('\n') if new_text else ['']
    
    for idx, line in enumerate(text_lines):
        if idx == 0:
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        else:
            para = cell.add_paragraph()
        
        para.paragraph_format.line_spacing = spacing
        
        if line.strip():
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size)

def generate_session_plan_docx(data):
    """
    Generate RTB-compliant session plan - ALWAYS uses RTB structure
    Never falls back to plain text
    """
    
    logger.info("Generating RTB Session Plan - Using RTB Template Structure")
    
    # First, try to use actual RTB template file
    template_path = os.path.join(os.path.dirname(__file__), 'RTB Templates', 'RTB Session plan template.docx')
    
    if os.path.exists(template_path):
        try:
            logger.info(f"Loading RTB template from: {template_path}")
            doc = Document(template_path)
            
            # Fill template cells with data
            if len(doc.tables) > 0:
                table = doc.tables[0]
                
                # Fill header information (exact cell positions from RTB template)
                if len(table.rows) >= 6:
                    preserve_cell_format(table.rows[1].cells[1], data.get('code', '') or data.get('module_code_title', ''))
                    preserve_cell_format(table.rows[1].cells[3], data.get('sector', ''))
                    preserve_cell_format(table.rows[2].cells[1], data.get('trade', ''))
                    preserve_cell_format(table.rows[2].cells[3], data.get('rqf_level', ''))
                    preserve_cell_format(table.rows[3].cells[1], data.get('trainer_name', ''))
                    preserve_cell_format(table.rows[3].cells[3], data.get('module_code_title', ''))
                    preserve_cell_format(table.rows[4].cells[1], data.get('class_name', ''))
                    preserve_cell_format(table.rows[4].cells[3], str(data.get('number_of_trainees', '')))
                    preserve_cell_format(table.rows[5].cells[1], data.get('topic_of_session', ''))
                    preserve_cell_format(table.rows[5].cells[3], str(data.get('duration', '')))
            
            # Save and return
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            logger.info(f"✅ RTB Session Plan generated successfully: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            logger.error(f"Failed to use RTB template: {e}")
    
    # If template not found, create RTB-structured document from scratch
    logger.warning(f"RTB template not found at {template_path}, creating from scratch with RTB structure")
    
    doc = Document()
    
    # Set RTB margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('RWANDA TECHNICAL BOARD (RTB)\nSESSION PLAN')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Book Antiqua'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Header table
    header_table = doc.add_table(rows=6, cols=4)
    header_table.style = 'Table Grid'
    
    # Row 1
    preserve_cell_format(header_table.rows[0].cells[0], "Code:")
    preserve_cell_format(header_table.rows[0].cells[1], data.get('code', '') or data.get('module_code_title', ''))
    preserve_cell_format(header_table.rows[0].cells[2], "Sector:")
    preserve_cell_format(header_table.rows[0].cells[3], data.get('sector', ''))
    
    # Row 2
    preserve_cell_format(header_table.rows[1].cells[0], "Trade:")
    preserve_cell_format(header_table.rows[1].cells[1], data.get('trade', ''))
    preserve_cell_format(header_table.rows[1].cells[2], "RQF Level:")
    preserve_cell_format(header_table.rows[1].cells[3], data.get('rqf_level', ''))
    
    # Row 3
    preserve_cell_format(header_table.rows[2].cells[0], "Trainer:")
    preserve_cell_format(header_table.rows[2].cells[1], data.get('trainer_name', ''))
    preserve_cell_format(header_table.rows[2].cells[2], "Module:")
    preserve_cell_format(header_table.rows[2].cells[3], data.get('module_code_title', ''))
    
    # Row 4
    preserve_cell_format(header_table.rows[3].cells[0], "Class:")
    preserve_cell_format(header_table.rows[3].cells[1], data.get('class_name', ''))
    preserve_cell_format(header_table.rows[3].cells[2], "Trainees:")
    preserve_cell_format(header_table.rows[3].cells[3], str(data.get('number_of_trainees', '')))
    
    # Row 5
    preserve_cell_format(header_table.rows[4].cells[0], "Topic:")
    preserve_cell_format(header_table.rows[4].cells[1], data.get('topic_of_session', ''))
    preserve_cell_format(header_table.rows[4].cells[2], "Duration:")
    preserve_cell_format(header_table.rows[4].cells[3], f"{data.get('duration', '')} min")
    
    # Row 6
    preserve_cell_format(header_table.rows[5].cells[0], "Term:")
    preserve_cell_format(header_table.rows[5].cells[1], data.get('term', ''))
    preserve_cell_format(header_table.rows[5].cells[2], "Week:")
    preserve_cell_format(header_table.rows[5].cells[3], data.get('week', ''))
    
    doc.add_paragraph()
    
    # Content table - 11 rows for RTB fields
    content_table = doc.add_table(rows=11, cols=2)
    content_table.style = 'Table Grid'
    content_table.columns[0].width = Cm(4)
    content_table.columns[1].width = Cm(13)
    
    content_data = [
        ("Learning Outcomes", data.get('learning_outcomes', '')),
        ("Indicative Contents", data.get('indicative_contents', '')),
        ("Facilitation Technique", data.get('facilitation_techniques', '')),
        ("Learning Activities", data.get('learning_activities', '') or "See facilitation technique"),
        ("Resources", data.get('resources', '')),
        ("Assessment", data.get('assessment_details', '')),
        ("References", data.get('references', '')),
        ("Notes", ""),
        ("Reflection", ""),
        ("Trainer Signature", ""),
        ("Date", data.get('date', '')),
    ]
    
    for i, (label, value) in enumerate(content_data):
        preserve_cell_format(content_table.rows[i].cells[0], label)
        preserve_cell_format(content_table.rows[i].cells[1], value)
    
    # Save to temp file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    logger.info(f"✅ RTB Session Plan created from scratch: {temp_file.name}")
    return temp_file.name


def generate_scheme_of_work_docx(data):
    """Generate RTB-compliant scheme of work document"""
    
    logger.info("Generating RTB Scheme of Work")
    
    # Try to use template
    template_path = os.path.join(os.path.dirname(__file__), 'RTB Templates', 'Scheme of work.docx')
    
    if os.path.exists(template_path):
        try:
            logger.info(f"Loading Scheme template from: {template_path}")
            doc = Document(template_path)
            
            # Fill template with data
            if len(doc.tables) > 0:
                table = doc.tables[0]
                if len(table.rows) >= 4:
                    preserve_cell_format(table.rows[1].cells[1], data.get('sector', ''))
                    preserve_cell_format(table.rows[1].cells[3], data.get('district', ''))
                    preserve_cell_format(table.rows[2].cells[1], data.get('department_trade', ''))
                    preserve_cell_format(table.rows[3].cells[1], data.get('trainer_name', ''))
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            logger.info(f"✅ RTB Scheme of Work generated: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            logger.error(f"Failed to use Scheme template: {e}")
    
    # Create from scratch
    logger.warning(f"Scheme template not found, creating from scratch")
    
    doc = Document()
    
    # RTB margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('RWANDA TECHNICAL BOARD (RTB)\nSCHEME OF WORK')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Book Antiqua'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    # Header
    header_table = doc.add_table(rows=4, cols=4)
    header_table.style = 'Table Grid'
    
    preserve_cell_format(header_table.rows[0].cells[0], "Sector:")
    preserve_cell_format(header_table.rows[0].cells[1], data.get('sector', ''))
    preserve_cell_format(header_table.rows[0].cells[2], "District:")
    preserve_cell_format(header_table.rows[0].cells[3], data.get('district', ''))
    
    preserve_cell_format(header_table.rows[1].cells[0], "Department:")
    preserve_cell_format(header_table.rows[1].cells[1], data.get('department_trade', ''))
    preserve_cell_format(header_table.rows[1].cells[2], "Module:")
    preserve_cell_format(header_table.rows[1].cells[3], data.get('module_code_title', ''))
    
    preserve_cell_format(header_table.rows[2].cells[0], "Trainer:")
    preserve_cell_format(header_table.rows[2].cells[1], data.get('trainer_name', ''))
    preserve_cell_format(header_table.rows[2].cells[2], "Year:")
    preserve_cell_format(header_table.rows[2].cells[3], data.get('school_year', ''))
    
    preserve_cell_format(header_table.rows[3].cells[0], "Class:")
    preserve_cell_format(header_table.rows[3].cells[1], data.get('class_name', ''))
    preserve_cell_format(header_table.rows[3].cells[2], "Level:")
    preserve_cell_format(header_table.rows[3].cells[3], data.get('rqf_level', ''))
    
    doc.add_paragraph()
    
    # Content table for terms
    for term_num in range(1, 4):
        doc.add_paragraph(f"\nTERM {term_num}")
        
        term_table = doc.add_table(rows=5, cols=2)
        term_table.style = 'Table Grid'
        term_table.columns[0].width = Cm(4)
        term_table.columns[1].width = Cm(13)
        
        preserve_cell_format(term_table.rows[0].cells[0], "Weeks")
        preserve_cell_format(term_table.rows[0].cells[1], data.get(f'term{term_num}_weeks', ''))
        
        preserve_cell_format(term_table.rows[1].cells[0], "Learning Outcomes")
        preserve_cell_format(term_table.rows[1].cells[1], data.get(f'term{term_num}_learning_outcomes', ''))
        
        preserve_cell_format(term_table.rows[2].cells[0], "Indicative Contents")
        preserve_cell_format(term_table.rows[2].cells[1], data.get(f'term{term_num}_indicative_contents', ''))
        
        preserve_cell_format(term_table.rows[3].cells[0], "Duration")
        preserve_cell_format(term_table.rows[3].cells[1], data.get(f'term{term_num}_duration', ''))
        
        preserve_cell_format(term_table.rows[4].cells[0], "Learning Place")
        preserve_cell_format(term_table.rows[4].cells[1], data.get(f'term{term_num}_learning_place', ''))
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    logger.info(f"✅ RTB Scheme of Work created: {temp_file.name}")
    return temp_file.name


def generate_session_plan_pdf(data):
    """Generate PDF from session plan"""
    docx_path = generate_session_plan_docx(data)
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        from docx2pdf.convert import convert as docx2pdf_convert
        docx2pdf_convert(docx_path, pdf_path)
        return pdf_path
    except:
        logger.warning("PDF conversion not available, returning DOCX path")
        return docx_path


def generate_scheme_of_work_pdf(data):
    """Generate PDF from scheme of work"""
    docx_path = generate_scheme_of_work_docx(data)
    pdf_path = docx_path.replace('.docx', '.pdf')
    
    try:
        from docx2pdf.convert import convert as docx2pdf_convert
        docx2pdf_convert(docx_path, pdf_path)
        return pdf_path
    except:
        logger.warning("PDF conversion not available, returning DOCX path")
        return docx_path
