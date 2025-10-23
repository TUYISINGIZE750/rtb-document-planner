from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import tempfile
import os
import logging

logger = logging.getLogger(__name__)

def generate_session_plan_docx(data):
    """Generate RTB Session Plan DOCX document"""
    try:
        doc = Document()
        
        # Helper to get value from dict or object with validation
        def get_val(key, default='N/A'):
            try:
                if isinstance(data, dict):
                    value = data.get(key, default)
                else:
                    value = getattr(data, key, default)
                return str(value) if value is not None else default
            except Exception as e:
                logger.error(f"Error getting value for {key}: {e}")
                return default
        
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "REPUBLIC OF RWANDA\nMINISTRY OF EDUCATION\nRWANDA TECHNICAL BOARD (RTB)"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title = doc.add_heading('SESSION PLAN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Basic Information Table
        table = doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Fill basic info with validation
        cells = table.rows[0].cells
        cells[0].text = 'Sector:'
        cells[1].text = get_val('sector')
        cells[2].text = 'Sub-Sector:'
        cells[3].text = get_val('sub_sector')
        
        cells = table.rows[1].cells
        cells[0].text = 'Trade:'
        cells[1].text = get_val('trade')
        cells[2].text = 'Qualification Title:'
        cells[3].text = get_val('qualification_title')
        
        cells = table.rows[2].cells
        cells[0].text = 'RQF Level:'
        cells[1].text = get_val('rqf_level')
        cells[2].text = 'Module Code & Title:'
        cells[3].text = get_val('module_code_title')
        
        cells = table.rows[3].cells
        cells[0].text = 'Term:'
        cells[1].text = get_val('term')
        cells[2].text = 'Week:'
        cells[3].text = get_val('week')
        
        cells = table.rows[4].cells
        cells[0].text = 'Date:'
        cells[1].text = get_val('date')
        cells[2].text = 'Trainer Name:'
        cells[3].text = get_val('trainer_name')
        
        cells = table.rows[5].cells
        cells[0].text = 'Class:'
        cells[1].text = get_val('class_name')
        cells[2].text = 'Number of Trainees:'
        cells[3].text = get_val('number_of_trainees')
        
        cells = table.rows[6].cells
        cells[0].text = 'Learning Outcomes:'
        cells[1].text = get_val('learning_outcomes')
        cells[2].text = 'Indicative Contents:'
        cells[3].text = get_val('indicative_contents')
        
        cells = table.rows[7].cells
        cells[0].text = 'Topic of Session:'
        cells[1].text = get_val('topic_of_session')
        cells[2].text = 'Duration:'
        cells[3].text = get_val('duration') + ' minutes'
        
        doc.add_paragraph()
        
        # Session Details
        doc.add_heading('Session Details', level=1)
        
        doc.add_heading('Objectives:', level=2)
        doc.add_paragraph(get_val('objectives', 'To be defined based on learning outcomes'))
        
        doc.add_heading('Facilitation Techniques:', level=2)
        doc.add_paragraph(get_val('facilitation_techniques', 'Interactive learning and practical exercises'))
        
        doc.add_heading('Learning Activities:', level=2)
        doc.add_paragraph(get_val('learning_activities', 'Hands-on practice and group discussions'))
        
        doc.add_heading('Resources:', level=2)
        doc.add_paragraph(get_val('resources', 'Textbooks, computers, and practical materials'))
        
        doc.add_heading('Assessment Details:', level=2)
        doc.add_paragraph(get_val('assessment_details', 'Continuous assessment and practical evaluation'))
        
        doc.add_heading('References:', level=2)
        doc.add_paragraph(get_val('references', 'RTB curriculum guidelines and approved textbooks'))
        
        # Save to temporary file with proper error handling
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        logger.info(f"Session plan DOCX generated: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating session plan DOCX: {e}")
        # Create a simple fallback document
        doc = Document()
        doc.add_heading('RTB Session Plan', 0)
        doc.add_paragraph('Error generating document. Please contact support.')
        doc.add_paragraph(f'Error details: {str(e)}')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name

def generate_scheme_of_work_docx(data):
    """Generate RTB Scheme of Work DOCX document"""
    try:
        doc = Document()
        
        # Helper to get value from dict or object with validation
        def get_val(key, default='N/A'):
            try:
                if isinstance(data, dict):
                    value = data.get(key, default)
                else:
                    value = getattr(data, key, default)
                return str(value) if value is not None else default
            except Exception as e:
                logger.error(f"Error getting value for {key}: {e}")
                return default
        
        # Header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "REPUBLIC OF RWANDA\nMINISTRY OF EDUCATION\nRWANDA TECHNICAL BOARD (RTB)"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Title
        title = doc.add_heading('SCHEME OF WORK', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Basic Information
        doc.add_heading('Basic Information', level=1)
        
        info_table = doc.add_table(rows=6, cols=4)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cells = info_table.rows[0].cells
        cells[0].text = 'Province:'
        cells[1].text = get_val('province')
        cells[2].text = 'District:'
        cells[3].text = get_val('district')
        
        cells = info_table.rows[1].cells
        cells[0].text = 'Sector:'
        cells[1].text = get_val('sector')
        cells[2].text = 'School:'
        cells[3].text = get_val('school')
        
        cells = info_table.rows[2].cells
        cells[0].text = 'Department/Trade:'
        cells[1].text = get_val('department_trade')
        cells[2].text = 'Qualification Title:'
        cells[3].text = get_val('qualification_title')
        
        cells = info_table.rows[3].cells
        cells[0].text = 'RQF Level:'
        cells[1].text = get_val('rqf_level')
        cells[2].text = 'Module Code & Title:'
        cells[3].text = get_val('module_code_title')
        
        cells = info_table.rows[4].cells
        cells[0].text = 'School Year:'
        cells[1].text = get_val('school_year')
        cells[2].text = 'Terms:'
        cells[3].text = get_val('terms')
        
        cells = info_table.rows[5].cells
        cells[0].text = 'Module Hours:'
        cells[1].text = get_val('module_hours')
        cells[2].text = 'Number of Classes:'
        cells[3].text = get_val('number_of_classes')
        
        # Term Details
        for term_num in [1, 2, 3]:
            if get_val(f'term{term_num}_weeks'):
                doc.add_heading(f'Term {term_num}', level=1)
                
                term_table = doc.add_table(rows=1, cols=4)
                term_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header row
                hdr_cells = term_table.rows[0].cells
                hdr_cells[0].text = 'Weeks'
                hdr_cells[1].text = 'Learning Outcomes'
                hdr_cells[2].text = 'Indicative Contents'
                hdr_cells[3].text = 'Duration'
                
                # Data row
                row_cells = term_table.add_row().cells
                row_cells[0].text = get_val(f'term{term_num}_weeks')
                row_cells[1].text = get_val(f'term{term_num}_learning_outcomes')
                row_cells[2].text = get_val(f'term{term_num}_indicative_contents')
                row_cells[3].text = get_val(f'term{term_num}_duration')
        
        # Signatures
        doc.add_heading('Signatures', level=1)
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cells = sig_table.rows[0].cells
        cells[0].text = 'Trainer Name:'
        cells[1].text = get_val('trainer_name')
        
        cells = sig_table.rows[1].cells
        cells[0].text = 'DOS Name:'
        cells[1].text = get_val('dos_name')
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        logger.info(f"Scheme DOCX generated: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating scheme DOCX: {e}")
        # Create a simple fallback document
        doc = Document()
        doc.add_heading('RTB Scheme of Work', 0)
        doc.add_paragraph('Error generating document. Please contact support.')
        doc.add_paragraph(f'Error details: {str(e)}')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name