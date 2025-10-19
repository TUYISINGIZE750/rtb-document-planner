from docx import Document
import os

def extract_header_info(file_path):
    try:
        doc = Document(file_path)
        
        print("=== HEADER CONTENT ANALYSIS ===")
        
        # Get header content
        header = doc.sections[0].header
        
        print(f"Header paragraphs: {len(header.paragraphs)}")
        for i, para in enumerate(header.paragraphs):
            print(f"Header para {i}: '{para.text}'")
            
        # Check for tables in header
        print(f"Header tables: {len(header.tables)}")
        for i, table in enumerate(header.tables):
            print(f"Header table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
            for j, row in enumerate(table.rows):
                print(f"  Row {j}:")
                for k, cell in enumerate(row.cells):
                    if cell.text.strip():
                        print(f"    Cell {k}: {cell.text.strip()}")
        
        # Extract images from media folder
        print(f"\nExtracting header image...")
        
        # Copy the image files to our backend folder
        import shutil
        
        # The image should be in the document's media folder
        # Let's copy it to our backend for use in generation
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    extract_header_info("../GENCP 401 Session plan.docx")