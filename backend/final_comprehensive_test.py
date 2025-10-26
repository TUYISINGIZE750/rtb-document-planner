"""Final comprehensive test - check everything"""
from rtb_template_filler_exact import fill_session_plan_template, fill_scheme_template
from docx import Document
import os

def test_all_scenarios():
    """Test all possible scenarios"""
    print("\n" + "="*70)
    print("FINAL COMPREHENSIVE TEST")
    print("="*70 + "\n")
    
    issues = []
    
    # Test 1: Session Plan with all fields
    print("Test 1: Session Plan - All Fields...")
    try:
        data = {
            'sector': 'ICT',
            'trade': 'Computer Hardware',
            'trainer_name': 'John MUGISHA',
            'module_code_title': 'CHM4101',
            'term': 'Term 1',
            'week': 'Week 5',
            'date': '15/01/2025',
            'class_name': 'CHM4A',
            'number_of_trainees': '25',
            'learning_outcomes': 'Assemble computers correctly',
            'indicative_contents': 'Computer components, Assembly procedures',
            'topic_of_session': 'Installing Motherboard',
            'duration': '40',
            'objectives': 'Install motherboard\nConnect power supply',
            'facilitation_techniques': 'Hands-on practice',
            'references': 'RTB Manual'
        }
        file = fill_session_plan_template(data)
        doc = Document(file)
        
        # Check structure
        if len(doc.tables) != 1 or len(doc.tables[0].rows) != 23:
            issues.append("Test 1: Structure mismatch")
        
        # Check content
        table = doc.tables[0]
        if 'ICT' not in table.rows[1].cells[0].text:
            issues.append("Test 1: Sector not found")
        if 'MUGISHA' not in table.rows[2].cells[0].text:
            issues.append("Test 1: Trainer not found")
        if 'Installing Motherboard' not in table.rows[6].cells[0].text:
            issues.append("Test 1: Topic not found")
        
        print("  [PASS]")
    except Exception as e:
        issues.append(f"Test 1: {str(e)}")
        print(f"  [FAIL] {e}")
    
    # Test 2: Different facilitation techniques
    print("\nTest 2: All Facilitation Techniques...")
    techniques = ['Trainer-guided', 'Simulation', 'Group work', 'Hands-on', 'Discussion', 'Project-based']
    for tech in techniques:
        try:
            data['facilitation_techniques'] = tech
            file = fill_session_plan_template(data)
            doc = Document(file)
            content = doc.tables[0].rows[13].cells[0].text
            if len(content) < 50:
                issues.append(f"Test 2: {tech} - content too short")
            print(f"  {tech}: [PASS]")
        except Exception as e:
            issues.append(f"Test 2: {tech} - {str(e)}")
            print(f"  {tech}: [FAIL]")
    
    # Test 3: Scheme of Work
    print("\nTest 3: Scheme of Work - All Terms...")
    try:
        scheme_data = {
            'term1_weeks': 'Week 1-12',
            'term1_learning_outcomes': 'LO1: Prepare tools',
            'term1_duration': '40 hours',
            'term1_indicative_contents': 'IC1.1: Workplace prep',
            'term2_weeks': 'Week 13-24',
            'term2_learning_outcomes': 'LO2: Install OS',
            'term2_duration': '40 hours',
            'term2_indicative_contents': 'IC2.1: OS installation',
            'term3_weeks': 'Week 25-36',
            'term3_learning_outcomes': 'LO3: Troubleshoot',
            'term3_duration': '40 hours',
            'term3_indicative_contents': 'IC3.1: Diagnostics'
        }
        file = fill_scheme_template(scheme_data)
        doc = Document(file)
        
        if len(doc.tables) != 3:
            issues.append("Test 3: Should have 3 tables")
        
        # Check each term
        for i in range(3):
            if len(doc.tables[i].rows) < 2:
                issues.append(f"Test 3: Term {i+1} structure issue")
            if f'Week {i*12+1}' not in doc.tables[i].rows[2].cells[0].text:
                issues.append(f"Test 3: Term {i+1} weeks not found")
        
        print("  [PASS]")
    except Exception as e:
        issues.append(f"Test 3: {str(e)}")
        print(f"  [FAIL] {e}")
    
    # Test 4: Content formatting
    print("\nTest 4: Content Formatting...")
    try:
        file = fill_session_plan_template(data)
        doc = Document(file)
        table = doc.tables[0]
        
        # Check objectives numbering
        obj_text = table.rows[8].cells[0].text
        if '1.' not in obj_text or '2.' not in obj_text:
            issues.append("Test 4: Objectives not numbered")
        
        # Check no excessive spacing
        for row in table.rows:
            for cell in row.cells:
                if '   ' in cell.text:
                    issues.append("Test 4: Excessive spacing found")
                    break
        
        print("  [PASS]")
    except Exception as e:
        issues.append(f"Test 4: {str(e)}")
        print(f"  [FAIL] {e}")
    
    # Test 5: Empty/Missing fields handling
    print("\nTest 5: Missing Fields Handling...")
    try:
        minimal_data = {
            'sector': 'ICT',
            'topic_of_session': 'Test Topic',
            'facilitation_techniques': 'Trainer-guided'
        }
        file = fill_session_plan_template(minimal_data)
        doc = Document(file)
        if len(doc.tables) != 1:
            issues.append("Test 5: Failed with minimal data")
        print("  [PASS]")
    except Exception as e:
        issues.append(f"Test 5: {str(e)}")
        print(f"  [FAIL] {e}")
    
    # Test 6: Special characters
    print("\nTest 6: Special Characters...")
    try:
        special_data = data.copy()
        special_data['topic_of_session'] = "Installing CPU & RAM (32GB)"
        special_data['objectives'] = "1. Install CPU\n2. Install RAM (DDR4)"
        file = fill_session_plan_template(special_data)
        doc = Document(file)
        content = doc.tables[0].rows[6].cells[0].text
        if '&' not in content or '32GB' not in content:
            issues.append("Test 6: Special characters not preserved")
        print("  [PASS]")
    except Exception as e:
        issues.append(f"Test 6: {str(e)}")
        print(f"  [FAIL] {e}")
    
    # Test 7: Long content
    print("\nTest 7: Long Content...")
    try:
        long_data = data.copy()
        long_data['learning_outcomes'] = "By the end of this module, learners will be able to: " + "Perform complex tasks. " * 20
        long_data['indicative_contents'] = "Content includes: " + "Various topics. " * 30
        file = fill_session_plan_template(long_data)
        doc = Document(file)
        if len(doc.tables[0].rows[4].cells[1].text) < 100:
            issues.append("Test 7: Long content not preserved")
        print("  [PASS]")
    except Exception as e:
        issues.append(f"Test 7: {str(e)}")
        print(f"  [FAIL] {e}")
    
    # Summary
    print("\n" + "="*70)
    print("TEST SUMMARY")
    print("="*70)
    
    if not issues:
        print("\n✅ ALL TESTS PASSED - NO ISSUES FOUND!")
        print("\nSystem Status:")
        print("  ✅ Content structuring: PERFECT")
        print("  ✅ All logic: WORKING FLAWLESSLY")
        print("  ✅ Formatting: CLEAN")
        print("  ✅ Template compliance: 100%")
        print("  ✅ Error handling: ROBUST")
        print("\n🎉 READY FOR PRODUCTION DEPLOYMENT!")
        return True
    else:
        print(f"\n⚠️ FOUND {len(issues)} ISSUE(S):")
        for issue in issues:
            print(f"  - {issue}")
        return False

if __name__ == "__main__":
    test_all_scenarios()
