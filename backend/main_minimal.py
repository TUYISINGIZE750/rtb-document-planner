from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
import hashlib
import tempfile
import os
from docx import Document

app = Flask(__name__)
CORS(app, origins=["https://tuyisingize750.github.io"])

# Simple in-memory storage for demo
users = {
    "+250796014803": {
        "name": "Test Teacher",
        "password": hashlib.sha256("teacher123".encode()).hexdigest(),
        "phone": "+250796014803"
    },
    "+250789751597": {
        "name": "Administrator", 
        "password": hashlib.sha256("admin123".encode()).hexdigest(),
        "phone": "+250789751597"
    }
}

documents = {}
doc_counter = 1

@app.route('/', methods=['GET', 'OPTIONS'])
def home():
    if request.method == 'OPTIONS':
        return '', 204
    return jsonify({"message": "RTB API", "status": "online", "version": "minimal"})

@app.route('/users/login', methods=['POST', 'OPTIONS'])
def login():
    if request.method == 'OPTIONS':
        return '', 204
    
    data = request.get_json()
    phone = data.get('phone')
    password = data.get('password')
    
    if phone in users:
        user = users[phone]
        if hashlib.sha256(password.encode()).hexdigest() == user['password']:
            return jsonify({
                "name": user['name'],
                "phone": user['phone'],
                "role": "user"
            })
    
    return jsonify({"detail": "Invalid credentials"}), 401

@app.route('/session-plans', methods=['POST', 'OPTIONS'])
def create_session_plan():
    if request.method == 'OPTIONS':
        return '', 204
    
    global doc_counter
    doc_id = doc_counter
    doc_counter += 1
    
    data = request.get_json()
    documents[doc_id] = {
        'type': 'session_plan',
        'data': data
    }
    
    return jsonify({"id": doc_id, "message": "Session plan created"})

@app.route('/schemes', methods=['POST', 'OPTIONS'])
def create_scheme():
    if request.method == 'OPTIONS':
        return '', 204
    
    global doc_counter
    doc_id = doc_counter
    doc_counter += 1
    
    data = request.get_json()
    documents[doc_id] = {
        'type': 'scheme',
        'data': data
    }
    
    return jsonify({"id": doc_id, "message": "Scheme created"})

@app.route('/session-plans/<int:plan_id>/download', methods=['GET'])
def download_session_plan(plan_id):
    if plan_id not in documents:
        return jsonify({"detail": "Not found"}), 404
    
    data = documents[plan_id]['data']
    
    # Create simple DOCX
    doc = Document()
    doc.add_heading('RTB SESSION PLAN', 0)
    doc.add_paragraph(f"Sector: {data.get('sector', 'N/A')}")
    doc.add_paragraph(f"Trade: {data.get('trade', 'N/A')}")
    doc.add_paragraph(f"Topic: {data.get('topic_of_session', 'N/A')}")
    doc.add_paragraph(f"Teacher: {data.get('trainer_name', 'N/A')}")
    doc.add_paragraph(f"Duration: {data.get('duration', '40')} minutes")
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    
    return send_file(temp_file.name, as_attachment=True, download_name=f"session_plan_{plan_id}.docx")

@app.route('/schemes/<int:scheme_id>/download', methods=['GET'])
def download_scheme(scheme_id):
    if scheme_id not in documents:
        return jsonify({"detail": "Not found"}), 404
    
    data = documents[scheme_id]['data']
    
    # Create simple DOCX
    doc = Document()
    doc.add_heading('RTB SCHEME OF WORK', 0)
    doc.add_paragraph(f"School: {data.get('school', 'N/A')}")
    doc.add_paragraph(f"Module: {data.get('module_code_title', 'N/A')}")
    doc.add_paragraph(f"Teacher: {data.get('trainer_name', 'N/A')}")
    doc.add_paragraph(f"Year: {data.get('school_year', 'N/A')}")
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    
    return send_file(temp_file.name, as_attachment=True, download_name=f"scheme_{scheme_id}.docx")

if __name__ == '__main__':
    app.run(debug=False)