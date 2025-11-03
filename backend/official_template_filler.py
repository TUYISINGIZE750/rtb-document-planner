"""Official RTB Template Filler - Uses templates from DOCS TO REFER TO folder"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import qn
import os
import tempfile
from datetime import datetime
import base64

def set_cell_font(cell, font_name='Bookman Old Style', font_size=12, bold=False):
    """Set font for all paragraphs and runs in a cell with 1.5 line spacing"""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
        if not paragraph.runs and paragraph.text:
            run = paragraph.add_run(paragraph.text)
            paragraph.text = ''
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold

def set_cell_text_with_bold_label(cell, label, value, font_name='Bookman Old Style', font_size=12):
    """Set cell text with bold label and normal value - 1.5 line spacing"""
    cell.text = ''
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    label_run = p.add_run(label)
    label_run.font.name = font_name
    label_run.font.size = Pt(font_size)
    label_run.font.bold = True
    
    value_run = p.add_run(value)
    value_run.font.name = font_name
    value_run.font.size = Pt(font_size)
    value_run.font.bold = False

def fill_session_plan_official(data):
    """Fill RTB Session plan template.docx from RTB Templates folder"""
    import sys
    import logging
    logger = logging.getLogger(__name__)
    
    if data is None:
        logger.error("❌ Data is None!")
        raise ValueError("Data cannot be None")
    
    logger.info(f"🐍 Python version: {sys.version}")
    logger.info(f"📂 Current file: {__file__}")
    logger.info(f"📂 Current dir: {os.getcwd()}")
    logger.info(f"📊 Data type: {type(data)}")
    logger.info(f"📊 Data keys: {list(data.keys()) if isinstance(data, dict) else 'NOT A DICT'}")
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    logger.info(f"📂 Base dir: {base_dir}")
    
    # List files in base directory
    try:
        files = os.listdir(base_dir)
        logger.info(f"📂 Files in base dir: {files[:10]}")
    except Exception as e:
        logger.error(f"❌ Cannot list base dir: {e}")
    
    # Check if RTB Templates folder exists
    rtb_folder = os.path.join(base_dir, 'RTB Templates')
    logger.info(f"📂 RTB Templates path: {rtb_folder}")
    logger.info(f"📂 RTB Templates exists: {os.path.exists(rtb_folder)}")
    
    if os.path.exists(rtb_folder):
        try:
            template_files = os.listdir(rtb_folder)
            logger.info(f"📂 Files in RTB Templates: {template_files}")
        except Exception as e:
            logger.error(f"❌ Cannot list RTB Templates: {e}")
    
    # Use edited template from OFFICIALRTBTEMPLATES (no shapes)
    parent_dir = os.path.dirname(base_dir)
    template_path = os.path.join(parent_dir, 'OFFICIALRTBTEMPLATES', 'SESSION PLAN TEMPLATES.docx')
    
    # Fallback to old template if new one doesn't exist
    if not os.path.exists(template_path):
        template_path = os.path.join(base_dir, 'RTB Templates', 'RTB Session plan template.docx')
        logger.info(f"⚠️ Using fallback template: {template_path}")
    logger.info(f"📂 Full template path: {template_path}")
    logger.info(f"📂 Template exists: {os.path.exists(template_path)}")
    
    if not os.path.exists(template_path):
        logger.warning(f"⚠️ Template not found, creating simple document")
        # Fallback: create simple document
        doc = Document()
        doc.add_heading('RTB Session Plan', 0)
        doc.add_paragraph(f"Topic: {data.get('topic_of_session', '')}")
        doc.add_paragraph(f"Objectives: {data.get('objectives', '')}")
        doc.add_paragraph(f"Learning Activities: {data.get('learning_activities', '')}")
        doc.add_paragraph(f"Assessment: {data.get('assessment_details', '')}")
        doc.add_paragraph(f"References: {data.get('references', '')}")
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        logger.info(f"✅ Fallback document created: {temp_file.name}")
        return temp_file.name
    
    try:
        doc = Document(template_path)
        logger.info(f"✅ Template loaded, tables: {len(doc.tables)}")
        
        if not doc.tables:
            raise Exception("No table found in template")
    except Exception as e:
        logger.error(f"❌ Error loading template: {str(e)}")
        raise
    
    if not doc.tables:
        logger.error("❌ No tables in template!")
        raise Exception("Template has no tables")
    
    # Add header section at top of document
    school_name = data.get('school_name', '')
    province = data.get('province', '')
    district = data.get('district', '')
    sector_loc = data.get('sector_location', '')
    cell_loc = data.get('cell', '')
    village = data.get('village', '')
    
    # Insert header table at beginning
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Remove spacing from header table
    for row in header_table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcMar xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/></w:tcMar>'))
    
    # LEFT: RTB Logo
    left_cell = header_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)
    left_para.paragraph_format.line_spacing = 1.5
    left_run = left_para.add_run('RWANDA\nTVET BOARD')
    left_run.font.bold = True
    left_run.font.size = Pt(9)
    left_run.font.name = 'Bookman Old Style'
    
    # CENTER: School info
    center_cell = header_table.rows[0].cells[1]
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_para.paragraph_format.space_before = Pt(0)
    center_para.paragraph_format.space_after = Pt(0)
    center_para.paragraph_format.line_spacing = 1.5
    
    name_run = center_para.add_run(school_name + '\n')
    name_run.font.bold = True
    name_run.font.size = Pt(11)
    name_run.font.name = 'Bookman Old Style'
    
    location_text = f"{province} - {district} - {sector_loc} - {cell_loc} - {village}"
    loc_run = center_para.add_run(location_text)
    loc_run.font.size = Pt(8)
    loc_run.font.name = 'Bookman Old Style'
    loc_run.font.bold = True
    
    # RIGHT: School logo
    right_cell = header_table.rows[0].cells[2]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    right_para.paragraph_format.line_spacing = 1.5
    
    school_logo_base64 = data.get('school_logo', '')
    if school_logo_base64 and 'base64' in school_logo_base64:
        try:
            logo_data = school_logo_base64.split(',')[1] if ',' in school_logo_base64 else school_logo_base64
            logo_bytes = base64.b64decode(logo_data)
            temp_logo = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            temp_logo.write(logo_bytes)
            temp_logo.close()
            
            right_run = right_para.add_run()
            right_run.add_picture(temp_logo.name, width=Inches(0.8))
            
            try:
                os.remove(temp_logo.name)
            except:
                pass
            logger.info('✅ School logo added')
        except Exception as e:
            logger.error(f"Logo error: {e}")
            right_run = right_para.add_run('SCHOOL\nLOGO')
            right_run.font.size = Pt(10)
            right_run.font.name = 'Bookman Old Style'
    else:
        right_run = right_para.add_run('SCHOOL\nLOGO')
        right_run.font.size = Pt(8)
        right_run.font.name = 'Bookman Old Style'
    
    # Move header table to beginning
    header_element = header_table._element
    doc._element.body.insert(0, header_element)
    
    # Remove any paragraphs between header and main table
    body = doc._element.body
    elements_to_remove = []
    found_header = False
    for i, element in enumerate(body):
        if element.tag.endswith('tbl') and not found_header:
            found_header = True
        elif found_header and element.tag.endswith('p'):
            elements_to_remove.append(element)
        elif found_header and element.tag.endswith('tbl'):
            break
    
    for element in elements_to_remove:
        body.remove(element)
    
    # Prevent page break between header and main table
    from docx.oxml.shared import OxmlElement
    from docx.oxml.ns import qn
    
    # Set keep_together for header table
    for row in header_table.rows:
        tr = row._element
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
    
    # Set keep_with_next for header table to keep it with main table
    header_para = header_table._element.getparent()
    if header_para is not None:
        pPr = header_para.get_or_add_pPr() if hasattr(header_para, 'get_or_add_pPr') else None
        if pPr is not None:
            keepNext = OxmlElement('w:keepNext')
            pPr.append(keepNext)
    
    logger.info('✅ Header table added at top with no spacing and no page break')
    
    # Main session plan table (first table after moving header)
    table = doc.tables[1]
    logger.info(f"✅ Main table found with {len(table.rows)} rows")
    
    # Row 0: Bold headers
    set_cell_text_with_bold_label(table.rows[0].cells[0], "Sector: ", data.get('sector', '').strip())
    set_cell_text_with_bold_label(table.rows[0].cells[1], "Trade: ", data.get('trade', '').strip())
    set_cell_text_with_bold_label(table.rows[0].cells[4], "Level: ", data.get('level', 'Level 4').strip())
    set_cell_text_with_bold_label(table.rows[0].cells[5], "Date: ", data.get('date', '').strip())
    
    # Row 1: Bold headers
    set_cell_text_with_bold_label(table.rows[1].cells[0], "Trainer name: ", data.get('trainer_name', '').strip())
    
    # School year and term in cell 5
    cell = table.rows[1].cells[5]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r1 = p.add_run("School year: ")
    r1.font.bold = True
    r1.font.name = 'Bookman Old Style'
    r1.font.size = Pt(12)
    r2 = p.add_run(data.get('school_year', '2024-2025').strip())
    r2.font.name = 'Bookman Old Style'
    r2.font.size = Pt(12)
    p.add_run('\n')
    r3 = p.add_run("Term: ")
    r3.font.bold = True
    r3.font.name = 'Bookman Old Style'
    r3.font.size = Pt(12)
    r4 = p.add_run(data.get('term', '').strip())
    r4.font.name = 'Bookman Old Style'
    r4.font.size = Pt(12)
    
    # Row 2: Bold headers
    set_cell_text_with_bold_label(table.rows[2].cells[0], "Module (Code&Name): ", data.get('module_code_title', '').strip())
    set_cell_text_with_bold_label(table.rows[2].cells[1], "Week: ", data.get('week', '1').strip())
    set_cell_text_with_bold_label(table.rows[2].cells[4], "No. Trainees: ", data.get('number_of_trainees', '').strip())
    set_cell_text_with_bold_label(table.rows[2].cells[5], "Class(es): ", data.get('class_name', '').strip())
    
    # Row 3: Learning Outcome
    set_cell_font(table.rows[3].cells[0], bold=True)
    table.rows[3].cells[1].text = data.get('learning_outcomes', '').strip()
    set_cell_font(table.rows[3].cells[1])
    
    # Row 4: Indicative content
    set_cell_font(table.rows[4].cells[0], bold=True)
    table.rows[4].cells[1].text = data.get('indicative_contents', '').strip()
    set_cell_font(table.rows[4].cells[1])
    
    # Row 5: Topic
    set_cell_text_with_bold_label(table.rows[5].cells[0], "Topic of the session: ", data.get('topic_of_session', '').strip())
    
    # Row 6: Range and Duration
    range_val = data.get('range', '').strip()
    set_cell_text_with_bold_label(table.rows[6].cells[0], "Range: ", range_val)
    set_cell_text_with_bold_label(table.rows[6].cells[2], "Duration of the session: ", data.get('duration', '').strip())
    
    # Row 7: Objectives
    objectives = data.get('objectives', '').strip()
    logger.info(f"📝 Filling objectives: {objectives[:100] if objectives else 'EMPTY'}")
    set_cell_text_with_bold_label(table.rows[7].cells[0], "Objectives:\n", objectives)
    
    # Row 8: Facilitation techniques
    facilitation = data.get('facilitation_techniques', '').strip()
    logger.info(f"📝 Facilitation: '{facilitation}'")
    if facilitation:
        set_cell_text_with_bold_label(table.rows[8].cells[0], "Facilitation technique(s): ", facilitation)
    else:
        set_cell_text_with_bold_label(table.rows[8].cells[0], "Facilitation technique(s): ", "Trainer Guided")
    
    # Parse learning activities and resources
    learning_acts = data.get('learning_activities', '')
    resources_text = data.get('resources', '')
    
    # Split activities into sections
    sections = learning_acts.split('\n\n') if learning_acts else []
    intro = ''
    dev = ''
    conclusion = ''
    
    for section in sections:
        if 'Introduction:' in section:
            intro = section.replace('Introduction:', '').strip()
        elif 'Development:' in section:
            dev = section.replace('Development:', '').strip()
        elif 'Conclusion:' in section:
            conclusion = section.replace('Conclusion:', '').strip()
    
    # Parse resources
    resource_lines = [r.strip() for r in resources_text.split('\n') if r.strip()] if resources_text else []
    all_resources = ', '.join(resource_lines) if resource_lines else 'Whiteboard, markers, handouts'
    
    # Row 9: Bold headers
    set_cell_font(table.rows[9].cells[0], bold=True)
    set_cell_font(table.rows[9].cells[3], bold=True)
    if len(table.rows[9].cells) > 5:
        set_cell_font(table.rows[9].cells[5], bold=True)
    
    # Row 10: Introduction
    cell = table.rows[10].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r1 = p.add_run("Trainer's activity: ")
    r1.font.bold = True
    r1.font.name = 'Bookman Old Style'
    r1.font.size = Pt(12)
    r2 = p.add_run(intro + '\n\n')
    r2.font.name = 'Bookman Old Style'
    r2.font.size = Pt(12)
    r3 = p.add_run("Learner's activity: ")
    r3.font.bold = True
    r3.font.name = 'Bookman Old Style'
    r3.font.size = Pt(12)
    r4 = p.add_run("Participate actively and ask questions")
    r4.font.name = 'Bookman Old Style'
    r4.font.size = Pt(12)
    table.rows[10].cells[3].text = all_resources
    set_cell_font(table.rows[10].cells[3])
    if len(table.rows[10].cells) > 5:
        table.rows[10].cells[5].text = "5 min"
        set_cell_font(table.rows[10].cells[5])
    
    # Row 11: Bold header
    set_cell_font(table.rows[11].cells[0], bold=True)
    
    # Row 12: Development
    cell = table.rows[12].cells[0]
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r1 = p.add_run("Step 1:\n")
    r1.font.bold = True
    r1.font.name = 'Bookman Old Style'
    r1.font.size = Pt(12)
    r2 = p.add_run("Trainer's activity: ")
    r2.font.bold = True
    r2.font.name = 'Bookman Old Style'
    r2.font.size = Pt(12)
    r3 = p.add_run(dev + '\n\n')
    r3.font.name = 'Bookman Old Style'
    r3.font.size = Pt(12)
    r4 = p.add_run("Learner's activity: ")
    r4.font.bold = True
    r4.font.name = 'Bookman Old Style'
    r4.font.size = Pt(12)
    r5 = p.add_run("Practice and apply concepts")
    r5.font.name = 'Bookman Old Style'
    r5.font.size = Pt(12)
    table.rows[12].cells[3].text = all_resources
    set_cell_font(table.rows[12].cells[3])
    if len(table.rows[12].cells) > 5:
        table.rows[12].cells[5].text = "30 min"
        set_cell_font(table.rows[12].cells[5])
    
    # Row 15: Bold header
    set_cell_font(table.rows[15].cells[0], bold=True)
    
    # Row 16: Conclusion
    set_cell_text_with_bold_label(table.rows[16].cells[0], "Summary: ", conclusion)
    table.rows[16].cells[3].text = all_resources
    set_cell_font(table.rows[16].cells[3])
    if len(table.rows[16].cells) > 5:
        table.rows[16].cells[5].text = "5 min"
        set_cell_font(table.rows[16].cells[5])
    
    # Row 17: Assessment - bold header
    assessment = data.get('assessment_details', '').strip()
    table.rows[17].cells[0].text = assessment
    set_cell_font(table.rows[17].cells[0])
    table.rows[17].cells[3].text = "Assessment sheets"
    set_cell_font(table.rows[17].cells[3])
    if len(table.rows[17].cells) > 5:
        table.rows[17].cells[5].text = "5 min"
        set_cell_font(table.rows[17].cells[5])
    
    # Row 19: References
    references = data.get('references', '').strip()
    set_cell_text_with_bold_label(table.rows[19].cells[0], "References:\n", references)
    
    # Row 18: Evaluation of the session
    evaluation = data.get('evaluation', '').strip()
    if not evaluation:
        evaluation = "Self-assessment form"
    set_cell_text_with_bold_label(table.rows[18].cells[0], "Evaluation of the session: ", evaluation)
    
    # Row 20: Appendices
    appendix = data.get('appendix', '').strip()
    if not appendix:
        appendix = "PPT, Task Sheets, Assessment"
    set_cell_text_with_bold_label(table.rows[20].cells[0], "Appendices:\n", appendix)
    
    # Save to temp file
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        logger.info(f"💾 Saving to: {temp_file.name}")
        doc.save(temp_file.name)
        temp_file.close()
        logger.info(f"✅ Document saved successfully")
        return temp_file.name
    except Exception as e:
        logger.error(f"❌ Error saving document: {str(e)}")
        raise

def fill_scheme_official(data):
    """Fill Scheme of work.docx from RTB Templates folder"""
    import logging
    logger = logging.getLogger(__name__)
    
    logger.info("=== SCHEME GENERATION START ===")
    logger.info(f"Current dir: {os.getcwd()}")
    logger.info(f"File location: {__file__}")
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    logger.info(f"Base dir: {base_dir}")
    
    template_path = os.path.join(base_dir, 'RTB Templates', 'Scheme of work.docx')
    logger.info(f"Template path: {template_path}")
    logger.info(f"Template exists: {os.path.exists(template_path)}")
    
    # List files in base dir
    try:
        files = os.listdir(base_dir)
        logger.info(f"Files in base dir: {files}")
    except Exception as e:
        logger.error(f"Cannot list base dir: {e}")
    
    # Check RTB Templates folder
    rtb_folder = os.path.join(base_dir, 'RTB Templates')
    if os.path.exists(rtb_folder):
        try:
            rtb_files = os.listdir(rtb_folder)
            logger.info(f"Files in RTB Templates: {rtb_files}")
        except Exception as e:
            logger.error(f"Cannot list RTB Templates: {e}")
    else:
        logger.error(f"RTB Templates folder not found: {rtb_folder}")
    
    if not os.path.exists(template_path):
        logger.error(f"Template not found: {template_path}")
        logger.error("Creating fallback document...")
        doc = Document()
        doc.add_heading('Scheme of Work', 0)
        doc.add_paragraph(f"Province: {data.get('province', '')}")
        doc.add_paragraph(f"District: {data.get('district', '')}")
        doc.add_paragraph(f"School: {data.get('school', '')}")
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        return temp_file.name
    
    doc = Document(template_path)
    logger.info(f"Scheme template loaded, tables: {len(doc.tables)}")
    
    # Create new document from scratch
    from docx import Document as NewDocument
    doc = NewDocument()
    
    school_name = data.get('school_name', '') or data.get('school', '')
    province = data.get('province', '')
    district = data.get('district', '')
    sector_loc = data.get('sector_location', '')
    cell_loc = data.get('cell', '')
    village = data.get('village', '')
    
    # 1. SCHOOL HEADER TABLE (3 columns)
    header_table = doc.add_table(rows=1, cols=3)
    header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    left_cell = header_table.rows[0].cells[0]
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run('RWANDA\nTVET BOARD')
    left_run.font.bold = True
    left_run.font.size = Pt(12)
    left_run.font.name = 'Bookman Old Style'
    
    center_cell = header_table.rows[0].cells[1]
    center_para = center_cell.paragraphs[0]
    center_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = center_para.add_run(school_name + '\n')
    name_run.font.bold = True
    name_run.font.size = Pt(14)
    name_run.font.name = 'Bookman Old Style'
    location_text = f"{province} - {district} - {sector_loc} - {cell_loc} - {village}"
    loc_run = center_para.add_run(location_text)
    loc_run.font.size = Pt(10)
    loc_run.font.name = 'Bookman Old Style'
    loc_run.font.bold = True
    
    right_cell = header_table.rows[0].cells[2]
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    school_logo_base64 = data.get('school_logo', '')
    if school_logo_base64 and ('base64' in school_logo_base64 or school_logo_base64.startswith('data:image')):
        try:
            # Extract base64 data
            if 'base64,' in school_logo_base64:
                logo_data = school_logo_base64.split('base64,')[1]
            elif ',' in school_logo_base64:
                logo_data = school_logo_base64.split(',')[1]
            else:
                logo_data = school_logo_base64
            
            logo_bytes = base64.b64decode(logo_data)
            temp_logo = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            temp_logo.write(logo_bytes)
            temp_logo.close()
            
            right_run = right_para.add_run()
            right_run.add_picture(temp_logo.name, width=Inches(1.2))
            
            try:
                os.remove(temp_logo.name)
            except:
                pass
            logger.info('School logo added to scheme')
        except Exception as e:
            logger.error(f"Logo error: {e}")
            right_run = right_para.add_run('SCHOOL\nLOGO')
            right_run.font.size = Pt(10)
            right_run.font.name = 'Bookman Old Style'
    else:
        right_run = right_para.add_run('SCHOOL\nLOGO')
        right_run.font.size = Pt(10)
        right_run.font.name = 'Bookman Old Style'
    
    # 2. INFO TABLE (8 rows x 6 columns) - Matching official RTB structure
    info_table = doc.add_table(rows=8, cols=6)
    info_table.style = 'Table Grid'
    
    # Row 0: Sector (cols 0-2) | Trainer (cols 3-5)
    info_table.rows[0].cells[0].merge(info_table.rows[0].cells[2])
    info_table.rows[0].cells[3].merge(info_table.rows[0].cells[5])
    set_cell_text_with_bold_label(info_table.rows[0].cells[0], "Sector: ", data.get('sector', ''))
    set_cell_text_with_bold_label(info_table.rows[0].cells[3], "Trainer: ", data.get('trainer_name', ''))
    
    # Row 1: Trade (cols 0-2) | School Year (cols 3-5)
    info_table.rows[1].cells[0].merge(info_table.rows[1].cells[2])
    info_table.rows[1].cells[3].merge(info_table.rows[1].cells[5])
    set_cell_text_with_bold_label(info_table.rows[1].cells[0], "Trade: ", data.get('trade', ''))
    set_cell_text_with_bold_label(info_table.rows[1].cells[3], "School Year: ", data.get('school_year', ''))
    
    # Row 2: Qualification Title (cols 0-2) | Term (cols 3-5)
    info_table.rows[2].cells[0].merge(info_table.rows[2].cells[2])
    info_table.rows[2].cells[3].merge(info_table.rows[2].cells[5])
    set_cell_text_with_bold_label(info_table.rows[2].cells[0], "Qualification Title: ", data.get('qualification_title', ''))
    set_cell_text_with_bold_label(info_table.rows[2].cells[3], "Term: ", data.get('terms', ''))
    
    # Row 3: RQF Level (cols 0-2, rowspan 4) | Module details (cols 3-5, rowspan 4)
    # First merge vertically for RQF Level (rows 3-6, col 0)
    info_table.rows[3].cells[0].merge(info_table.rows[6].cells[0])
    # Merge Module details cell vertically (rows 3-6) and horizontally (cols 1-5)
    info_table.rows[3].cells[1].merge(info_table.rows[6].cells[5])
    set_cell_text_with_bold_label(info_table.rows[3].cells[0], "RQF Level: ", data.get('rqf_level', ''))
    
    # Center-align and format Module details to span entire right section
    module_details_cell = info_table.rows[3].cells[1]
    module_details_para = module_details_cell.paragraphs[0]
    module_details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add Module details header
    run1 = module_details_para.add_run('Module details\n\n')
    run1.font.bold = True
    run1.font.name = 'Bookman Old Style'
    run1.font.size = Pt(12)
    
    # Add module information
    run2 = module_details_para.add_run('Module code and title: ')
    run2.font.bold = True
    run2.font.name = 'Bookman Old Style'
    run2.font.size = Pt(12)
    run3 = module_details_para.add_run(data.get('module_code_title', '') + '\n')
    run3.font.bold = False
    run3.font.name = 'Bookman Old Style'
    run3.font.size = Pt(12)
    
    run4 = module_details_para.add_run('Learning hours: ')
    run4.font.bold = True
    run4.font.name = 'Bookman Old Style'
    run4.font.size = Pt(12)
    run5 = module_details_para.add_run(data.get('module_hours', '') + '\n')
    run5.font.bold = False
    run5.font.name = 'Bookman Old Style'
    run5.font.size = Pt(12)
    
    run6 = module_details_para.add_run('Number of Classes: ')
    run6.font.bold = True
    run6.font.name = 'Bookman Old Style'
    run6.font.size = Pt(12)
    run7 = module_details_para.add_run(data.get('number_of_classes', ''))
    run7.font.bold = False
    run7.font.name = 'Bookman Old Style'
    run7.font.size = Pt(12)
    
    # Row 7: Date (cols 0-2) | Class Name label | value (cols 3-5)
    info_table.rows[7].cells[0].merge(info_table.rows[7].cells[2])
    date_value = data.get('date', '').strip()
    if not date_value:
        date_value = datetime.now().strftime('%d/%m/%Y')
    set_cell_text_with_bold_label(info_table.rows[7].cells[0], "Date: ", date_value)
    info_table.rows[7].cells[3].text = 'Class Name:'
    set_cell_font(info_table.rows[7].cells[3], bold=True)
    info_table.rows[7].cells[4].merge(info_table.rows[7].cells[5])
    info_table.rows[7].cells[4].text = data.get('class_name', '')
    set_cell_font(info_table.rows[7].cells[4], bold=False)
    
    # 3. TERM 1 TABLE (matching official RTB structure)
    # Create table with proper structure: 2 header rows + data rows
    term1_table = doc.add_table(rows=2, cols=9)
    term1_table.style = 'Table Grid'
    
    # Header Row 0: Weeks | Competence (colspan 3) | Activities | Resources | Assessment | Place | Observation
    term1_table.rows[0].cells[0].merge(term1_table.rows[1].cells[0])  # Weeks spans 2 rows
    term1_table.rows[0].cells[1].merge(term1_table.rows[0].cells[3])  # Competence spans 3 cols
    term1_table.rows[0].cells[4].merge(term1_table.rows[1].cells[4])  # Activities spans 2 rows
    term1_table.rows[0].cells[5].merge(term1_table.rows[1].cells[5])  # Resources spans 2 rows
    term1_table.rows[0].cells[6].merge(term1_table.rows[1].cells[6])  # Assessment spans 2 rows
    term1_table.rows[0].cells[7].merge(term1_table.rows[1].cells[7])  # Place spans 2 rows
    term1_table.rows[0].cells[8].merge(term1_table.rows[1].cells[8])  # Observation spans 2 rows
    
    # Set header texts and formatting
    term1_table.rows[0].cells[0].text = 'Weeks'
    term1_table.rows[0].cells[1].text = 'Competence code and name'
    term1_table.rows[0].cells[4].text = 'Learning Activities'
    term1_table.rows[0].cells[5].text = 'Resources (Equipment, tools, and materials)'
    term1_table.rows[0].cells[6].text = 'Evidences of formative assessment'
    term1_table.rows[0].cells[7].text = 'Learning Place'
    term1_table.rows[0].cells[8].text = 'Observation'
    
    # Header Row 1: LO | Duration | IC (under Competence)
    term1_table.rows[1].cells[1].text = 'Learning outcome (LO)'
    term1_table.rows[1].cells[2].text = 'Duration'
    term1_table.rows[1].cells[3].text = 'Indicative content (IC)'
    
    # Apply light green background and bold to headers
    from docx.oxml.shared import OxmlElement
    def set_cell_background(cell, color):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._element.get_or_add_tcPr().append(shading)
    
    for row_idx in [0, 1]:
        for cell in term1_table.rows[row_idx].cells:
            set_cell_background(cell, 'D4EDDA')
            set_cell_font(cell, bold=True)
    
    # Add data rows for Term 1
    term1_weeks = data.get('term1_weeks', '')
    term1_competence = data.get('term1_competence', '')
    term1_outcomes_raw = data.get('term1_learning_outcomes', '')
    term1_contents_raw = data.get('term1_indicative_contents', '')
    term1_duration = data.get('term1_duration', '')
    
    term1_outcomes = [lo.strip() for lo in term1_outcomes_raw.replace('\r\n', '\n').split('\n') if lo.strip()]
    term1_contents = [ic.strip() for ic in term1_contents_raw.replace('\r\n', '\n').split('\n') if ic.strip()]
    
    # Add rows for each LO/IC pair
    for i, (lo, ic) in enumerate(zip(term1_outcomes, term1_contents)):
        row = term1_table.add_row()
        row.cells[0].text = term1_weeks if i == 0 else ''
        row.cells[1].text = lo
        row.cells[2].text = term1_duration
        row.cells[3].text = ic
        row.cells[4].text = 'Practical exercises, Group work, Demonstrations'
        row.cells[5].text = 'Textbooks, Computers, Tools, Materials'
        row.cells[6].text = 'Observation, Practical tests, Written tests'
        row.cells[7].text = 'Workshop'
        row.cells[8].text = ''
        for cell in row.cells:
            set_cell_font(cell, bold=False)
    
    # Add integrated assessment row
    assess_row = term1_table.add_row()
    assess_row.cells[0].text = ''
    assess_row.cells[1].text = f'Integrated Assessment (for {term1_competence})'
    assess_row.cells[2].text = ''
    assess_row.cells[3].text = ''
    assess_row.cells[4].text = ''
    assess_row.cells[5].text = ''
    assess_row.cells[6].text = ''
    assess_row.cells[7].text = ''
    assess_row.cells[8].text = ''
    set_cell_font(assess_row.cells[1], bold=True)
    
    if False and len(doc.tables) > 1:
        h = doc.tables[1]
        try:
            # Row 0: Sector (cell 1), Trainer (cell 3)
            h.rows[0].cells[1].text = data.get('sector') or ''
            set_cell_font(h.rows[0].cells[1], bold=False)
            h.rows[0].cells[3].text = data.get('trainer_name') or ''
            set_cell_font(h.rows[0].cells[3], bold=False)
            
            # Row 1: Trade (cell 1), School Year (cell 3)
            h.rows[1].cells[1].text = data.get('department_trade') or data.get('trade') or ''
            set_cell_font(h.rows[1].cells[1], bold=False)
            h.rows[1].cells[3].text = data.get('school_year') or ''
            set_cell_font(h.rows[1].cells[3], bold=False)
            
            # Row 2: Qualification (cell 1), Term (cell 3)
            h.rows[2].cells[1].text = data.get('qualification_title') or ''
            set_cell_font(h.rows[2].cells[1], bold=False)
            h.rows[2].cells[3].text = data.get('terms') or ''
            set_cell_font(h.rows[2].cells[3], bold=False)
            
            # Row 3: RQF Level (cell 1)
            h.rows[3].cells[1].text = data.get('rqf_level') or ''
            set_cell_font(h.rows[3].cells[1], bold=False)
            
            # Row 4: Module code and title (cell 3)
            h.rows[4].cells[3].text = data.get('module_code_title') or ''
            set_cell_font(h.rows[4].cells[3], bold=False)
            
            # Row 5: Learning hours (cell 3)
            h.rows[5].cells[3].text = data.get('module_hours') or ''
            set_cell_font(h.rows[5].cells[3], bold=False)
            
            # Row 6: Number of Classes (cell 3)
            h.rows[6].cells[3].text = data.get('number_of_classes') or ''
            set_cell_font(h.rows[6].cells[3], bold=False)
            
            # Row 7: Date (cell 1), Class Name (cell 3)
            date_value = data.get('date', '').strip()
            if not date_value:
                date_value = datetime.now().strftime('%d/%m/%Y')
            h.rows[7].cells[1].text = date_value
            set_cell_font(h.rows[7].cells[1], bold=False)
            h.rows[7].cells[3].text = data.get('class_name') or ''
            set_cell_font(h.rows[7].cells[3], bold=False)
            
            logger.info("Header table filled successfully")
        except Exception as e:
            logger.error(f"Error filling header: {e}")
            import traceback
            logger.error(traceback.format_exc())
    
    # 4. TERM 2 TABLE (same structure as Term 1)
    term2_table = doc.add_table(rows=2, cols=9)
    term2_table.style = 'Table Grid'
    
    # Header setup for Term 2
    term2_table.rows[0].cells[0].merge(term2_table.rows[1].cells[0])
    term2_table.rows[0].cells[1].merge(term2_table.rows[0].cells[3])
    term2_table.rows[0].cells[4].merge(term2_table.rows[1].cells[4])
    term2_table.rows[0].cells[5].merge(term2_table.rows[1].cells[5])
    term2_table.rows[0].cells[6].merge(term2_table.rows[1].cells[6])
    term2_table.rows[0].cells[7].merge(term2_table.rows[1].cells[7])
    term2_table.rows[0].cells[8].merge(term2_table.rows[1].cells[8])
    
    term2_table.rows[0].cells[0].text = 'Weeks'
    term2_table.rows[0].cells[1].text = 'Competence code and name'
    term2_table.rows[0].cells[4].text = 'Learning Activities'
    term2_table.rows[0].cells[5].text = 'Resources (Equipment, tools, and materials)'
    term2_table.rows[0].cells[6].text = 'Evidences of formative assessment'
    term2_table.rows[0].cells[7].text = 'Learning Place'
    term2_table.rows[0].cells[8].text = 'Observation'
    term2_table.rows[1].cells[1].text = 'Learning outcome (LO)'
    term2_table.rows[1].cells[2].text = 'Duration'
    term2_table.rows[1].cells[3].text = 'Indicative content (IC)'
    
    for row_idx in [0, 1]:
        for cell in term2_table.rows[row_idx].cells:
            set_cell_background(cell, 'D4EDDA')
            set_cell_font(cell, bold=True)
    
    term2_weeks = data.get('term2_weeks', '')
    term2_competence = data.get('term2_competence', '')
    term2_outcomes_raw = data.get('term2_learning_outcomes', '')
    term2_contents_raw = data.get('term2_indicative_contents', '')
    term2_duration = data.get('term2_duration', '')
    term2_outcomes = [lo.strip() for lo in term2_outcomes_raw.replace('\r\n', '\n').split('\n') if lo.strip()]
    term2_contents = [ic.strip() for ic in term2_contents_raw.replace('\r\n', '\n').split('\n') if ic.strip()]
    
    for i, (lo, ic) in enumerate(zip(term2_outcomes, term2_contents)):
        row = term2_table.add_row()
        row.cells[0].text = term2_weeks if i == 0 else ''
        row.cells[1].text = lo
        row.cells[2].text = term2_duration
        row.cells[3].text = ic
        row.cells[4].text = 'Practical exercises, Group work, Demonstrations'
        row.cells[5].text = 'Textbooks, Computers, Tools, Materials'
        row.cells[6].text = 'Observation, Practical tests, Written tests'
        row.cells[7].text = 'Workshop'
        row.cells[8].text = ''
        for cell in row.cells:
            set_cell_font(cell, bold=False)
    
    assess_row = term2_table.add_row()
    assess_row.cells[1].text = f'Integrated Assessment (for {term2_competence})'
    set_cell_font(assess_row.cells[1], bold=True)
    
    if False and len(doc.tables) > 3:
        table2 = doc.tables[3]
        
        pass
    
    # 5. TERM 3 TABLE (same structure)
    term3_table = doc.add_table(rows=2, cols=9)
    term3_table.style = 'Table Grid'
    
    term3_table.rows[0].cells[0].merge(term3_table.rows[1].cells[0])
    term3_table.rows[0].cells[1].merge(term3_table.rows[0].cells[3])
    term3_table.rows[0].cells[4].merge(term3_table.rows[1].cells[4])
    term3_table.rows[0].cells[5].merge(term3_table.rows[1].cells[5])
    term3_table.rows[0].cells[6].merge(term3_table.rows[1].cells[6])
    term3_table.rows[0].cells[7].merge(term3_table.rows[1].cells[7])
    term3_table.rows[0].cells[8].merge(term3_table.rows[1].cells[8])
    
    term3_table.rows[0].cells[0].text = 'Weeks'
    term3_table.rows[0].cells[1].text = 'Competence code and name'
    term3_table.rows[0].cells[4].text = 'Learning Activities'
    term3_table.rows[0].cells[5].text = 'Resources (Equipment, tools, and materials)'
    term3_table.rows[0].cells[6].text = 'Evidences of formative assessment'
    term3_table.rows[0].cells[7].text = 'Learning Place'
    term3_table.rows[0].cells[8].text = 'Observation'
    term3_table.rows[1].cells[1].text = 'Learning outcome (LO)'
    term3_table.rows[1].cells[2].text = 'Duration'
    term3_table.rows[1].cells[3].text = 'Indicative content (IC)'
    
    for row_idx in [0, 1]:
        for cell in term3_table.rows[row_idx].cells:
            set_cell_background(cell, 'D4EDDA')
            set_cell_font(cell, bold=True)
    
    term3_weeks = data.get('term3_weeks', '')
    term3_competence = data.get('term3_competence', '')
    term3_outcomes_raw = data.get('term3_learning_outcomes', '')
    term3_contents_raw = data.get('term3_indicative_contents', '')
    term3_duration = data.get('term3_duration', '')
    term3_outcomes = [lo.strip() for lo in term3_outcomes_raw.replace('\r\n', '\n').split('\n') if lo.strip()]
    term3_contents = [ic.strip() for ic in term3_contents_raw.replace('\r\n', '\n').split('\n') if ic.strip()]
    
    for i, (lo, ic) in enumerate(zip(term3_outcomes, term3_contents)):
        row = term3_table.add_row()
        row.cells[0].text = term3_weeks if i == 0 else ''
        row.cells[1].text = lo
        row.cells[2].text = term3_duration
        row.cells[3].text = ic
        row.cells[4].text = 'Practical exercises, Group work, Demonstrations'
        row.cells[5].text = 'Textbooks, Computers, Tools, Materials'
        row.cells[6].text = 'Observation, Practical tests, Written tests'
        row.cells[7].text = 'Workshop'
        row.cells[8].text = ''
        for cell in row.cells:
            set_cell_font(cell, bold=False)
    
    assess_row = term3_table.add_row()
    assess_row.cells[1].text = f'Integrated Assessment (for {term3_competence})'
    set_cell_font(assess_row.cells[1], bold=True)
    
    # 6. SIGNATURE TABLE
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'
    
    # Row 0: Prepared by
    sig_table.rows[0].cells[0].text = 'Prepared by: (Name, position and Signature)'
    set_cell_font(sig_table.rows[0].cells[0], bold=True)
    prepared_name = data.get('prepared_by_name', '') or data.get('trainer_name', '')
    prepared_pos = data.get('prepared_by_position', '') or 'Trainer'
    sig_table.rows[0].cells[1].text = f"{prepared_name}, {prepared_pos}"
    set_cell_font(sig_table.rows[0].cells[1], bold=False)
    
    # Row 1: Verified by
    sig_table.rows[1].cells[0].text = 'Verified by: (Name, position and Signature)'
    set_cell_font(sig_table.rows[1].cells[0], bold=True)
    verified_name = data.get('verified_by_name', '') or data.get('dos_name', '')
    verified_pos = data.get('verified_by_position', '') or 'Director of Studies'
    sig_table.rows[1].cells[1].text = f"{verified_name}, {verified_pos}"
    set_cell_font(sig_table.rows[1].cells[1], bold=False)
    
    # Row 2: Approved by
    sig_table.rows[2].cells[0].text = 'Approved by: (Name, position and Signature)'
    set_cell_font(sig_table.rows[2].cells[0], bold=True)
    approved_name = data.get('approved_by_name', '') or data.get('manager_name', '')
    approved_pos = data.get('approved_by_position', '') or 'School Manager'
    sig_table.rows[2].cells[1].text = f"{approved_name}, {approved_pos}"
    set_cell_font(sig_table.rows[2].cells[1], bold=False)
    
    # Save document
    logger.info('Scheme of work created with proper table structure')
    try:
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        logger.info(f"Saving to: {temp_file.name}")
        doc.save(temp_file.name)
        temp_file.close()
        
        if os.path.exists(temp_file.name):
            size = os.path.getsize(temp_file.name)
            logger.info(f"Scheme saved successfully: {size} bytes")
        else:
            logger.error("Saved file does not exist!")
        
        return temp_file.name
    except Exception as e:
        logger.error(f"Error saving scheme: {e}")
        import traceback
        logger.error(traceback.format_exc())
        raise
