from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import tempfile
import logging

logger = logging.getLogger(__name__)

def generate_rtb_session_plan(data):
    """Generate RTB-compliant session plan matching official template"""
    try:
        doc = Document()
        
        # Helper function
        def get_val(key, default=''):
            if hasattr(data, key):
                return str(getattr(data, key) or default)
            elif isinstance(data, dict):
                return str(data.get(key, default))
            return default
        
        # Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run("REPUBLIC OF RWANDA\nMINISTRY OF EDUCATION\nRWANDA TECHNICAL BOARD (RTB)")
        header_run.bold = True
        header_run.font.size = Pt(12)
        
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("SESSION PLAN")
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        doc.add_paragraph()
        
        # Basic Information Table
        table = doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'
        
        # Row 1
        cells = table.rows[0].cells
        cells[0].text = 'Sector:'
        cells[1].text = get_val('sector')
        cells[2].text = 'Sub-Sector:'
        cells[3].text = get_val('sub_sector')
        
        # Row 2
        cells = table.rows[1].cells
        cells[0].text = 'Trade:'
        cells[1].text = get_val('trade')
        cells[2].text = 'Qualification Title:'
        cells[3].text = get_val('qualification_title')
        
        # Row 3
        cells = table.rows[2].cells
        cells[0].text = 'RQF Level:'
        cells[1].text = get_val('rqf_level')
        cells[2].text = 'Module Code & Title:'
        cells[3].text = get_val('module_code_title')
        
        # Row 4
        cells = table.rows[3].cells
        cells[0].text = 'Term:'
        cells[1].text = get_val('term')
        cells[2].text = 'Week:'
        cells[3].text = get_val('week')
        
        # Row 5
        cells = table.rows[4].cells
        cells[0].text = 'Date:'
        cells[1].text = get_val('date')
        cells[2].text = 'Trainer Name:'
        cells[3].text = get_val('trainer_name')
        
        # Row 6
        cells = table.rows[5].cells
        cells[0].text = 'Class:'
        cells[1].text = get_val('class_name')
        cells[2].text = 'Number of Trainees:'
        cells[3].text = get_val('number_of_trainees')
        
        # Row 7
        cells = table.rows[6].cells
        cells[0].text = 'Learning Outcomes:'
        cells[1].text = get_val('learning_outcomes')
        cells[2].text = 'Indicative Contents:'
        cells[3].text = get_val('indicative_contents')
        
        # Row 8
        cells = table.rows[7].cells
        cells[0].text = 'Topic of Session:'
        cells[1].text = get_val('topic_of_session')
        cells[2].text = 'Duration:'
        cells[3].text = f"{get_val('duration', '40')} minutes"
        
        doc.add_paragraph()
        
        # Session Details Section
        doc.add_heading('SESSION DETAILS', level=1)
        
        # Objectives
        obj_para = doc.add_paragraph()
        obj_para.add_run('Objectives: ').bold = True
        obj_para.add_run(get_val('objectives', 'By the end of this session, trainees will be able to understand and apply the concepts covered.'))
        
        # Facilitation Techniques
        fac_para = doc.add_paragraph()
        fac_para.add_run('Facilitation Techniques: ').bold = True
        fac_para.add_run(get_val('facilitation_techniques', 'Interactive discussions, practical demonstrations, and hands-on activities.'))
        
        # Learning Activities
        act_para = doc.add_paragraph()
        act_para.add_run('Learning Activities: ').bold = True
        act_para.add_run(get_val('learning_activities', 'Group work, individual practice, and collaborative problem-solving exercises.'))
        
        # Resources
        res_para = doc.add_paragraph()
        res_para.add_run('Resources: ').bold = True
        res_para.add_run(get_val('resources', 'Textbooks, computers, projector, and relevant materials.'))
        
        # Assessment
        ass_para = doc.add_paragraph()
        ass_para.add_run('Assessment: ').bold = True
        ass_para.add_run(get_val('assessment_details', 'Continuous assessment through observation, questioning, and practical tasks.'))
        
        # References
        ref_para = doc.add_paragraph()
        ref_para.add_run('References: ').bold = True
        ref_para.add_run(get_val('references', 'RTB curriculum guidelines and approved learning materials.'))
        
        # Save document
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        logger.info(f"RTB Session Plan generated: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating RTB session plan: {e}")
        # Fallback
        doc = Document()
        doc.add_heading('RTB SESSION PLAN', 0)
        doc.add_paragraph('Error generating document. Please contact support.')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name

def generate_rtb_scheme_of_work(data):
    """Generate RTB-compliant scheme of work matching official template"""
    try:
        doc = Document()
        
        # Helper function
        def get_val(key, default=''):
            if hasattr(data, key):
                return str(getattr(data, key) or default)
            elif isinstance(data, dict):
                return str(data.get(key, default))
            return default
        
        # Header
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_para.add_run("REPUBLIC OF RWANDA\nMINISTRY OF EDUCATION\nRWANDA TECHNICAL BOARD (RTB)")
        header_run.bold = True
        header_run.font.size = Pt(12)
        
        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("SCHEME OF WORK")
        title_run.bold = True
        title_run.font.size = Pt(16)
        
        doc.add_paragraph()
        
        # Basic Information Table
        info_table = doc.add_table(rows=6, cols=4)
        info_table.style = 'Table Grid'
        
        # Row 1
        cells = info_table.rows[0].cells
        cells[0].text = 'Province:'
        cells[1].text = get_val('province')
        cells[2].text = 'District:'
        cells[3].text = get_val('district')
        
        # Row 2
        cells = info_table.rows[1].cells
        cells[0].text = 'Sector:'
        cells[1].text = get_val('sector')
        cells[2].text = 'School:'
        cells[3].text = get_val('school')
        
        # Row 3
        cells = info_table.rows[2].cells
        cells[0].text = 'Department/Trade:'
        cells[1].text = get_val('department_trade')
        cells[2].text = 'Qualification Title:'
        cells[3].text = get_val('qualification_title')
        
        # Row 4
        cells = info_table.rows[3].cells
        cells[0].text = 'RQF Level:'
        cells[1].text = get_val('rqf_level')
        cells[2].text = 'Module Code & Title:'
        cells[3].text = get_val('module_code_title')
        
        # Row 5
        cells = info_table.rows[4].cells
        cells[0].text = 'School Year:'
        cells[1].text = get_val('school_year')
        cells[2].text = 'Terms:'
        cells[3].text = get_val('terms', '3')
        
        # Row 6
        cells = info_table.rows[5].cells
        cells[0].text = 'Module Hours:'
        cells[1].text = get_val('module_hours')
        cells[2].text = 'Number of Classes:'
        cells[3].text = get_val('number_of_classes')
        
        doc.add_paragraph()
        
        # Term Details Table
        doc.add_heading('TERM BREAKDOWN', level=1)
        
        term_table = doc.add_table(rows=4, cols=4)
        term_table.style = 'Table Grid'
        
        # Header row
        hdr_cells = term_table.rows[0].cells
        hdr_cells[0].text = 'Term'
        hdr_cells[1].text = 'Weeks'
        hdr_cells[2].text = 'Learning Outcomes'
        hdr_cells[3].text = 'Indicative Contents'
        
        # Make header bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Term 1
        cells = term_table.rows[1].cells
        cells[0].text = 'Term 1'
        cells[1].text = get_val('term1_weeks', '1-12')
        cells[2].text = get_val('term1_learning_outcomes')
        cells[3].text = get_val('term1_indicative_contents')
        
        # Term 2
        cells = term_table.rows[2].cells
        cells[0].text = 'Term 2'
        cells[1].text = get_val('term2_weeks', '13-24')
        cells[2].text = get_val('term2_learning_outcomes')
        cells[3].text = get_val('term2_indicative_contents')
        
        # Term 3
        cells = term_table.rows[3].cells
        cells[0].text = 'Term 3'
        cells[1].text = get_val('term3_weeks', '25-36')
        cells[2].text = get_val('term3_learning_outcomes')
        cells[3].text = get_val('term3_indicative_contents')
        
        doc.add_paragraph()
        
        # Signatures Section
        doc.add_heading('SIGNATURES', level=1)
        
        sig_table = doc.add_table(rows=3, cols=3)
        sig_table.style = 'Table Grid'
        
        # Headers
        hdr_cells = sig_table.rows[0].cells
        hdr_cells[0].text = 'Role'
        hdr_cells[1].text = 'Name'
        hdr_cells[2].text = 'Signature & Date'
        
        # Trainer
        cells = sig_table.rows[1].cells
        cells[0].text = 'Trainer:'
        cells[1].text = get_val('trainer_name')
        cells[2].text = ''
        
        # DOS
        cells = sig_table.rows[2].cells
        cells[0].text = 'DOS:'
        cells[1].text = get_val('dos_name')
        cells[2].text = ''
        
        # Save document
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        logger.info(f"RTB Scheme of Work generated: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating RTB scheme: {e}")
        # Fallback
        doc = Document()
        doc.add_heading('RTB SCHEME OF WORK', 0)
        doc.add_paragraph('Error generating document. Please contact support.')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name