from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """Set cell borders"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), 'single')
            edge_el.set(qn('w:sz'), '4')
            edge_el.set(qn('w:space'), '0')
            edge_el.set(qn('w:color'), '000000')
            tcBorders.append(edge_el)
    tcPr.append(tcBorders)

def merge_cells_in_row(table, row_idx, start_col, end_col):
    """Merge cells horizontally"""
    cell = table.rows[row_idx].cells[start_col]
    for col in range(start_col + 1, end_col + 1):
        cell.merge(table.rows[row_idx].cells[col])
    return cell

def generate_rtb_session_plan(data):
    """Generate RTB-compliant session plan matching official template"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    # Main table: 23 rows x 6 columns
    table = doc.add_table(rows=23, cols=6)
    table.style = 'Table Grid'
    
    # Row 0: Empty header row
    
    # Row 1: Sector, Sub-sector (trade), Date
    merge_cells_in_row(table, 1, 1, 3).text = f"Sub-sector: {data.get('trade', 'N/A')}"
    merge_cells_in_row(table, 1, 4, 5).text = f"Date : {data.get('date', 'N/A')}"
    table.rows[1].cells[0].text = f"Sector :    {data.get('sector', 'N/A')}"
    
    # Row 2: Trainer name, Term
    merge_cells_in_row(table, 2, 0, 3).text = f"Lead Trainer's name : {data.get('trainer_name', 'N/A')}"
    merge_cells_in_row(table, 2, 4, 5).text = f"TERM : {data.get('term', 'I')}"
    
    # Row 3: Module, Week, Trainees, Class
    table.rows[3].cells[0].text = f"Module(Code&Name): {data.get('module_code_title', 'N/A')}"
    merge_cells_in_row(table, 3, 1, 2).text = f"Week : {data.get('week', 'I')}"
    table.rows[3].cells[3].text = f"No. Trainees: {data.get('number_of_trainees', '0')}"
    merge_cells_in_row(table, 3, 4, 5).text = f"Class(es): {data.get('class_name', '1')}"
    
    # Row 4: Learning outcome
    table.rows[4].cells[0].text = "Learning outcome:"
    merge_cells_in_row(table, 4, 1, 5).text = data.get('learning_outcomes', 'N/A')
    
    # Row 5: Indicative contents
    table.rows[5].cells[0].text = "Indicative contents:"
    merge_cells_in_row(table, 5, 1, 5).text = data.get('indicative_contents', 'N/A')
    
    # Row 6: Topic
    merge_cells_in_row(table, 6, 0, 5).text = f"Topic of the session: {data.get('topic_of_session', 'N/A')}"
    
    # Row 7: Range and Duration
    table.rows[7].cells[0].text = f"Range: \n{data.get('rqf_level', 'N/A')}"
    merge_cells_in_row(table, 7, 1, 5).text = f"Duration of the session: {data.get('duration', '40')}min"
    
    # Row 8: Objectives
    merge_cells_in_row(table, 8, 0, 5).text = f"Objectives: By the end of this session every learner should be able to:\n{data.get('learning_outcomes', 'N/A')}"
    
    # Row 9: Facilitation technique
    merge_cells_in_row(table, 9, 0, 5).text = f"Facilitation technique(s):   {data.get('facilitation_techniques', 'N/A')}"
    
    # Row 10: Headers for Introduction
    table.rows[10].cells[0].text = "Introduction"
    merge_cells_in_row(table, 10, 1, 1).text = "Introduction"
    merge_cells_in_row(table, 10, 2, 4).text = "Resources"
    table.rows[10].cells[5].text = "Duration"
    
    # Row 11: Introduction content
    merge_cells_in_row(table, 11, 0, 1).text = "Trainer's activity: \nGreets and makes roll calls\nIntroduces topic and objectives"
    merge_cells_in_row(table, 11, 2, 4).text = "Attendance sheet\nPPT\nProjector\nWhiteboard"
    table.rows[11].cells[5].text = "5 minutes"
    
    # Row 12: Development header
    merge_cells_in_row(table, 12, 0, 5).text = "Development/Body"
    
    # Rows 13-15: Development steps (simplified)
    steps = [
        ("Introduction to topic", "Explains key concepts and demonstrates examples"),
        ("Practical application", "Guides learners through hands-on activities"),
        ("Group work and practice", "Assigns tasks and monitors learner progress")
    ]
    for i, (title, activity) in enumerate(steps):
        row_idx = 13 + i
        merge_cells_in_row(table, row_idx, 0, 1).text = f"Step {i+1}: {title}\nTrainer's activity: \n{activity}"
        merge_cells_in_row(table, row_idx, 2, 4).text = "Computer\nProjector\nPPT\nLearning materials"
        table.rows[row_idx].cells[5].text = f"{int(int(data.get('duration', 40)) * 0.6 / 3)}\nminutes"
    
    # Row 16: Conclusion header
    merge_cells_in_row(table, 16, 0, 5).text = "Conclusion"
    
    # Row 17: Summary
    merge_cells_in_row(table, 17, 0, 1).text = "Summary:\nTrainer involves learners to summarize key points learned"
    merge_cells_in_row(table, 17, 2, 4).text = "Computer\nProjector"
    table.rows[17].cells[5].text = "3 minutes"
    
    # Row 18: Assessment
    merge_cells_in_row(table, 18, 0, 1).text = "Assessment/Assignment\nTrainer's activity: \nProvides assessment questions or assignment"
    merge_cells_in_row(table, 18, 2, 4).text = "Assessment sheets\nQuestion papers"
    table.rows[18].cells[5].text = "5 minutes"
    
    # Row 19: Evaluation
    merge_cells_in_row(table, 19, 0, 1).text = "Evaluation of the session:\nTrainer's activity: \nCollects feedback from learners"
    merge_cells_in_row(table, 19, 2, 4).text = "Self-assessment form\nFeedback forms"
    table.rows[19].cells[5].text = "2 minutes"
    
    # Row 20: References
    merge_cells_in_row(table, 20, 0, 5).text = "References:\nRTB Curriculum Guidelines\nModule Learning Materials"
    
    # Row 21: Appendices
    merge_cells_in_row(table, 21, 0, 5).text = "Appendices: PPT, Task Sheets, Assessment, Learning Materials"
    
    # Row 22: Reflection
    merge_cells_in_row(table, 22, 0, 5).text = "Reflection :\n"
    
    # Format all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
    
    return doc

def generate_rtb_scheme_of_work(data):
    """Generate RTB-compliant scheme of work matching official template"""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
        section.page_width = Cm(29.7)  # A4 landscape
        section.page_height = Cm(21.0)
    
    # Header information
    doc.add_heading('SCHEME OF WORK', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # School info table
    info_table = doc.add_table(rows=8, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Province:', data.get('province', 'N/A')),
        ('District:', data.get('district', 'N/A')),
        ('School:', data.get('school', 'N/A')),
        ('Sector:', data.get('sector', 'N/A')),
        ('Module Code & Title:', data.get('module_code_title', 'N/A')),
        ('RQF Level:', data.get('rqf_level', 'N/A')),
        ('Trainer Name:', data.get('trainer_name', 'N/A')),
        ('School Year:', data.get('school_year', 'N/A'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Main scheme table: 9 columns
    scheme_table = doc.add_table(rows=2, cols=9)
    scheme_table.style = 'Table Grid'
    
    # Header row 1
    headers1 = ['Weeks', 'Competence code and name', '', '', 'Learning Activities', 
                'Resources (Equipment, tools, and materials)', 'Evidences of formative assessment', 
                'Learning Place', 'Observation']
    for i, header in enumerate(headers1):
        scheme_table.rows[0].cells[i].text = header
    
    # Header row 2
    headers2 = ['Weeks', 'Learning outcome (LO)', 'Duration', 'Indicative content (IC)', 
                'Learning Activities', 'Resources (Equipment, tools, and materials)', 
                'Evidences of formative assessment', 'Learning Place', 'Observation']
    for i, header in enumerate(headers2):
        scheme_table.rows[1].cells[i].text = header
        scheme_table.rows[1].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add term data from wizard
    terms = [
        ('term1', 'Term 1'),
        ('term2', 'Term 2'),
        ('term3', 'Term 3')
    ]
    
    for term_key, term_name in terms:
        weeks = data.get(f'{term_key}_weeks', '')
        los = data.get(f'{term_key}_learning_outcomes', '')
        duration = data.get(f'{term_key}_duration', '')
        ics = data.get(f'{term_key}_indicative_contents', '')
        
        if weeks and los:  # Only add if data exists
            row = scheme_table.add_row()
            row.cells[0].text = weeks
            row.cells[1].text = los
            row.cells[2].text = duration
            row.cells[3].text = ics
            row.cells[4].text = data.get('delivery_approach', 'Lectures, Practicals, Group work')
            row.cells[5].text = data.get('resource_inventory', 'Computer, Projector, Learning materials')
            row.cells[6].text = data.get('formative_assessment', 'Continuous assessment')
            row.cells[7].text = 'Lab/Classroom'
            row.cells[8].text = term_name
    
    # Format all cells
    for row in scheme_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)
    
    return doc
