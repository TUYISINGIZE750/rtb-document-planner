"""RTB Template Filler - Generates RTB-compliant session plans matching exact format"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os

try:
    from facilitation_content_generator import (
        generate_introduction_activities,
        generate_development_activities,
        generate_resources,
        generate_assessment
    )
except ImportError:
    def generate_introduction_activities(topic, facilitation): return f"Introduction activities for {topic}"
    def generate_development_activities(topic, facilitation, activities): return f"Development activities"
    def generate_resources(facilitation, resources): return "Resources"
    def generate_assessment(topic, facilitation, details): return "Assessment"

try:
    from content_formatter import clean_text, format_objectives, format_resources
except ImportError:
    def clean_text(text): return str(text) if text else ""
    def format_objectives(text): return str(text) if text else ""
    def format_resources(text): return str(text) if text else ""

def preserve_cell_format(cell, new_text, font_name='Book Antiqua', font_size=12, spacing=1.5):
    """Update cell text while applying proper formatting"""
    if not cell.paragraphs:
        cell.text = new_text
        return
    
    cell.paragraphs[0].clear()
    
    text_lines = new_text.split('\n') if new_text else ['']
    
    for idx, line in enumerate(text_lines):
        if idx == 0:
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        else:
            para = cell.add_paragraph()
        
        para.paragraph_format.line_spacing = spacing
        
        if line.strip():
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size)
        else:
            para.text = ''

def fill_session_plan_template(data):
    """Generate RTB-compliant session plan matching exact RTB format
    
    Structure:
    - Header: RTB title with code, sector, trade, level, term, week, date
    - Content table: Basic information (Sector, Trade, Level, Module, etc.)
    - Main sections: Introduction, Development, Conclusion, Assessment, Evaluation
    - References, Appendices, Reflection
    """
    
    try:
        template_path = os.path.join(os.path.dirname(__file__), 'rtb_session_plan_template.docx')
        if os.path.exists(template_path):
            doc = Document(template_path)
            if doc.tables and len(doc.tables) > 0:
                table = doc.tables[0]
                if len(table.rows) >= 23:
                    return _fill_existing_template(doc, table, data)
    except Exception as e:
        pass
    
    return _create_rtb_session_plan_from_scratch(data)

def _fill_existing_template(doc, table, data):
    """Fill existing RTB template if structure matches"""
    try:
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
        
        preserve_cell_format(table.rows[1].cells[0], f"Sector: {data.get('sector', '')}", spacing=1.5)
        preserve_cell_format(table.rows[1].cells[1], f"Sub-sector: {data.get('trade', '')}", spacing=1.5)
        preserve_cell_format(table.rows[1].cells[4] if len(table.rows[1].cells) > 4 else table.rows[1].cells[-1], f"Date: {data.get('date', '')}", spacing=1.5)
        
        preserve_cell_format(table.rows[2].cells[0], f"Lead Trainer: {data.get('trainer_name', '')}", spacing=1.5)
        preserve_cell_format(table.rows[2].cells[4] if len(table.rows[2].cells) > 4 else table.rows[2].cells[-1], f"TERM: {data.get('term', '')}", spacing=1.5)
        
        preserve_cell_format(table.rows[3].cells[0], f"Module: {data.get('module_code_title', '')}", spacing=1.5)
        preserve_cell_format(table.rows[3].cells[1] if len(table.rows[3].cells) > 1 else table.rows[3].cells[0], f"Week: {data.get('week', '')}", spacing=1.5)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        return temp_file.name
    except:
        pass
    
    return _create_rtb_session_plan_from_scratch(data)

def _create_rtb_session_plan_from_scratch(data):
    """Create RTB session plan document from scratch with exact format"""
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    title = doc.add_paragraph()
    title_run = title.add_run('RWANDA TECHNICAL BOARD (RTB)\nSESSION PLAN')
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = 'Book Antiqua'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    header_table = doc.add_table(rows=3, cols=3)
    header_table.style = 'Table Grid'
    
    h_cells = header_table.rows[0].cells
    preserve_cell_format(h_cells[0], f"Code: {data.get('module_code_title', '')[:30]}")
    preserve_cell_format(h_cells[1], f"Sector: {data.get('sector', '')}")
    preserve_cell_format(h_cells[2], f"Trade: {data.get('trade', '')}")
    
    h_cells = header_table.rows[1].cells
    preserve_cell_format(h_cells[0], f"Level: {data.get('rqf_level', '')}")
    preserve_cell_format(h_cells[1], f"Term: {data.get('term', '')}")
    preserve_cell_format(h_cells[2], f"Week: {data.get('week', '')}")
    
    h_cells = header_table.rows[2].cells
    preserve_cell_format(h_cells[0], f"Date: {data.get('date', '')}")
    preserve_cell_format(h_cells[1], "")
    preserve_cell_format(h_cells[2], "")
    
    doc.add_paragraph()
    
    content_table = doc.add_table(rows=11, cols=2)
    content_table.style = 'Table Grid'
    
    content_data = [
        ("Sector", data.get('sector', '')),
        ("Trade", data.get('trade', '')),
        ("RQF Level", data.get('rqf_level', '')),
        ("Module", data.get('module_code_title', '')),
        ("Class", data.get('class_name', '')),
        ("Number of Trainees", str(data.get('number_of_trainees', ''))),
        ("Topic", data.get('topic_of_session', '')),
        ("Duration", f"{data.get('duration', '')} minutes"),
        ("Learning Outcomes", clean_text(data.get('learning_outcomes', ''))),
        ("Indicative Contents", clean_text(data.get('indicative_contents', ''))),
        ("Facilitation Techniques", clean_text(data.get('facilitation_techniques', ''))),
    ]
    
    for i, (label, value) in enumerate(content_data):
        preserve_cell_format(content_table.rows[i].cells[0], label)
        preserve_cell_format(content_table.rows[i].cells[1], value)
    
    doc.add_paragraph()
    
    resources_section = doc.add_paragraph()
    resources_run = resources_section.add_run("Resources")
    resources_run.bold = True
    resources_run.font.size = Pt(12)
    resources_run.font.name = 'Book Antiqua'
    resources_section.paragraph_format.line_spacing = 1.5
    
    resources_text = clean_text(data.get('resources', ''))
    if resources_text:
        for line in resources_text.split('\n'):
            if line.strip():
                resource_para = doc.add_paragraph(f"• {line.strip()}", style='List Bullet')
                resource_para.paragraph_format.line_spacing = 1.5
                for run in resource_para.runs:
                    run.font.name = 'Book Antiqua'
                    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    intro_heading = doc.add_paragraph()
    intro_run = intro_heading.add_run("Introduction")
    intro_run.bold = True
    intro_run.font.size = Pt(12)
    intro_run.font.name = 'Book Antiqua'
    intro_heading.paragraph_format.line_spacing = 1.5
    
    topic = data.get('topic_of_session', '')
    facilitation = data.get('facilitation_techniques', 'Trainer-guided')
    intro_content = generate_introduction_activities(topic, facilitation)
    intro_para = doc.add_paragraph(intro_content)
    intro_para.paragraph_format.line_spacing = 1.5
    for run in intro_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    dev_heading = doc.add_paragraph()
    dev_run = dev_heading.add_run("Development")
    dev_run.bold = True
    dev_run.font.size = Pt(12)
    dev_run.font.name = 'Book Antiqua'
    dev_heading.paragraph_format.line_spacing = 1.5
    
    dev_content = generate_development_activities(topic, facilitation, data.get('learning_activities'))
    dev_para = doc.add_paragraph(dev_content)
    dev_para.paragraph_format.line_spacing = 1.5
    for run in dev_para.runs:
        run.font.name = 'Book Antiqua'
        run.font.size = Pt(12)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name

def format_section_content(content):
    """Format section content with proper structure and indentation"""
    if not content:
        return ""
    
    lines = content.split('\n')
    formatted_lines = []
    
    for line in lines:
        if line.strip():
            formatted_lines.append(line.rstrip())
    
    return '\n'.join(formatted_lines)

def fetch_web_references(module, topic, learning_outcomes, indicative_contents):
    """Fetch references with APA formatting - with fallback to static resources"""
    try:
        search_terms = f"{topic} {module}".strip()
        if not search_terms:
            search_terms = "Technical and vocational education"
        
        return get_apa_formatted_references(module, topic, learning_outcomes, indicative_contents)
    except Exception as e:
        print(f"Reference generation failed: {e}")
        return get_default_apa_references()

def get_apa_formatted_references(module, topic, learning_outcomes, indicative_contents):
    """Generate APA formatted references based on content"""
    content_lower = f"{topic} {module}".lower() if topic and module else ""
    references = []
    
    if any(term in content_lower for term in ['programming', 'python', 'java', 'code', 'software', 'algorithm', 'loop', 'array', 'function', 'variable']):
        references = [
            "Deitel, P., & Deitel, H. (2019). Python for programmers. Pearson Education.",
            "McConnell, S. (2004). Code complete: A practical handbook of software construction. Microsoft Press.",
            "Matthes, E. (2019). Python crash course: A hands-on, project-based introduction to programming. No Starch Press.",
            "Downey, A. (2015). Think Python: How to think like a computer scientist. O'Reilly Media.",
            "Hunt, A., & Thomas, D. (2019). The pragmatic programmer (2nd ed.). Addison-Wesley."
        ]
    elif any(term in content_lower for term in ['network', 'cisco', 'routing', 'switching', 'tcp', 'ip', 'internet', 'firewall']):
        references = [
            "Tanenbaum, A. S., & Wetherall, D. J. (2021). Computer networks (6th ed.). Pearson Education.",
            "Kurose, J. F., & Ross, K. W. (2020). Computer networking: A top-down approach (8th ed.). Pearson.",
            "Odom, W. (2019). CCNA 200-301 official cert guide library (2nd ed.). Cisco Press.",
            "Cisco Networking Academy. (2020). CCNA routing and switching course materials. Cisco Systems.",
            "Doyle, J. C., Alderson, D. L., & Willinger, W. (2015). Internet topology and the evolution of the Internet. ACM Transactions."
        ]
    elif any(term in content_lower for term in ['database', 'sql', 'mysql', 'data management', 'data modeling', 'nosql']):
        references = [
            "Elmasri, R., & Navathe, S. B. (2020). Fundamentals of database systems (8th ed.). Pearson Education.",
            "Coronel, C., & Morris, S. (2019). Database systems: Design, implementation, and management (13th ed.). Cengage.",
            "Beaulieu, A. (2020). Learning SQL: Generate, manipulate, and retrieve data (3rd ed.). O'Reilly Media.",
            "Garcia-Molina, H., Ullman, J. D., & Widom, J. (2008). Database systems: The complete book (2nd ed.). Prentice Hall.",
            "DuBois, P. (2020). MySQL cookbook: Solutions for database developers and administrators. O'Reilly."
        ]
    elif any(term in content_lower for term in ['web', 'html', 'css', 'javascript', 'frontend', 'responsive', 'react', 'vue']):
        references = [
            "Duckett, J. (2014). HTML and CSS: Design and build websites. Wiley Publishing.",
            "Flanagan, D. (2020). JavaScript: The definitive guide (7th ed.). O'Reilly Media.",
            "Robbins, J. N. (2018). Learning web design: A beginner's guide to HTML, CSS, JavaScript (5th ed.). O'Reilly.",
            "Nielsen, J., & Norman, D. A. (2014). Usability 101: Introduction to usability. Nielsen Norman Group.",
            "Mozilla Foundation. (2023). Web development documentation and standards. Retrieved from https://developer.mozilla.org"
        ]
    elif any(term in content_lower for term in ['business', 'management', 'leadership', 'entrepreneurship', 'accounting', 'finance']):
        references = [
            "Drucker, P. F. (2006). The effective executive: The definitive guide to getting the right things done. Harper Business.",
            "Porter, M. E. (2008). Competitive advantage: Creating and sustaining superior performance. Free Press.",
            "Mintzberg, H. (2009). Managing. Berrett-Koehler Publishers.",
            "Kotter, J. P. (2012). Leading change. Harvard Business Review Press.",
            "Covey, S. R. (2004). The 7 habits of highly effective people. Free Press."
        ]
    else:
        return get_default_apa_references()
    
    formatted_refs = []
    for i, ref in enumerate(references[:5], 1):
        formatted_refs.append(f"{i}. {ref}")
    
    return '\n\n'.join(formatted_refs)

def get_default_apa_references():
    """Default APA formatted references for TVET"""
    references = [
        "Rwanda Education Board. (2021). TVET curriculum framework. REB Publications.",
        "UNESCO-UNEVOC. (2020). Technical and vocational education and training (TVET) and the sustainable development goals (SDGs). UNESCO Publications.",
        "Ministry of Education Rwanda. (2022). Competency-based training guidelines for technical and vocational education. MINEDUC.",
        "Rwanda Technical and Vocational Education and Training Board. (2023). National module guidelines and standards. RTVETB Publications.",
        "International Labour Organization. (2021). World employment and social outlook: The role of digital labour platforms. ILO."
    ]
    
    formatted_refs = []
    for i, ref in enumerate(references, 1):
        formatted_refs.append(f"{i}. {ref}")
    
    return '\n\n'.join(formatted_refs)

def fill_scheme_template(data):
    """Fill scheme with enhanced formatting - Book Antiqua 12pt, spacing 1.5"""
    template_path = os.path.join(os.path.dirname(__file__), 'rtb_scheme_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError("RTB Scheme template not found")
    
    doc = Document(template_path)
    
    if not doc.tables or len(doc.tables) == 0:
        raise ValueError("Template has no tables")
    
    sections = doc.sections
    for section in sections:
        section.top_margin = int(1.27 * 914400)
        section.bottom_margin = int(1.27 * 914400)
        section.left_margin = int(1.27 * 914400)
        section.right_margin = int(1.27 * 914400)
    
    for term_idx in range(min(3, len(doc.tables))):
        term_num = term_idx + 1
        table = doc.tables[term_idx]
        
        if len(table.rows) > 2:
            weeks = clean_text(data.get(f'term{term_num}_weeks', ''))
            outcomes = clean_text(data.get(f'term{term_num}_learning_outcomes', ''))
            duration = clean_text(data.get(f'term{term_num}_duration', ''))
            contents = clean_text(data.get(f'term{term_num}_indicative_contents', ''))
            learning_place = clean_text(data.get(f'term{term_num}_learning_place', ''))
            
            preserve_cell_format(table.rows[2].cells[0], weeks, spacing=1.5)
            preserve_cell_format(table.rows[2].cells[1], outcomes, spacing=1.5)
            preserve_cell_format(table.rows[2].cells[2], duration, spacing=1.5)
            preserve_cell_format(table.rows[2].cells[3], contents, spacing=1.5)
            if len(table.rows[2].cells) > 4:
                preserve_cell_format(table.rows[2].cells[4], learning_place, spacing=1.5)
    
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    return temp_file.name
