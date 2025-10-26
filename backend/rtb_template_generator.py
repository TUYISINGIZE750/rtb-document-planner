from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
import logging

logger = logging.getLogger(__name__)

def preserve_cell_format(cell, new_text, font_name='Book Antiqua', font_size=12, spacing=1.5):
    """Update cell text while preserving RTB formatting"""
    if not cell.paragraphs:
        cell.text = str(new_text) if new_text else ''
        return
    
    cell.paragraphs[0].clear()
    
    text_lines = str(new_text).split('\n') if new_text else ['']
    
    for idx, line in enumerate(text_lines):
        if idx == 0:
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        else:
            para = cell.add_paragraph()
        
        para.paragraph_format.line_spacing = spacing
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        
        if line.strip():
            run = para.add_run(line)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            para.text = ''

def get_template_path(template_name):
    """Get path to RTB template - checks multiple locations"""
    base_dir = os.path.dirname(__file__)
    
    template_path = os.path.join(base_dir, template_name)
    if os.path.exists(template_path):
        logger.info(f"Found template at: {template_path}")
        return template_path
    
    prod_templates_dir = os.path.join(base_dir, '..', 'RTB Templates')
    if template_name == 'rtb_session_plan_template.docx':
        alt_path = os.path.join(prod_templates_dir, 'RTB Session plan template.docx')
    else:
        alt_path = os.path.join(prod_templates_dir, 'Scheme of work.docx')
    
    if os.path.exists(alt_path):
        logger.info(f"Found template at: {alt_path}")
        return alt_path
    
    logger.warning(f"Template not found: {template_name}")
    return None

def generate_session_plan_from_template(data):
    """
    Generate RTB-compliant session plan by filling official template with user data.
    
    Args:
        data (dict): Session plan data including:
            - sector, trade, trainer_name, module_code_title, week, term, date
            - topic_of_session, duration, learning_outcomes, facilitation_techniques
            - indicative_contents, resources, number_of_trainees, class_name, rqf_level
    
    Returns:
        str: Path to generated DOCX file
    """
    template_path = get_template_path('rtb_session_plan_template.docx')
    
    try:
        if template_path and os.path.exists(template_path):
            doc = Document(template_path)
            doc = _fill_session_plan_table(doc, data)
        else:
            logger.error("Session plan template not found")
            raise FileNotFoundError("RTB Session plan template not found")
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating session plan: {e}")
        raise

def _fill_session_plan_table(doc, data):
    """Fill RTB session plan template table with data"""
    
    for section in doc.sections:
        section.top_margin = Cm(1.27)
        section.bottom_margin = Cm(1.27)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)
    
    if not doc.tables or len(doc.tables) == 0:
        logger.error("No tables found in template")
        return doc
    
    table = doc.tables[0]
    
    try:
        if len(table.rows) > 1 and len(table.rows[1].cells) >= 6:
            preserve_cell_format(table.rows[1].cells[0], data.get('sector', ''))
            preserve_cell_format(table.rows[1].cells[1], data.get('trade', ''))
            preserve_cell_format(table.rows[1].cells[4], data.get('date', ''))
        
        if len(table.rows) > 2 and len(table.rows[2].cells) >= 6:
            preserve_cell_format(table.rows[2].cells[0], data.get('trainer_name', ''))
            preserve_cell_format(table.rows[2].cells[4], data.get('term', ''))
        
        if len(table.rows) > 3 and len(table.rows[3].cells) >= 6:
            preserve_cell_format(table.rows[3].cells[0], data.get('module_code_title', ''))
            preserve_cell_format(table.rows[3].cells[1], data.get('week', ''))
            preserve_cell_format(table.rows[3].cells[3], str(data.get('number_of_trainees', '')))
            preserve_cell_format(table.rows[3].cells[4], data.get('class_name', ''))
        
        if len(table.rows) > 4 and len(table.rows[4].cells) >= 2:
            preserve_cell_format(table.rows[4].cells[1], data.get('learning_outcomes', ''))
        
        if len(table.rows) > 5 and len(table.rows[5].cells) >= 2:
            preserve_cell_format(table.rows[5].cells[1], data.get('indicative_contents', ''))
        
        if len(table.rows) > 6 and len(table.rows[6].cells) >= 1:
            preserve_cell_format(table.rows[6].cells[0], data.get('topic_of_session', ''))
        
        if len(table.rows) > 7 and len(table.rows[7].cells) >= 2:
            preserve_cell_format(table.rows[7].cells[0], data.get('rqf_level', ''))
            duration = data.get('duration', '')
            preserve_cell_format(table.rows[7].cells[1], f"{duration} minutes" if duration else '')
        
        if len(table.rows) > 8 and len(table.rows[8].cells) >= 1:
            preserve_cell_format(table.rows[8].cells[0], data.get('learning_outcomes', ''))
        
        if len(table.rows) > 9 and len(table.rows[9].cells) >= 1:
            preserve_cell_format(table.rows[9].cells[0], data.get('facilitation_techniques', ''))
        
        if len(table.rows) > 11 and len(table.rows[11].cells) >= 6:
            resources = data.get('resources', '')
            preserve_cell_format(table.rows[11].cells[2], resources)
            preserve_cell_format(table.rows[11].cells[5], '5 minutes')
        
        for row_num in [13, 14, 15]:
            if len(table.rows) > row_num and len(table.rows[row_num].cells) >= 6:
                resources = data.get('resources', '')
                preserve_cell_format(table.rows[row_num].cells[2], resources)
                preserve_cell_format(table.rows[row_num].cells[5], '25 minutes')
        
        if len(table.rows) > 17 and len(table.rows[17].cells) >= 6:
            preserve_cell_format(table.rows[17].cells[2], data.get('resources', ''))
            preserve_cell_format(table.rows[17].cells[5], '3 minutes')
        
        if len(table.rows) > 18 and len(table.rows[18].cells) >= 6:
            preserve_cell_format(table.rows[18].cells[2], 'Assessment sheets')
            preserve_cell_format(table.rows[18].cells[5], '5 minutes')
        
        if len(table.rows) > 19 and len(table.rows[19].cells) >= 6:
            preserve_cell_format(table.rows[19].cells[2], 'Self-assessment form')
            preserve_cell_format(table.rows[19].cells[5], '2 minutes')
        
        if len(table.rows) > 20 and len(table.rows[20].cells) >= 1:
            preserve_cell_format(table.rows[20].cells[0], 'Appendices: PPT, Task Sheets, Assessment tools, Materials')
        
        if len(table.rows) > 21 and len(table.rows[21].cells) >= 1:
            preserve_cell_format(table.rows[21].cells[0], 'Reflection: Space for trainer reflection')
            
    except Exception as e:
        logger.error(f"Error filling session plan table: {e}")
    
    return doc

def generate_scheme_of_work_from_template(data):
    """
    Generate RTB-compliant scheme of work by filling official template with user data.
    
    Args:
        data (dict): Scheme of work data with term information
    
    Returns:
        str: Path to generated DOCX file
    """
    template_path = get_template_path('rtb_scheme_template.docx')
    
    try:
        if not template_path or not os.path.exists(template_path):
            base_dir = os.path.dirname(__file__)
            alt_template = os.path.join(base_dir, '..', 'RTB Templates', 'Scheme of work.docx')
            if os.path.exists(alt_template):
                template_path = alt_template
            else:
                logger.error("Scheme of work template not found")
                raise FileNotFoundError("RTB Scheme of work template not found")
        
        doc = Document(template_path)
        
        for term_idx in range(min(3, len(doc.tables))):
            term_num = term_idx + 1
            table = doc.tables[term_idx]
            
            term_key = f'term{term_num}'
            weeks = data.get(f'{term_key}_weeks', '')
            outcomes = data.get(f'{term_key}_learning_outcomes', '')
            duration = data.get(f'{term_key}_duration', '')
            contents = data.get(f'{term_key}_indicative_contents', '')
            
            if weeks or outcomes:
                if len(table.rows) < 3:
                    table.add_row()
                
                row = table.rows[2] if len(table.rows) > 2 else table.rows[-1]
                cells = row.cells
                
                if len(cells) > 0:
                    preserve_cell_format(cells[0], weeks)
                if len(cells) > 1:
                    preserve_cell_format(cells[1], outcomes)
                if len(cells) > 2:
                    preserve_cell_format(cells[2], duration)
                if len(cells) > 3:
                    preserve_cell_format(cells[3], contents)
        
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error generating scheme of work: {e}")
        raise
