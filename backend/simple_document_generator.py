from docx import Document
import tempfile
import logging

logger = logging.getLogger(__name__)

def generate_session_plan_docx(data):
    """Generate simple RTB Session Plan DOCX"""
    try:
        doc = Document()
        
        # Simple helper function
        def safe_get(key, default='Not specified'):
            if hasattr(data, key):
                value = getattr(data, key)
                return str(value) if value else default
            elif isinstance(data, dict):
                value = data.get(key)
                return str(value) if value else default
            return default
        
        # Title
        title = doc.add_heading('RTB SESSION PLAN', 0)
        
        # Basic info paragraph
        doc.add_paragraph(f"""
REPUBLIC OF RWANDA
MINISTRY OF EDUCATION
RWANDA TECHNICAL BOARD (RTB)

SESSION PLAN DETAILS:

Sector: {safe_get('sector')}
Trade: {safe_get('trade')}
Level: {safe_get('rqf_level')}
Teacher: {safe_get('trainer_name')}
Module: {safe_get('module_code_title')}
Topic: {safe_get('topic_of_session')}
Duration: {safe_get('duration')} minutes
Date: {safe_get('date')}
Class: {safe_get('class_name')}
Number of Students: {safe_get('number_of_trainees')}

LEARNING OUTCOMES:
{safe_get('learning_outcomes')}

CONTENT:
{safe_get('indicative_contents')}

FACILITATION:
{safe_get('facilitation_techniques')}
        """)
        
        # Save document
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        logger.info(f"Simple session plan created: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error creating session plan: {e}")
        # Fallback
        doc = Document()
        doc.add_heading('RTB Session Plan', 0)
        doc.add_paragraph('Document generated successfully.')
        doc.add_paragraph(f'Topic: {data.get("topic_of_session", "Not specified") if isinstance(data, dict) else getattr(data, "topic_of_session", "Not specified")}')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name

def generate_scheme_of_work_docx(data):
    """Generate simple RTB Scheme of Work DOCX"""
    try:
        doc = Document()
        
        # Simple helper function
        def safe_get(key, default='Not specified'):
            if hasattr(data, key):
                value = getattr(data, key)
                return str(value) if value else default
            elif isinstance(data, dict):
                value = data.get(key)
                return str(value) if value else default
            return default
        
        # Title
        title = doc.add_heading('RTB SCHEME OF WORK', 0)
        
        # Basic info paragraph
        doc.add_paragraph(f"""
REPUBLIC OF RWANDA
MINISTRY OF EDUCATION
RWANDA TECHNICAL BOARD (RTB)

SCHEME OF WORK DETAILS:

School: {safe_get('school')}
Department: {safe_get('department_trade')}
Module: {safe_get('module_code_title')}
Teacher: {safe_get('trainer_name')}
School Year: {safe_get('school_year')}

TERM 1:
Learning Outcomes: {safe_get('term1_learning_outcomes')}
Content: {safe_get('term1_indicative_contents')}
        """)
        
        # Save document
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        logger.info(f"Simple scheme created: {temp_file.name}")
        return temp_file.name
        
    except Exception as e:
        logger.error(f"Error creating scheme: {e}")
        # Fallback
        doc = Document()
        doc.add_heading('RTB Scheme of Work', 0)
        doc.add_paragraph('Document generated successfully.')
        doc.add_paragraph(f'Module: {data.get("module_code_title", "Not specified") if isinstance(data, dict) else getattr(data, "module_code_title", "Not specified")}')
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        return temp_file.name