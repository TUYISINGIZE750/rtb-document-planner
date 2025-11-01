#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""Quick test of document generation"""

import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from rtb_template_generator import generate_session_plan_from_template

# Test data
test_data = {
    'sector': 'ICT and Multimedia',
    'trade': 'Computer system and architecture',
    'trainer_name': 'John Doe',
    'module_code_title': 'Introduction to PLC',
    'week': 'Week 11',
    'term': 'Term 2',
    'date': '2025-10-26',
    'class_name': 'L4 CSA',
    'number_of_trainees': '35',
    'topic_of_session': 'Do while loops in C',
    'duration': '40',
    'learning_outcomes': 'Understand Do while loops',
    'indicative_contents': 'Syntax and examples',
    'facilitation_techniques': 'Simulation',
    'resources': 'Projector, Whiteboard',
    'rqf_level': 'Level 3'
}

print("Testing Session Plan Generation...")
print("=" * 80)

try:
    # Generate document
    gen_file = generate_session_plan_from_template(test_data)
    print(f"✓ Document generated: {gen_file}")
    print(f"✓ File exists: {os.path.exists(gen_file)}")
    print(f"✓ File size: {os.path.getsize(gen_file)} bytes")
    
    # Load generated document
    doc = Document(gen_file)
    
    # Check table
    if len(doc.tables) > 0:
        table = doc.tables[0]
        print(f"\n✓ Table found")
        print(f"  - Rows: {len(table.rows)}")
        print(f"  - Columns: {len(table.rows[0].cells)}")
        
        # Check specific cells for data
        print(f"\nCell Contents Check:")
        print(f"  Row 1, Col 0 (Sector): {table.rows[1].cells[0].text.strip()}")
        print(f"  Row 1, Col 1 (Trade): {table.rows[1].cells[1].text.strip()}")
        print(f"  Row 1, Col 4 (Date): {table.rows[1].cells[4].text.strip()}")
        print(f"  Row 2, Col 0 (Trainer): {table.rows[2].cells[0].text.strip()}")
        print(f"  Row 3, Col 0 (Module): {table.rows[3].cells[0].text.strip()}")
        print(f"  Row 3, Col 1 (Week): {table.rows[3].cells[1].text.strip()}")
        
        # Check margins
        print(f"\nFormatting Check:")
        for section in doc.sections:
            print(f"  - Top margin: {section.top_margin.cm:.2f} cm")
            print(f"  - Left margin: {section.left_margin.cm:.2f} cm")
        
        print("\n✓ Document generation appears to be working correctly!")
        print(f"\nTest file saved to: {gen_file}")
        
    else:
        print("✗ No tables found in document!")
        
except Exception as e:
    print(f"✗ ERROR: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)

print("\n" + "=" * 80)
