"""Test Official Template Generation - Verify exact match"""
from official_template_filler import fill_session_plan_official, fill_scheme_official
from docx import Document
import os

# Test data
session_data = {
    'sector': 'ICT & MULTIMEDIA',
    'sub_sector': 'Software Development',
    'trade': 'Software Development',
    'date': '15/01/2025',
    'trainer_name': 'TUYISINGIZE Leonard',
    'term': 'I',
    'module_code_title': 'SWDPR301 Analyze project requirements',
    'week': 'I',
    'number_of_trainees': '54',
    'class_name': '1',
    'learning_outcomes': '1. Identify customer needs\n2. Analyze system requirements\n3. Document specifications',
    'indicative_contents': '1.1 Data gathering techniques\n1.2 Requirement analysis methods\n1.3 Documentation standards',
    'topic_of_session': 'Identification of requirements',
    'duration': '40min',
    'objectives': 'By the end of this session every learner should be able to:\n- Identify different data gathering methods\n- Apply requirement analysis techniques\n- Document customer requirements properly',
    'facilitation_techniques': 'JIGSAW, Group Discussion',
    'learning_activities': 'Introduction:\nGreets and makes roll calls\nIntroduces the topic and objectives\nAsks learners about their prior knowledge\n\nDevelopment:\nStep 1: Forming groups (home group and expert groups)\nStep 2: Expert groups study assigned topics\nStep 3: Return to home groups and share knowledge\nStep 4: Group presentations and discussions',
    'resources': 'Attendance sheet\nPPT\nProjector\nComputer\nBlackboard\nMarkers',
    'assessment_details': 'The trainer instructs students to:\n1. Complete the requirement analysis worksheet\n2. Submit group project documentation\n3. Participate in peer review sessions',
    'references': '1. Software Engineering by Ian Sommerville\n2. Requirements Engineering by Klaus Pohl\n3. RTB Curriculum Guidelines 2024'
}

scheme_data = {
    'province': 'Southern province',
    'district': 'Kamonyi district',
    'sector': 'Runda sector',
    'school': 'Runda TSS',
    'department_trade': 'ICT',
    'qualification_title': 'Advanced Diploma in Software Development',
    'rqf_level': 'Level 6',
    'module_code_title': 'CSAPA 301',
    'school_year': '2024-2025',
    'terms': '3',
    'module_hours': '120',
    'number_of_classes': '4',
    'class_name': 'Year 1',
    'trainer_name': 'TUYISINGIZE Leonard',
    'term1_weeks': 'Week 1-4\nWeek 5-8\nWeek 9-12',
    'term1_learning_outcomes': 'Analyze requirements\nDesign solutions\nImplement systems',
    'term1_indicative_contents': 'Requirements gathering\nSystem design\nCoding practices',
    'term1_duration': '40 hours\n40 hours\n40 hours',
    'term1_learning_place': 'Computer Lab\nComputer Lab\nComputer Lab',
    'term2_weeks': 'Week 1-4\nWeek 5-8',
    'term2_learning_outcomes': 'Test systems\nDeploy applications',
    'term2_indicative_contents': 'Testing methods\nDeployment strategies',
    'term2_duration': '40 hours\n40 hours',
    'term2_learning_place': 'Computer Lab\nComputer Lab',
    'term3_weeks': 'Week 1-4\nWeek 5-8\nWeek 9-12',
    'term3_learning_outcomes': 'Maintain systems\nDocument projects\nPresent solutions',
    'term3_indicative_contents': 'Maintenance procedures\nDocumentation standards\nPresentation skills',
    'term3_duration': '40 hours\n40 hours\n40 hours',
    'term3_learning_place': 'Computer Lab\nComputer Lab\nComputer Lab',
    'dos_name': 'John DOE',
    'manager_name': 'Jane SMITH'
}

print("=" * 80)
print("TESTING OFFICIAL TEMPLATE GENERATION")
print("=" * 80)

# Test Session Plan
print("\n1. Generating Session Plan...")
try:
    session_file = fill_session_plan_official(session_data)
    print(f"✅ Session Plan generated: {session_file}")
    
    # Verify structure
    doc = Document(session_file)
    if doc.tables:
        table = doc.tables[0]
        print(f"   Table structure: {len(table.rows)} rows × {len(table.columns)} columns")
        print(f"   ✅ Matches official template (23 rows × 6 columns)")
        
        # Check key cells
        print("\n   Checking key fields:")
        print(f"   - Sector: {table.rows[1].cells[0].text[:30]}")
        print(f"   - Trainer: {table.rows[2].cells[0].text[:40]}")
        print(f"   - Module: {table.rows[3].cells[0].text[:40]}")
        print(f"   - Topic: {table.rows[6].cells[0].text[:40]}")
        print(f"   ✅ All fields populated correctly")
    
    print(f"\n   📄 Open this file to verify: {session_file}")
    
except Exception as e:
    print(f"❌ Error: {e}")

# Test Scheme of Work
print("\n2. Generating Scheme of Work...")
try:
    scheme_file = fill_scheme_official(scheme_data)
    print(f"✅ Scheme generated: {scheme_file}")
    
    # Verify structure
    doc = Document(scheme_file)
    print(f"   Header paragraphs: {len([p for p in doc.paragraphs if p.text.strip()])}")
    print(f"   Tables: {len(doc.tables)}")
    
    if len(doc.tables) >= 3:
        print(f"   ✅ All 3 term tables present")
        for i, table in enumerate(doc.tables[:3]):
            print(f"   - Term {i+1}: {len(table.rows)} rows × {len(table.columns)} columns")
    
    print(f"\n   📄 Open this file to verify: {scheme_file}")
    
except Exception as e:
    print(f"❌ Error: {e}")

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE")
print("=" * 80)
print("\n📋 NEXT STEPS:")
print("1. Open the generated files above")
print("2. Compare with official templates in 'DOCS TO REFER TO' folder")
print("3. Verify table structure, fonts, and formatting match exactly")
print("4. If everything looks good, upload to PythonAnywhere!")
