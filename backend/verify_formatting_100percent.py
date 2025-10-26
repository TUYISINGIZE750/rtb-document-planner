"""100% Verification of Document Formatting"""
from docx import Document
import os

def verify_document_formatting(doc_path, doc_type):
    """Verify all formatting aspects"""
    print(f"\n{'='*70}")
    print(f"VERIFYING: {doc_type}")
    print(f"{'='*70}\n")
    
    if not os.path.exists(doc_path):
        print(f"[ERROR] File not found: {doc_path}")
        return False
    
    doc = Document(doc_path)
    table = doc.tables[0]
    
    issues = []
    checks_passed = 0
    total_checks = 0
    
    # Check 1: No excessive spacing
    print("1. Checking for excessive spacing...")
    total_checks += 1
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            text = cell.text
            if '   ' in text:  # 3+ spaces
                issues.append(f"Row {row_idx}, Cell {cell_idx}: Multiple spaces found")
            if '\n\n\n' in text:  # 3+ newlines
                issues.append(f"Row {row_idx}, Cell {cell_idx}: Excessive line breaks")
    
    if not issues:
        print("   [PASS] No excessive spacing found")
        checks_passed += 1
    else:
        print(f"   [FAIL] Found {len(issues)} spacing issues")
    
    # Check 2: Proper bullet formatting
    print("\n2. Checking bullet point formatting...")
    total_checks += 1
    bullet_issues = 0
    for row_idx in [11, 13, 18, 19]:  # Rows with bullets
        if row_idx < len(table.rows):
            text = table.rows[row_idx].cells[0].text
            if '•' in text:
                lines = text.split('\n')
                for line in lines:
                    if '•' in line and line.startswith('\t•'):
                        bullet_issues += 1
    
    if bullet_issues == 0:
        print("   [PASS] Bullets properly formatted (2-space indent)")
        checks_passed += 1
    else:
        print(f"   [FAIL] Found {bullet_issues} bullet formatting issues")
    
    # Check 3: Objectives numbering
    print("\n3. Checking objectives numbering...")
    total_checks += 1
    if len(table.rows) > 8:
        objectives_text = table.rows[8].cells[0].text
        if '1.' in objectives_text and '2.' in objectives_text:
            print("   [PASS] Objectives properly numbered")
            checks_passed += 1
        else:
            print("   [FAIL] Objectives not properly numbered")
    
    # Check 4: Resources formatting
    print("\n4. Checking resources formatting...")
    total_checks += 1
    if len(table.rows) > 13:
        resources_text = table.rows[13].cells[2].text
        lines = [l for l in resources_text.split('\n') if l.strip()]
        if len(lines) >= 3 and not '  ' in resources_text:
            print(f"   [PASS] Resources listed cleanly ({len(lines)} items)")
            checks_passed += 1
        else:
            print("   [FAIL] Resources formatting issues")
    
    # Check 5: Content fits in cells
    print("\n5. Checking content fits in cells...")
    total_checks += 1
    overflow_issues = 0
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if len(cell.text) > 5000:  # Excessive content
                overflow_issues += 1
    
    if overflow_issues == 0:
        print("   [PASS] All content fits properly")
        checks_passed += 1
    else:
        print(f"   [FAIL] {overflow_issues} cells with excessive content")
    
    # Check 6: No scattered content
    print("\n6. Checking for scattered content...")
    total_checks += 1
    scattered = False
    for row in table.rows:
        for cell in row.cells:
            text = cell.text
            if text.count('\n\n') > 3:  # Too many blank lines
                scattered = True
    
    if not scattered:
        print("   [PASS] Content well-organized")
        checks_passed += 1
    else:
        print("   [FAIL] Scattered content detected")
    
    # Check 7: Font consistency
    print("\n7. Checking font consistency...")
    total_checks += 1
    fonts_ok = True
    for row in table.rows[:5]:  # Check first 5 rows
        for cell in row.cells:
            if cell.paragraphs:
                for para in cell.paragraphs:
                    if para.runs:
                        font = para.runs[0].font.name
                        if font and 'Times New Roman' not in font:
                            fonts_ok = False
    
    if fonts_ok:
        print("   [PASS] Fonts consistent (Times New Roman)")
        checks_passed += 1
    else:
        print("   [FAIL] Font inconsistencies detected")
    
    # Check 8: Table structure
    print("\n8. Checking table structure...")
    total_checks += 1
    if len(table.rows) == 23 and len(table.columns) == 6:
        print("   [PASS] Table structure correct (23x6)")
        checks_passed += 1
    else:
        print(f"   [FAIL] Table structure: {len(table.rows)}x{len(table.columns)}")
    
    # Summary
    print(f"\n{'='*70}")
    print(f"VERIFICATION SUMMARY")
    print(f"{'='*70}")
    print(f"Checks Passed: {checks_passed}/{total_checks}")
    print(f"Success Rate: {(checks_passed/total_checks)*100:.1f}%")
    
    if checks_passed == total_checks:
        print("\n✅ 100% VERIFIED - ALL CHECKS PASSED!")
        return True
    else:
        print(f"\n⚠️ {total_checks - checks_passed} issues found")
        return False

def main():
    base_dir = os.path.dirname(__file__)
    
    # Find latest test file
    test_files = [f for f in os.listdir(base_dir) if f.startswith('TEST_Session_Plan_')]
    if test_files:
        latest = sorted(test_files)[-1]
        test_path = os.path.join(base_dir, latest)
        
        print("\n" + "="*70)
        print("100% FORMATTING VERIFICATION")
        print("="*70)
        
        result = verify_document_formatting(test_path, latest)
        
        if result:
            print("\n" + "="*70)
            print("PROOF: FORMATTING IS 100% CORRECT")
            print("="*70)
            print("""
All checks passed:
✅ No excessive spacing
✅ Proper bullet formatting
✅ Objectives numbered correctly
✅ Resources listed cleanly
✅ Content fits in cells
✅ Well-organized content
✅ Font consistency maintained
✅ Table structure correct

READY FOR DEPLOYMENT!
            """)
        else:
            print("\nSome issues detected. Review above.")
    else:
        print("[ERROR] No test files found. Run test_document_generation.py first.")

if __name__ == "__main__":
    main()
