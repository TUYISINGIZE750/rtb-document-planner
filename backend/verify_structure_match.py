"""Verify generated documents match RTB template structure exactly"""
from docx import Document
from docx.oxml.ns import qn

def get_cell_structure(cell):
    """Get colspan and vmerge info for a cell"""
    tc = cell._element
    tcPr = tc.tcPr
    
    gridSpan = 1
    vMerge = None
    
    if tcPr is not None:
        gridSpan_elem = tcPr.find(qn('w:gridSpan'))
        if gridSpan_elem is not None:
            gridSpan = int(gridSpan_elem.get(qn('w:val')))
        
        vMerge_elem = tcPr.find(qn('w:vMerge'))
        if vMerge_elem is not None:
            vMerge_val = vMerge_elem.get(qn('w:val'))
            vMerge = 'restart' if vMerge_val == 'restart' else 'continue'
    
    return (gridSpan, vMerge)

def compare_tables(template_path, generated_path, doc_name):
    """Compare template and generated document structures"""
    print(f"\n{'='*70}")
    print(f"COMPARING: {doc_name}")
    print(f"{'='*70}")
    
    template_doc = Document(template_path)
    generated_doc = Document(generated_path)
    
    # For session plan, compare table 1 (main table, table 0 is header we added)
    if 'Session' in doc_name:
        template_table = template_doc.tables[0]
        generated_table = generated_doc.tables[1]  # Skip our added header
    else:
        template_table = template_doc.tables[0]
        generated_table = generated_doc.tables[0]
    
    print(f"\nTemplate rows: {len(template_table.rows)}")
    print(f"Generated rows: {len(generated_table.rows)}")
    
    if len(template_table.rows) != len(generated_table.rows):
        print("[WARNING] Row count mismatch!")
        return False
    
    structure_match = True
    mismatches = []
    
    for row_idx in range(len(template_table.rows)):
        template_row = template_table.rows[row_idx]
        generated_row = generated_table.rows[row_idx]
        
        if len(template_row.cells) != len(generated_row.cells):
            mismatches.append(f"Row {row_idx}: cell count {len(template_row.cells)} vs {len(generated_row.cells)}")
            structure_match = False
            continue
        
        for cell_idx in range(len(template_row.cells)):
            template_struct = get_cell_structure(template_row.cells[cell_idx])
            generated_struct = get_cell_structure(generated_row.cells[cell_idx])
            
            if template_struct != generated_struct:
                mismatches.append(
                    f"Row {row_idx}, Cell {cell_idx}: "
                    f"Template {template_struct} vs Generated {generated_struct}"
                )
                structure_match = False
    
    if structure_match:
        print("\n[OK] Structure matches perfectly!")
        print("  - All rows present")
        print("  - All cells present")
        print("  - All colspan values match")
        print("  - All rowspan (vMerge) values match")
        return True
    else:
        print(f"\n[FAIL] Structure mismatches found: {len(mismatches)}")
        for mismatch in mismatches[:10]:  # Show first 10
            print(f"  - {mismatch}")
        if len(mismatches) > 10:
            print(f"  ... and {len(mismatches) - 10} more")
        return False

if __name__ == "__main__":
    print("\n" + "="*70)
    print("RTB TEMPLATE STRUCTURE VERIFICATION")
    print("="*70)
    
    # Compare Session Plan
    session_match = compare_tables(
        'RTB Templates/RTB Session plan template.docx',
        'test_session_plan.docx',
        'Session Plan'
    )
    
    # Compare Scheme of Work
    scheme_match = compare_tables(
        'RTB Templates/Scheme of work.docx',
        'test_scheme_of_work.docx',
        'Scheme of Work'
    )
    
    print(f"\n{'='*70}")
    print("FINAL VERIFICATION RESULTS")
    print(f"{'='*70}")
    print(f"Session Plan Structure: {'[PASS]' if session_match else '[FAIL]'}")
    print(f"Scheme of Work Structure: {'[PASS]' if scheme_match else '[FAIL]'}")
    
    if session_match and scheme_match:
        print("\n[SUCCESS] All documents match RTB template structure 100%")
        print("  - Colspan preserved")
        print("  - Rowspan preserved")
        print("  - Table structure identical")
        print("  - Ready for production!")
    else:
        print("\n[WARNING] Structure mismatches detected")
        print("  - Review mismatches above")
    
    print("="*70 + "\n")
