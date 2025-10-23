from docx import Document
from docx.oxml.ns import qn
import sys

def analyze_rtb_template(file_path, output_file):
    """Deep analysis of RTB template structure"""
    doc = Document(file_path)
    
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(f"ANALYZING: {file_path}\n")
        f.write("="*80 + "\n\n")
        
        # Analyze tables
        for table_idx, table in enumerate(doc.tables):
            f.write(f"\nTABLE {table_idx + 1}\n")
            f.write(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}\n")
            f.write("-"*80 + "\n")
            
            for row_idx, row in enumerate(table.rows):
                f.write(f"\nRow {row_idx + 1}:\n")
                
                for cell_idx, cell in enumerate(row.cells):
                    tc = cell._element
                    tcPr = tc.tcPr
                    
                    # Get cell text
                    text = cell.text.strip()[:60]
                    f.write(f"  Cell {cell_idx + 1}: '{text}'\n")
                    
                    if tcPr is not None:
                        # Check gridSpan (colspan)
                        gridSpan = tcPr.find(qn('w:gridSpan'))
                        if gridSpan is not None:
                            span_val = gridSpan.get(qn('w:val'))
                            f.write(f"    COLSPAN: {span_val}\n")
                        
                        # Check vMerge (rowspan)
                        vMerge = tcPr.find(qn('w:vMerge'))
                        if vMerge is not None:
                            merge_val = vMerge.get(qn('w:val'))
                            if merge_val:
                                f.write(f"    ROWSPAN START\n")
                            else:
                                f.write(f"    ROWSPAN CONTINUE\n")
                        
                        # Check shading (background color)
                        shd = tcPr.find(qn('w:shd'))
                        if shd is not None:
                            fill = shd.get(qn('w:fill'))
                            f.write(f"    BACKGROUND: {fill}\n")
                        
                        # Check borders
                        tcBorders = tcPr.find(qn('w:tcBorders'))
                        if tcBorders is not None:
                            f.write(f"    HAS BORDERS\n")
                    
                    # Check text formatting
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.bold:
                                f.write(f"    TEXT: BOLD\n")
                            if run.font.size:
                                f.write(f"    FONT SIZE: {run.font.size.pt}pt\n")
                            if run.font.name:
                                f.write(f"    FONT: {run.font.name}\n")
                            break
                        break

# Analyze both templates
print("Analyzing Session Plan template...")
analyze_rtb_template(
    r'RTB Templates\TUYISINGIZE LEONARD_SESSION PLAN_S4F3.docx',
    'session_plan_analysis.txt'
)

print("Analyzing Scheme of Work template...")
analyze_rtb_template(
    r'RTB Templates\CSAPA 301 Scheme of work.docx',
    'scheme_analysis.txt'
)

print("\nAnalysis complete! Check:")
print("- session_plan_analysis.txt")
print("- scheme_analysis.txt")
