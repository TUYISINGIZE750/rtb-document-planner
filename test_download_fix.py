#!/usr/bin/env python3
"""
Test script to verify download fixes work correctly.
Run locally to test before deploying to PythonAnywhere.
"""

import sys
import os
import tempfile
import logging
from pathlib import Path

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Test data
test_data_session = {
    'sector': 'Agriculture',
    'sub_sector': 'Crop Production',
    'trade': 'Coffee Farming',
    'qualification_title': 'Diploma in Coffee Production',
    'rqf_level': 'Level 4',
    'module_code_title': 'AGRO-401: Advanced Coffee Production',
    'term': '1',
    'week': '5',
    'date': '2024-10-25',
    'trainer_name': 'Mr. Jean Mukiza',
    'class_name': 'AGR-4A',
    'number_of_trainees': '35',
    'learning_outcomes': 'Understand coffee production techniques',
    'indicative_contents': 'Soil preparation, planting, irrigation',
    'topic_of_session': 'Coffee Bean Harvesting Techniques',
    'duration': '45',
    'objectives': 'S.M.A.R.T objectives for the session',
    'facilitation_techniques': 'Hands-on practical demonstration',
    'learning_activities': 'Group work and practical exercises',
    'resources': 'Coffee plants, tools, notebooks',
    'assessment_details': 'Practical assessment and observation',
    'references': 'RTB curriculum guidelines'
}

test_data_scheme = {
    'province': 'Southern Province',
    'district': 'Huye',
    'sector': 'Agriculture',
    'school': 'Technical School Huye',
    'department_trade': 'Coffee Production',
    'qualification_title': 'Diploma in Coffee Production',
    'rqf_level': 'Level 4',
    'module_code_title': 'AGRO-401: Advanced Coffee Production',
    'school_year': '2024/2025',
    'terms': '3',
    'module_hours': '180',
    'number_of_classes': '3',
    'trainer_name': 'Mr. Jean Mukiza',
    'dos_name': 'Dr. Marie Habimana',
    'term1_weeks': '15',
    'term1_learning_outcomes': 'Master coffee cultivation',
    'term1_indicative_contents': 'Soil, irrigation, pest management',
    'term1_duration': '60 hours',
    'term1_learning_place': 'Farm and classroom',
    'term2_weeks': '15',
    'term2_learning_outcomes': 'Coffee processing techniques',
    'term2_indicative_contents': 'Harvesting, processing, quality control',
    'term2_duration': '60 hours',
    'term2_learning_place': 'Processing facility',
    'term3_weeks': '15',
    'term3_learning_outcomes': 'Coffee marketing and business',
    'term3_indicative_contents': 'Market analysis, branding, sales',
    'term3_duration': '60 hours',
    'term3_learning_place': 'Market and classroom'
}

def test_header_paragraph_fix():
    """Test that header paragraphs are handled correctly"""
    logger.info("Testing header paragraph fix...")
    try:
        from docx import Document
        
        doc = Document()
        header = doc.sections[0].header
        
        # This should not fail even if paragraphs is empty
        if header.paragraphs:
            header_para = header.paragraphs[0]
        else:
            header_para = header.add_paragraph()
        
        header_para.text = "Test Header"
        
        logger.info("✓ Header paragraph fix works correctly")
        return True
    except Exception as e:
        logger.error(f"✗ Header paragraph test failed: {e}")
        return False

def test_session_plan_generation():
    """Test session plan document generation with fallback"""
    logger.info("Testing session plan generation...")
    try:
        # Add backend to path
        backend_path = Path(__file__).parent / "backend"
        sys.path.insert(0, str(backend_path))
        
        from document_generator import generate_session_plan_docx
        
        file_path = generate_session_plan_docx(test_data_session)
        
        if file_path and os.path.exists(file_path):
            size = os.path.getsize(file_path)
            logger.info(f"✓ Session plan generated successfully: {file_path} ({size} bytes)")
            os.remove(file_path)  # Cleanup
            return True
        else:
            logger.error("✗ Session plan generation failed: no file created")
            return False
            
    except Exception as e:
        logger.error(f"✗ Session plan test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_scheme_generation():
    """Test scheme of work document generation with fallback"""
    logger.info("Testing scheme of work generation...")
    try:
        # Add backend to path
        backend_path = Path(__file__).parent / "backend"
        sys.path.insert(0, str(backend_path))
        
        from document_generator import generate_scheme_of_work_docx
        
        file_path = generate_scheme_of_work_docx(test_data_scheme)
        
        if file_path and os.path.exists(file_path):
            size = os.path.getsize(file_path)
            logger.info(f"✓ Scheme generated successfully: {file_path} ({size} bytes)")
            os.remove(file_path)  # Cleanup
            return True
        else:
            logger.error("✗ Scheme generation failed: no file created")
            return False
            
    except Exception as e:
        logger.error(f"✗ Scheme test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def test_error_handling():
    """Test that errors are properly caught and logged"""
    logger.info("Testing error handling...")
    try:
        # Add backend to path
        backend_path = Path(__file__).parent / "backend"
        sys.path.insert(0, str(backend_path))
        
        from document_generator import generate_session_plan_docx
        
        # Test with empty data
        empty_data = {}
        file_path = generate_session_plan_docx(empty_data)
        
        if file_path and os.path.exists(file_path):
            size = os.path.getsize(file_path)
            logger.info(f"✓ Empty data handling works: {file_path} ({size} bytes)")
            os.remove(file_path)
            return True
        else:
            logger.error("✗ Empty data handling failed")
            return False
            
    except Exception as e:
        logger.error(f"✗ Error handling test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

def run_all_tests():
    """Run all tests"""
    logger.info("=" * 60)
    logger.info("Starting Download Fix Tests")
    logger.info("=" * 60)
    
    tests = [
        ("Header Paragraph Fix", test_header_paragraph_fix),
        ("Error Handling", test_error_handling),
        ("Session Plan Generation", test_session_plan_generation),
        ("Scheme Generation", test_scheme_generation),
    ]
    
    results = []
    for test_name, test_func in tests:
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            logger.error(f"Unexpected error in {test_name}: {e}")
            results.append((test_name, False))
        logger.info("")
    
    # Summary
    logger.info("=" * 60)
    logger.info("TEST SUMMARY")
    logger.info("=" * 60)
    passed = sum(1 for _, result in results if result)
    total = len(results)
    
    for test_name, result in results:
        status = "✓ PASS" if result else "✗ FAIL"
        logger.info(f"{status}: {test_name}")
    
    logger.info("-" * 60)
    logger.info(f"Results: {passed}/{total} tests passed")
    logger.info("=" * 60)
    
    return passed == total

if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
