import os
import sys
from docx import Document

# Official templates
official_session = r"C:\Users\PC\Music\Scheme of work and session plan planner\RTB Templates\TUYISINGIZE LEONARD_SESSION PLAN_S4F3.docx"
official_scheme = r"C:\Users\PC\Music\Scheme of work and session plan planner\RTB Templates\CSAPA 301 Scheme of work.docx"

# Backend templates
backend_session = r"C:\Users\PC\Music\Scheme of work and session plan planner\PRODUCTION_READY\backend\rtb_session_plan_template.docx"
backend_scheme = r"C:\Users\PC\Music\Scheme of work and session plan planner\PRODUCTION_READY\backend\rtb_scheme_template.docx"

print("=" * 70)
print("TEMPLATE VERIFICATION")
print("=" * 70)

# Official templates
print("\n1. OFFICIAL TEMPLATES IN RTB Templates/")
print("-" * 70)
if os.path.exists(official_session):
    doc = Document(official_session)
    print(f"✓ Session Plan template found")
    print(f"  - Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"  - Table {i}: {len(table.rows)} rows × {len(table.columns)} columns")
else:
    print(f"✗ Session Plan template NOT found")

if os.path.exists(official_scheme):
    doc = Document(official_scheme)
    print(f"✓ Scheme of Work template found")
    print(f"  - Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"  - Table {i}: {len(table.rows)} rows × {len(table.columns)} columns")
else:
    print(f"✗ Scheme of Work template NOT found")

# Backend templates
print("\n2. BACKEND TEMPLATES IN PRODUCTION_READY/backend/")
print("-" * 70)
if os.path.exists(backend_session):
    doc = Document(backend_session)
    print(f"✓ Session Plan template found")
    print(f"  - Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"  - Table {i}: {len(table.rows)} rows × {len(table.columns)} columns")
else:
    print(f"✗ Session Plan template NOT found")

if os.path.exists(backend_scheme):
    doc = Document(backend_scheme)
    print(f"✓ Scheme of Work template found")
    print(f"  - Tables: {len(doc.tables)}")
    for i, table in enumerate(doc.tables):
        print(f"  - Table {i}: {len(table.rows)} rows × {len(table.columns)} columns")
else:
    print(f"✗ Scheme of Work template NOT found")

# Compare file sizes
print("\n3. FILE COMPARISON")
print("-" * 70)
if os.path.exists(official_session) and os.path.exists(backend_session):
    size1 = os.path.getsize(official_session)
    size2 = os.path.getsize(backend_session)
    print(f"Session Plan:")
    print(f"  - Official: {size1} bytes")
    print(f"  - Backend:  {size2} bytes")
    if size1 == size2:
        print(f"  - Status: IDENTICAL")
    else:
        print(f"  - Status: DIFFERENT - Need to update backend template")

if os.path.exists(official_scheme) and os.path.exists(backend_scheme):
    size1 = os.path.getsize(official_scheme)
    size2 = os.path.getsize(backend_scheme)
    print(f"\nScheme of Work:")
    print(f"  - Official: {size1} bytes")
    print(f"  - Backend:  {size2} bytes")
    if size1 == size2:
        print(f"  - Status: IDENTICAL")
    else:
        print(f"  - Status: DIFFERENT - Need to update backend template")
